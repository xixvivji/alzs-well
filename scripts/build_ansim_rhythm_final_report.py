from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import build_ansim_rhythm_report as master


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "Ansim_Rhythm_2026_Financial_AI_Challenge_Final_Report.docx"


def add_final_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = master.Pt(16)
    p.paragraph_format.space_after = master.Pt(8)
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(p.add_run("2026 금융 AI Challenge"), size=11, bold=True, color=master.TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = master.Pt(54)
    p.paragraph_format.space_after = master.Pt(6)
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(p.add_run("안심리듬"), size=40, bold=True, color=master.NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = master.Pt(8)
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(p.add_run("ANSIM RHYTHM"), size=16, bold=True, color=master.BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = master.Pt(24)
    p.paragraph_format.line_spacing = 1.30
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(
        p.add_run("금융생활 변화 조기알림 및 보호업무 코파일럿\n공모전 전략·기능명세·MVP 구현 최종 통합본"),
        size=15,
        bold=True,
        color=master.INK,
    )

    master.base.callout(
        doc,
        "CORE DEMO",
        "같은 경보, 다른 맥락, 다른 다음 행동",
        fill=master.PALE_GOLD,
        accent=master.GOLD,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = master.Pt(26)
    p.paragraph_format.space_after = master.Pt(4)
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(
        p.add_run("질병을 예측하지 않고, 설명이 필요한 금융변화만 사람의 보호업무로 연결한다."),
        size=11,
        bold=True,
        color=master.NAVY,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = master.Pt(58)
    p.paragraph_format.space_after = master.Pt(3)
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(p.add_run("기준일  2026. 8. 13."), size=10, color=master.GRAY)
    p = doc.add_paragraph()
    p.alignment = master.WD_ALIGN_PARAGRAPH.CENTER
    master.base.set_font(
        p.add_run("팀 공유·공모전 제출 준비용 최종본 | 사실·해석·팀 설계·내부 검증지표를 구분"),
        size=8.7,
        color=master.GRAY,
    )


def add_final_exec_summary(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "Executive Summary",
        "안심리듬 한눈에 보기",
        "대회의 요구를 하나의 사용자 여정으로 묶는 인간 중심 금융안전 코파일럿",
    )
    master.base.callout(
        doc,
        "최종 정의",
        "안심리듬은 치매를 예측하지 않는다. 개인별 금융생활 변화를 발견하고, 생활맥락을 확인해 정상 변화는 해제하며, 설명이 필요한 사건만 행원의 보호업무로 연결한다.",
        fill=master.SKY,
        accent=master.BLUE,
    )
    doc.add_heading("대회 적합성", level=2)
    master.base.table(
        doc,
        ["대회가 요구하는 축", "안심리듬의 구현"],
        [
            ["고령층·취약소비자 포용금융", "낙인 없는 쉬운 설명, 본인 중심 동의, 생활맥락 확인"],
            ["이상금융거래 탐지", "개인별 기준선·MAD·추세·변화점·사유코드 기반 변화 탐지"],
            ["맞춤형 행동요령", "본인 확인→행원 검토→검증된 보호수단 안내의 단계적 흐름"],
            ["임직원 업무 효율화", "사건 묶음, 중립 질문, 근거, 상담기록·후속관리 초안 제공"],
        ],
        [3400, 5960],
        header_fill=master.NAVY,
    )
    doc.add_heading("수상 전략의 중심", level=2)
    master.add_bullet(doc, bullet_id, "같은 alertId·동일 거래·동일 사전판정을 유지하고, 생활맥락 응답만 바꾸어 다음 행동이 갈리는 장면을 시연한다.")
    master.add_bullet(doc, bullet_id, "AI는 점수·차단·상품가입을 결정하지 않는다. 탐지는 규칙·통계 엔진, 보호조치는 정책엔진, 최종결정은 고객과 행원이 담당한다.")
    master.add_bullet(doc, bullet_id, "로그인·API 키 없이 90초 안에 완주하고, LLM이 없어도 결정론적 템플릿으로 동일한 결과를 설명한다.")
    master.add_bullet(doc, bullet_id, "합성데이터로 개인 기준선이 전역 임계값보다 정상 생활변화 오탐을 줄이는지 동일 조건에서 비교한다.")
    master.base.callout(
        doc,
        "세 가지 성공조건",
        "작동하는 웹서비스 > 설명 가능한 오탐 감소 > 인간의 안전한 최종판단",
        fill=master.MINT,
        accent=master.TEAL,
    )


def add_problem_and_evidence(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "02 Problem & Evidence",
        "조기 발견과 안전한 다음 행동 사이의 공백",
        "보호수단은 존재하지만, 개인의 장기 변화와 생활맥락을 행원 후속업무까지 연결하는 층은 분절돼 있다",
    )
    master.paragraph(
        doc,
        "금융권에는 FDS·ASAP 같은 사기 탐지수단과 신탁·후견·안심차단 같은 보호수단이 있다. "
        "그러나 개인의 장기 금융생활 변화와 생활맥락을 확인해 적합한 보호수단·행원 후속업무로 잇는 과정은 여전히 분절돼 있다.",
        style="Lead",
    )
    master.base.add_process(
        doc,
        [
            ("평소 생활", "작은 변화가 흩어져 보임"),
            ("변화 누적", "중복·미납·낯선 거래"),
            ("맥락 확인", "정상 변화와 위험 구분"),
            ("사람 검토", "질문·설명·기록"),
            ("보호 연결", "검증된 다음 행동"),
        ],
        colors=[master.LIGHT_GRAY, master.PALE_GOLD, master.SKY, master.PALE_RED, master.MINT],
    )
    doc.add_heading("고객과 행원의 미충족 업무", level=2)
    master.base.table(
        doc,
        ["고객", "행원·금융회사"],
        [
            ["여러 은행·카드·자동이체에 흩어진 작은 변화를 스스로 종합하기 어렵다.", "FDS 경보 이후 사실확인·중립 질문·제도검색·상담기록·재연락이 반복된다."],
            ["위험을 알아차려도 안심차단·어카운트인포·두낫콜·신탁·후견 중 무엇을 언제 쓸지 판단하기 어렵다.", "모든 고객의 장기 거래 변화를 사람이 전수점검할 수 없고, 거래·서류·규정·법률 데이터도 파편화돼 있다."],
        ],
        [4680, 4680],
        compact=True,
    )
    doc.add_heading("연구가 말하는 범위와 제품의 안전선", level=2)
    master.base.table(
        doc,
        ["금융행동", "근거 수준", "안심리듬에서의 사용"],
        [
            ["청구서 누락·연체 / 계좌관리 실수", "상대적으로 강함", "지속·반복과 데이터 누락 여부를 함께 확인"],
            ["과잉지출·현금흐름 악화 / 사기 취약성", "부분·연관 근거", "소득·필수지출·생활사건과 결합해 설명"],
            ["중복결제·잊은 구독", "합리적 가설", "진단 신호가 아닌 확인이 필요한 오류로 표시"],
            ["시간·지역·ATM / 동일상품 반복가입", "직접 근거 제한", "여행·이사 등 생활맥락 또는 소비자보호 검토 신호로만 사용"],
        ],
        [2780, 1840, 4740],
        compact=True,
    )
    master.base.callout(
        doc,
        "출력 원칙",
        "‘치매 가능성 78%’가 아니라 ‘최근 평소와 다른 금융활동이 확인되어 본인 확인이 필요합니다.’라고 안내한다.",
        fill=master.PALE_RED,
        accent=master.RED,
    )
    source_p = master.source(
        doc,
        [
            ("JAMA Internal Medicine 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
            ("Neurology 2009", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
            ("JAMA Network Open 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
        ],
    )
    source_p.paragraph_format.space_before = master.Pt(2)
    source_p.paragraph_format.space_after = master.Pt(0)


def add_final_korea(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "03 Korea",
        "대한민국의 대응과 남은 공백",
        "사기 FDS·신탁·후견·사전신청형 보호수단은 작동하지만 목적과 시점이 다르다",
    )
    master.base.table(
        doc,
        ["현재 수단", "작동 방식", "안심리듬과의 경계"],
        [
            ["은행 FDS·ASAP", "접속·거래·공유 사기정보로 의심거래를 탐지하고 확인·지연·지급정지", "보이스피싱 중심. 장기 개인 금융생활 변화 확인은 별도 층"],
            ["치매안심 재산관리서비스", "NPS가 신탁받은 현금성 자산을 계획에 따라 지급·관리", "2026년 시범사업. 전 금융거래 모니터링이 아니며 자격은 상담 확인"],
            ["공공후견·성년후견", "법원 결정 범위에서 재산·복지·의료·주거 업무 지원", "실시간 조기탐지가 아니라 신청·심판 기반 법적 지원"],
            ["민간 치매안심신탁", "사전 계약·지정대리인·발동조건에 따라 생활·의료비 지급", "맡긴 자산만 보호하고 가입능력·수수료·대리인 신뢰 이슈 존재"],
            ["지정인 알림", "일부 고위험 상품·카드대출 이용을 선택 지정인에게 통지", "옵트인·특정 상품·사후 알림 중심이며 일상거래 변화는 미포함"],
            ["안심차단·어카운트인포·두낫콜", "거래 사전차단, 자동이체 일괄조회·해지, 영업연락 차단", "변화 탐지 후 실행할 수 있는 국내 행동수단으로 연결"],
        ],
        [2250, 3650, 3460],
        compact=True,
    )
    doc.add_heading("사실을 정확히 말하는 방법", level=2)
    master.add_bullet(doc, bullet_id, "ASAP은 2025년 10월 출범 후 2026년 4월까지 31.7만 건의 정보공유를 통해 5,261건의 계좌 지급정지 조치와 474.6억 원의 피해예방 성과를 냈다. AI 단독 자동조치라고 표현하지 않는다.")
    master.add_bullet(doc, bullet_id, "은행 FDS도 고령자·신규 수취인·단시간 다회 이체 등을 일부 탐지한다. 따라서 ‘한국에는 이상거래 탐지가 없다’고 주장하지 않는다.")
    master.add_bullet(doc, bullet_id, "언론 보도에 따르면 신한은행도 2026년 고령층 금융거래 이상징후 모니터링을 개발·고도화 중이다. 공개자료상 정식 출시·구체 기능은 확인되지 않아 ‘국내 최초·유일’ 주장을 하지 않는다.")
    master.source(
        doc,
        [
            ("금융위 FDS·ASAP", "https://www.fsc.go.kr/po010102/86997"),
            ("NPS 치매안심 재산관리서비스", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
            ("보건복지부 안내", "https://www.mohw.go.kr/menu.es?mid=a10712010600"),
            ("금융보안원 FDS 가이드", "https://www.fsec.or.kr/bbs/detail?bbsNo=11355&menuNo=69"),
            ("신한은행 동향(언론 보도)", "https://www.mk.co.kr/news/special-edition/12088908"),
        ],
    )


def add_global_and_positioning(doc: Document, number_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "04 Global & Positioning",
        "해외 지원 방식과 안심리듬의 연결층",
        "진단이 아니라 모니터링·제한형 지갑·신뢰연락인·법적 대리·사람 개입을 조합한다",
    )
    master.base.table(
        doc,
        ["유형", "대표 사례", "작동 방식", "시사점"],
        [
            ["읽기전용 모니터링", "EverSafe·Carefull", "여러 계좌를 연결해 개인 패턴 대비 이상을 본인·선택 가족에게 알림", "탐지 후 사람의 확인이 필요하며 자체 차단권한은 없음"],
            ["제한형 지갑", "Sibstar·True Link", "한도·ATM·업종·채널을 사전에 정해 승인단계에서 허용·거절", "선불카드 밖 거래는 보호하지 못하고 과잉통제 위험"],
            ["신뢰연락·거래별 보류", "미국 FINRA·캐나다 증권", "증권계좌 착취 의심 특정 거래를 제한적으로 보류하고 신뢰연락인에게 확인", "연락인은 거래권한이 없고 업권·조건이 제한됨"],
            ["취약고객·사기대응", "영국 FCA·지급지연 제도", "쉬운 지원과 사람 확인, 합리적 사기 의심 시 지급을 제한기간 지연", "고객지원·지연·경찰·보상은 서로 다른 제도"],
            ["일상관리·자금잠금", "일본 일상생활자립지원·싱가포르 Money Lock", "일상비 지원과 큰 자산 관리 분리, 고객이 일부 자금을 이체 불가 상태로 잠금", "정적 보호수단과 개인 변화탐지는 별개"],
        ],
        [1320, 1900, 3760, 2380],
        compact=True,
    )
    doc.add_heading("기존 체계와의 역할 분담", level=2)
    master.base.table(
        doc,
        ["구분", "기존 FDS·ASAP", "안심리듬", "신탁·후견 등"],
        [
            ["주 목적", "사기·의심계좌 탐지와 조치", "금융생활 변화 확인과 후속업무 지원", "맡긴 자산·법원 권한 범위의 지속 관리"],
            ["차별 초점", "공개 운영목적상 사기·계정탈취 대응 중심", "생활맥락 재평가·설명·기록·재연락", "사전계약 또는 판단능력 저하 이후"],
            ["결과", "추가인증·지연·지급정지", "쉬운 설명·본인 확인·행원 검토", "계획지급·대리·감독"],
        ],
        [1350, 2750, 2900, 2360],
        compact=True,
        header_fill=master.TEAL,
    )
    master.base.callout(
        doc,
        "차별성",
        "개인 기준선 탐지 → 생활맥락 재확인 → 동의된 공동확인 → 직원 판단 → 검증된 보호수단 안내를 하나의 사건 흐름으로 묶는다.",
        fill=master.MINT,
        accent=master.TEAL,
    )
    master.source(
        doc,
        [
            ("FINRA Rule 2165", "https://www.finra.org/rules-guidance/rulebooks/finra-rules/2165"),
            ("FCA FG21/1", "https://www.fca.org.uk/publication/finalised-guidance/fg21-1.pdf"),
            ("EverSafe", "https://www.eversafe.com/for-families/"),
            ("Carefull", "https://getcarefull.com/"),
            ("Sibstar", "https://www.sibstar.co.uk/"),
            ("True Link", "https://www.truelinkfinancial.com/"),
            ("CIRO Trusted Contact", "https://www.ciro.ca/office-investor/avoiding-fraud-and-protecting-your-investments/why-you-should-consider-appointing-trusted-contact-person"),
            ("Singapore Money Lock", "https://www.moneysense.gov.sg/scams/moneylock/"),
        ],
    )


def add_journey_and_spec(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "06 User Journey",
        "3개 화면·하나의 사건·사람의 최종판단",
        "고객 화면이 본체이고 행원 화면은 탐지 이후 보호업무를 완결하는 도입 모듈이다",
    )
    master.base.table(
        doc,
        ["화면", "보는 것", "핵심 행동", "수용기준"],
        [
            ["1. 금융생활 변화", "12개월 변화 요약, 비교근거 3개, 동일 alertId", "경보 상세 열기", "질병 라벨 0개, 첫 상호작용 5초 이내"],
            ["2. 생활맥락 확인", "근거 거래, 중립 질문, 동의 상태", "본인 거래 확인·확인 불가·응답 보류", "preDecision 보존, postDecision 별도 기록"],
            ["3. 행원 사건검토", "사건 묶음, 확인질문, 공식 보호수단, 기록초안", "승인·오탐 종결·재연락", "승인 전 차단·한도변경·제3자 통보 0건"],
        ],
        [1750, 3420, 2020, 2170],
        compact=True,
    )
    doc.add_heading("상태기계와 화면 문구", level=2)
    master.base.add_process(
        doc,
        [
            ("OPEN", "변화사건 생성"),
            ("AWAITING_CONTEXT", "생활맥락 질문"),
            ("CLOSED_NORMAL", "정상 변화로 종결"),
            ("BANK_REVIEW", "행원 검토 대기"),
        ],
        colors=[master.PALE_GOLD, master.SKY, master.MINT, master.PALE_RED],
    )
    doc.add_heading("핵심 하드 규칙", level=2)
    master.add_bullet(doc, bullet_id, "고객의 단순 ‘괜찮아요’만으로 고액 신규수취인·사기의심 hard signal을 자동 해제하지 않는다.")
    master.add_bullet(doc, bullet_id, "정상으로 확정된 생활맥락만 유효기간을 두고 기준선 갱신 후보로 사용하며, 강한 신규증거가 생기면 다시 확인한다.")
    master.add_bullet(doc, bullet_id, "신뢰연락인 미동의 상태에서는 UI와 API가 모두 차단되고, 동의해도 최소 확인정보만 제공한다.")
    master.add_bullet(doc, bullet_id, "최초 판정을 덮어쓰지 않고 탐지시점 근거 봉인본·맥락 출처·알고리즘·직원결정 버전을 함께 남긴다.")
    master.base.callout(
        doc,
        "MVP 범위",
        "보호수단 목록·연락인 설정·감사 타임라인은 세 화면 안의 패널 또는 모달로 구현한다. 실제 여러 금융회사 연결·거래실행은 하지 않는다.",
        fill=master.SKY,
        accent=master.BLUE,
    )


def add_detection_and_evaluation(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "08 Detection & Evaluation",
        "개인 기준선이 오탐을 줄이는지 증명",
        "설명문 생성보다 동일 holdout·동일 경보예산에서 계산과 반례를 비교한다",
    )
    master.base.table(
        doc,
        ["신호", "MVP 필수", "예외·안전처리"],
        [
            ["금액·횟수 급증", "일·주 중앙값과 MAD의 단측 modified z", "MAD=0·희소값이면 IQR·분위수·절대규칙 또는 LOW_CONFIDENCE"],
            ["신규 수취인·반복송금·중복결제", "수취인 사전과 금액·시간창의 결정적 규칙", "표기변형·본인계좌·취소·환불·pending 분리"],
            ["정기납부 누락", "주기 추정과 grace period", "데이터 단절·연결장애를 미납으로 오인하지 않음"],
            ["점진 추세·수준 전환", "연속창 추세와 간단 온라인 변화점", "고도화 시 Theil-Sen·Page-Hinkley 비교"],
            ["사건 융합", "reasonCode와 정책표로 incident grouping", "중복억제·cooldown·월 경보예산"],
        ],
        [2050, 4060, 3250],
        compact=True,
    )
    doc.add_heading("A/B/C 검증과 합성 골든셋", level=2)
    master.base.table(
        doc,
        ["비교군", "구성", "확인 질문"],
        [
            ["A. 전역 규칙", "전 고객 공통 금액·횟수 임계값", "개인차 때문에 정상 생활이 얼마나 과다 경보되는가"],
            ["B. 개인 기준선", "규칙 + 중앙값/MAD + 준비상태", "개인화가 사용자-월당 오탐을 얼마나 줄이는가"],
            ["C. 기준선+맥락", "B + 확인된 정상 생활사건", "정상경보를 줄이면서 위험사건 검토를 유지하는가"],
        ],
        [1800, 3300, 4260],
        compact=True,
        header_fill=master.TEAL,
    )
    master.add_bullet(doc, bullet_id, "시연용 A/B 페르소나와 평가셋을 분리하고, 고객 단위 holdout·여러 seed·미래정보 없는 시점재현으로 검증한다.")
    master.add_bullet(doc, bullet_id, "정상인데 이상처럼 보이는 어려운 반례로 여행·이사·가족지원·큰 합법구매·급여일 변화·데이터 단절·취소거래를 포함한다.")
    master.add_bullet(doc, bullet_id, "지표는 사건 단위 precision/recall, false alerts/user-month, 탐지지연, unsafe downgrade, 미동의 통보, no-key fallback, 90초 완주율이다.")
    master.base.callout(
        doc,
        "증거의 한계",
        "합성데이터 결과는 코드 회귀와 설계 비교를 위한 시뮬레이션이다. 실제 금융회사 성능·고령층 효과·임상 예측력으로 주장하지 않는다.",
        fill=master.PALE_RED,
        accent=master.RED,
    )


def add_policy_and_safety(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "09 Policy & Safety",
        "단계적 개입·동의·소비자보호 안전장치",
        "탐지보다 중요한 것은 누가 어떤 정보를 보고 어떤 행동을 할 수 있는지다",
    )
    master.base.table(
        doc,
        ["단계", "조건", "출력", "결정권자"],
        [
            ["L0 기록", "낮은 단일 변화", "대시보드 기록", "고객"],
            ["L1 확인", "설명 필요한 단일 사건", "근거 3개·중립 질문", "고객"],
            ["L2 계획", "반복·지속 또는 다중 신호", "검증된 행동 체크리스트·상담 권유", "고객"],
            ["L3 검토", "본인 거래 확인 불가·무응답·hard signal", "행원 사건 큐·기록 초안", "행원"],
            ["L4 외부연계", "사기 임박·별도 상담 의사", "은행·112·1332·선택 공공상담 안내", "권한기관"],
        ],
        [1200, 2670, 3500, 1990],
        compact=True,
    )
    doc.add_heading("신뢰연락인과 자동결정 경계", level=2)
    master.add_bullet(doc, bullet_id, "제공받는 자·목적·항목·기간·철회방법을 고객이 별도로 선택하고, 연락인도 초대 수락과 본인확인을 한다.")
    master.add_bullet(doc, bullet_id, "연락인은 후견인이나 법정대리인이 아니며 송금·해지·차단·전체내역 조회 권한을 자동 취득하지 않는다.")
    master.add_bullet(doc, bullet_id, "가족이 금융착취 당사자일 수 있으므로 의심 연락인은 제외하고 행원 전용 검토·복수 연락인·철회·열람로그를 둔다.")
    master.add_bullet(doc, bullet_id, "AI는 알림·설명·업무초안만 제공하고, 권리·재산에 영향을 주는 조치는 본인 재확인 또는 담당자의 실질 검토를 거친다.")
    master.base.callout(
        doc,
        "정량 안전 게이트",
        "미동의 제3자 제공 0건 · 행원 승인 전 계좌조치 0건 · 질병·인지상태 점수 0개 · 근거 없는 보호수단 안내 0건",
        fill=master.MINT,
        accent=master.TEAL,
    )
    master.source(
        doc,
        [
            ("개인정보보호법 제23조", "https://www.law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1027416043"),
            ("자동화된 결정 안내", "https://m.pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=10611"),
            ("금융위 금융분야 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142"),
        ],
    )


def add_architecture_and_ai(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "10 Architecture & AI",
        "Spring 모듈형 모놀리스 기본경로",
        "탐지·정책결정과 생성형 설명을 분리하고 키가 없어도 전체 서비스가 작동한다",
    )
    master.base.add_process(
        doc,
        [
            ("ledger", "합성거래·정규화"),
            ("detection", "규칙·MAD·추세"),
            ("case", "맥락·동의·정책"),
            ("explanation", "템플릿·선택 LLM"),
            ("demo/audit", "3화면·감사로그"),
        ],
        colors=[master.LIGHT_GRAY, master.SKY, master.PALE_GOLD, master.MINT, master.LIGHT_GRAY],
    )
    master.base.table(
        doc,
        ["계층", "MVP 구현", "금융회사 확장"],
        [
            ["배치", "단일 Spring 앱 + PostgreSQL, 공개 URL, 합성데이터", "은행 VPC/내부망, 승인된 전용 AI 중계"],
            ["탐지·정책", "Java 규칙·통계·상태기계, LLM 의존 없음", "모델 registry·shadow/canary·kill switch"],
            ["설명", "구조화 ExplanationFacts → 템플릿 기본", "공개문서 외부 LLM / 내부정보 내부·전용 LLM"],
            ["RAG", "필요한 공식 문서만 승인 수집", "문서 ID·페이지·조항·효력일·버전·checksum·권한필터"],
            ["감사", "동의·신호·결정·직원수정 append-only", "업무 DB 분리·변경불가 저장·보존/파기 통제"],
        ],
        [1600, 3900, 3860],
        compact=True,
    )
    doc.add_heading("LLM·데이터 라우팅", level=2)
    master.add_bullet(doc, bullet_id, "PUBLIC·SYNTHETIC만 외부 LLM 또는 템플릿을 허용하고, 개인신용·인증정보는 MVP에서 차단한다.")
    master.add_bullet(doc, bullet_id, "검색 문서가 입력보다 민감하면 effectiveDataClass = max(입력 등급, 검색 문서 등급)으로 상향 라우팅한다.")
    master.add_bullet(doc, bullet_id, "키 없음·800ms 초과·429·5xx·schema 오류는 즉시 템플릿으로 fallback하며 reasonCode·actionCode·연락권한은 바뀌지 않는다.")
    master.add_bullet(doc, bullet_id, "원문 거래·상담·prompt/completion은 일반 로그에 저장하지 않고 사건 근거·모델·프롬프트·문서 버전만 감사한다.")
    master.base.callout(
        doc,
        "권장 고정 스택",
        "Java 21 · Spring Boot 4.1.0 · Spring AI 2.0.0(선택) · Spring MVC/Security · Thymeleaf/HTMX · PostgreSQL · Flyway · Testcontainers",
        fill=master.SKY,
        accent=master.BLUE,
    )


def add_adoption_and_roadmap(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "11 Adoption & Roadmap",
        "금융회사 도입가치와 제출까지의 실행계획",
        "고객 추천 AI가 아니라 소비자보호 후속업무를 줄이는 단일 Agent로 시작한다",
    )
    master.base.table(
        doc,
        ["이해관계자", "도입가치", "실증 KPI"],
        [
            ["소비자보호·FDS팀", "경보 후 확인·설명·기록·재연락 품질을 표준화", "사건 검토시간, 기록 작성시간, 재연락 누락률"],
            ["고객", "정상 생활변화를 설명하고 불필요한 통제를 피하며 필요한 도움을 선택", "경보 이해도, 정상종결률, 이의제기 처리시간"],
            ["준법·보안", "동의·근거·알고리즘·직원결정의 감사 가능성", "미동의 제공 0건, 승인 전 실행 0건, 로그 완전성"],
        ],
        [2100, 4490, 2770],
        compact=True,
    )
    doc.add_heading("도입 단계", level=2)
    master.base.add_process(
        doc,
        [
            ("1. 공모전", "합성데이터·3화면·90초"),
            ("2. 내부 PoC", "비식별 사건·직원 사용성"),
            ("3. 제한 실증", "동의 고객·단일 기관"),
            ("4. 제휴", "금융회사·허가 마이데이터"),
            ("5. 확장", "신탁·후견·공공상담 선택연계"),
        ],
        colors=[master.SKY, master.MINT, master.PALE_GOLD, master.SKY, master.LIGHT_GRAY],
    )
    doc.add_heading("제출까지 우선순위", level=2)
    master.base.table(
        doc,
        ["기간", "우선순위", "완료조건"],
        [
            ["8/13~8/19", "문제정의·세 화면 기능명세", "상태·동의 차단·동일 alertId 정책 고정"],
            ["8/20~8/30", "합성데이터·탐지엔진·A/B 데모", "MAD·예외·Reset·행원승인·템플릿 fallback"],
            ["8/31~9/3", "골든셋·성능비교", "전역/개인/맥락 비교와 안전건수 검증"],
            ["9/4~9/7", "문서·배포 안정화", "PDF·기능명세·공개 URL smoke test·복구절차"],
        ],
        [1700, 3300, 4360],
        compact=True,
        header_fill=master.TEAL,
    )
    master.add_bullet(doc, bullet_id, "마이데이터 2.0은 상용제도지만 스타트업이 전 금융거래를 곧바로 가져오는 권한은 아니다. 상용화는 금융회사 또는 허가사업자 제휴를 전제로 한다.")
    master.add_bullet(doc, bullet_id, "전사 AI 플랫폼을 새로 만드는 것이 아니라 향후 금융사의 전사 Agent 플랫폼에 탑재 가능한 소비자보호 업무 Agent로 제안한다.")


def add_final_appendix(doc: Document, bullet_id: int) -> None:
    master.new_page(doc)
    master.section_title(
        doc,
        "Appendix",
        "표현 금지선·출시 게이트·주요 출처",
        "과장과 낙인을 피하고, 구현과 문서의 상태를 일치시킨다",
    )
    master.base.table(
        doc,
        ["피해야 할 주장", "최종본 표현"],
        [
            ["치매 위험도·인지저하 가능성 산출", "개인 기준선 대비 금융생활 변화와 본인 확인 필요 사유"],
            ["국내 탐지체계가 없다 / 국내 최초·유일", "FDS 이후 장기 맥락확인·후속업무 층을 결합한 차별성"],
            ["AI가 거래 차단·연락인 통보·상품가입 실행", "본인 재확인·별도 동의·직원 승인·법적 권한 이후의 외부 절차"],
            ["마이데이터로 즉시 전 금융사 연결", "MVP는 합성데이터, 상용화는 금융회사·허가사업자 제휴"],
            ["합성 정확도=실제 성능", "코드 회귀와 설계 비교를 위한 시뮬레이션 결과"],
            ["완료송금 취소 가능", "은행 즉시 연락·착오송금 반환지원 또는 사기신고 가능성을 조건부 안내"],
        ],
        [3670, 5690],
        compact=True,
    )
    doc.add_heading("출시 전 필수 게이트", level=2)
    gates = [
        "동일 alertId·원거래·근거·사전판정을 고정한 맥락 비교",
        "MAD=0·cold-start·결측·지연·취소·환불·기준선 오염 처리",
        "incident 병합·중복억제·cooldown과 미래정보 없는 시점재현",
        "미동의 연락 0건·승인 전 계좌조치 0건·no-key fallback 100%",
        "익명 세션 격리·Reset 멱등성·모듈검증·통합테스트·공개 URL 감시",
    ]
    for gate in gates:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = master.Pt(3.5)
        master.base.set_font(p.add_run("□ "), size=10, bold=True, color=master.TEAL)
        master.base.set_font(p.add_run(gate), size=10)
    doc.add_heading("주요 출처", level=2)
    links = [
        ("금융보안원 2026 금융 AI Challenge", "https://www.fsec.or.kr/bbs/detail?bbsNo=11997&menuNo=66"),
        ("금융위원회 FDS·ASAP 운영성과", "https://www.fsc.go.kr/po010102/86997"),
        ("국민연금공단 치매안심 재산관리서비스", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
        ("금융위원회 마이데이터 2.0", "https://fsc.go.kr/po010101/84780"),
        ("Spring Boot System Requirements", "https://docs.spring.io/spring-boot/system-requirements.html"),
        ("Spring AI Getting Started", "https://docs.spring.io/spring-ai/reference/getting-started.html"),
    ]
    for label, url in links:
        p = doc.add_paragraph()
        master.apply_num(p, bullet_id)
        master.base.hyperlink(p, label, url, size=8.7)
    master.paragraph(
        doc,
        "본 문서는 팀 공유·공모전 출품 준비용 통합 최종본이며 의료·법률·투자자문이 아니다. "
        "실제 제출은 대회 제공 양식과 최신 원문에 맞춰 재배치하고, 제도·법령·상품 상태는 제출 직전에 다시 확인한다.",
        style="Small",
    )


def renumber_kickers(doc: Document) -> None:
    replacements = {
        "04 GLOBAL & POSITIONING": "04 GLOBAL & POSITIONING",
        "07 FIELD VIEW": "05 FIELD VIEW",
        "06 USER JOURNEY": "06 USER JOURNEY",
        "09 DEMO": "07 DEMO",
        "08 DETECTION & EVALUATION": "08 DETECTION & EVALUATION",
        "09 POLICY & SAFETY": "09 POLICY & SAFETY",
        "10 ARCHITECTURE & AI": "10 ARCHITECTURE & AI",
        "16 DATA & AUDIT": "11 DATA & AUDIT",
        "11 ADOPTION & ROADMAP": "12 ADOPTION & ROADMAP",
    }
    for paragraph in doc.paragraphs:
        replacement = replacements.get(paragraph.text)
        if replacement is None or replacement == paragraph.text:
            continue
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = replacement
        else:
            master.base.set_font(paragraph.add_run(replacement), size=9, bold=True, color=master.TEAL)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    master.patch_base_tokens()
    master.base.set_cell_margins = master.set_cell_margins
    master.base.set_table_geometry = master.set_table_geometry

    doc = Document()
    bullet_id, number_id = master.configure_styles(doc)
    master.set_header_footer(doc)
    doc.core_properties.title = "안심리듬 2026 금융 AI Challenge 최종 통합본"
    doc.core_properties.subject = "금융생활 변화 조기알림 및 보호업무 코파일럿"
    doc.core_properties.author = "Ansim Rhythm Project Team"
    doc.core_properties.keywords = "안심리듬, 금융 AI, 개인 기준선, 생활맥락, 금융소비자보호, 행원 코파일럿"

    # standard_business_brief preset with the existing editorial-cover treatment.
    add_final_cover(doc)
    add_final_exec_summary(doc, bullet_id)
    master.add_competition(doc, bullet_id)
    add_problem_and_evidence(doc, bullet_id)
    add_final_korea(doc, bullet_id)
    add_global_and_positioning(doc, number_id)
    master.add_field_view(doc, bullet_id)
    add_journey_and_spec(doc, bullet_id)
    master.add_demo(doc)
    add_detection_and_evaluation(doc, bullet_id)
    add_policy_and_safety(doc, bullet_id)
    add_architecture_and_ai(doc, bullet_id)
    master.add_data_and_audit(doc)
    add_adoption_and_roadmap(doc, bullet_id)
    add_final_appendix(doc, bullet_id)
    renumber_kickers(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUT)
    master.audit(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
