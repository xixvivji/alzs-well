from __future__ import annotations

from pathlib import Path
from urllib.parse import urlparse
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_ansim_rhythm_report as master


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "Ansim_Rhythm_2026_Financial_AI_Challenge_Integrated_Final.docx"

# narrative_proposal preset, with named Korean brand overrides.
# Page: Letter, 1-inch margins, 9360 DXA usable width.
# Body: 11 pt, 8 pt after, 1.333 lines. H1/H2/H3 use preset spacing.
# Lists: 0.181/0.375/0.194 in, 4 pt after, 1.208 lines.
# Tables: 9360 DXA, indent 120, margins 80/120/80/120, header fill #F4F6F9.
FONT = "AppleGothic"
NAVY = "17324D"
BLUE = "2B6F9F"
TEAL = "2C8C82"
SKY = "EAF4F8"
MINT = "EAF6F3"
GOLD = "D69B2D"
PALE_GOLD = "FFF5DF"
RED = "A64242"
PALE_RED = "FAECEC"
INK = "202B35"
GRAY = "5F6B75"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D9E0E5"
WHITE = "FFFFFF"


def patch_tokens() -> None:
    for name, value in {
        "FONT": FONT,
        "NAVY": NAVY,
        "BLUE": BLUE,
        "TEAL": TEAL,
        "SKY": SKY,
        "MINT": MINT,
        "GOLD": GOLD,
        "PALE_GOLD": PALE_GOLD,
        "RED": RED,
        "PALE_RED": PALE_RED,
        "INK": INK,
        "GRAY": GRAY,
        "LIGHT_GRAY": LIGHT_GRAY,
        "MID_GRAY": MID_GRAY,
        "WHITE": WHITE,
    }.items():
        setattr(master, name, value)
        setattr(master.base, name, value)
    master.base.set_cell_margins = master.set_cell_margins
    master.base.set_table_geometry = master.set_table_geometry


def set_style_font(style, name: str) -> None:
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)


def set_run_font(run, *, size=None, bold=None, italic=None, color=None, name=FONT) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def patch_numbering(doc: Document, bullet_id: int, number_id: int) -> None:
    numbering = doc.part.numbering_part.element
    for num_id in (bullet_id, number_id):
        num = next((n for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId")) == str(num_id)), None)
        if num is None:
            continue
        aid = num.find(qn("w:abstractNumId"))
        if aid is None:
            continue
        abstract_id = aid.get(qn("w:val"))
        abstract = next(
            (a for a in numbering.findall(qn("w:abstractNum")) if a.get(qn("w:abstractNumId")) == abstract_id),
            None,
        )
        if abstract is None:
            continue
        lvl = abstract.find(qn("w:lvl"))
        ppr = lvl.find(qn("w:pPr"))
        ind = ppr.find(qn("w:ind"))
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        tabs = ppr.find(qn("w:tabs"))
        tab = tabs.find(qn("w:tab"))
        tab.set(qn("w:pos"), "540")
        spacing = ppr.find(qn("w:spacing"))
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")


def configure_styles(doc: Document):
    bullet_id, number_id = master.configure_styles(doc)
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333
    pf.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        set_style_font(style, FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Source", "Lead", "Small", "Table Text", "Kicker"):
        set_style_font(styles[name], FONT)
    styles["Source"].font.size = Pt(8.2)
    styles["Source"].paragraph_format.space_before = Pt(4)
    styles["Source"].paragraph_format.space_after = Pt(4)
    styles["Source"].paragraph_format.line_spacing = 1.10
    styles["Lead"].font.size = Pt(12)
    styles["Lead"].paragraph_format.space_after = Pt(8)
    styles["Lead"].paragraph_format.line_spacing = 1.20
    styles["Small"].font.size = Pt(8.6)
    styles["Small"].paragraph_format.space_after = Pt(4)
    styles["Table Text"].font.size = Pt(8.8)
    styles["Table Text"].paragraph_format.line_spacing = 1.10
    styles["Kicker"].font.size = Pt(9)

    patch_numbering(doc, bullet_id, number_id)
    return bullet_id, number_id


def setup_page(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header
        p = header.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run("안심리듬 | 치매머니 프로젝트"), size=8, color=GRAY)
        footer = section.footer
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        set_run_font(p.add_run("2026.08.13  |  "), size=8.2, color=GRAY)
        master.base.add_page_field(p)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def section_title(doc: Document, kicker: str, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(style="Kicker")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(kicker.upper()), size=9, bold=True, color=TEAL)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(title), size=23, bold=True, color=NAVY)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        p.paragraph_format.line_spacing = 1.18
        set_run_font(p.add_run(subtitle), size=11.5, color=GRAY)


def paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    set_run_font(p.add_run(text))


def add_bullet(doc: Document, bullet_id: int, text: str) -> None:
    p = doc.add_paragraph()
    master.apply_num(p, bullet_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text))


def add_number(doc: Document, number_id: int, text: str) -> None:
    p = doc.add_paragraph()
    master.apply_num(p, number_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    set_run_font(p.add_run(text))


def callout(doc: Document, label: str, text: str, *, fill=SKY, accent=BLUE) -> None:
    master.base.callout(doc, label, text, fill=fill, accent=accent)


def table(doc: Document, headers, rows, widths, *, header_fill=NAVY, compact=False):
    return master.base.table(doc, headers, rows, widths, header_fill=header_fill, compact=compact)


def source(doc: Document, items) -> None:
    master.source(doc, items)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("2026 금융 AI Challenge"), size=11, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("안심리듬"), size=38, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("ANSIM RHYTHM"), size=15, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)
    p.paragraph_format.line_spacing = 1.30
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        p.add_run("인지취약 고객의 금융의사결정을 보호하는\nAI 금융안전 코파일럿"),
        size=16,
        bold=True,
        color=INK,
    )
    callout(
        doc,
        "CORE PROMISE",
        "치매를 판별하지 않는다. 평소와 다른 금융생활을 이해 가능한 확인과 사람의 보호업무로 연결한다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(35)
    p.paragraph_format.space_after = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("프로젝트 주제  치매머니  |  대외 서비스명  안심리듬"), size=10.5, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(76)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("통합 최종 보고서  |  기준일 2026. 8. 13."), size=10, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("국내외 조사 · 현업 음성자문 · 망분리 대응 · MVP 기능명세 반영"), size=8.7, color=GRAY)


def add_exec(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(
        doc,
        "Executive Summary",
        "최종 제안: 탐지를 넘어 사건처리로",
        "기존 FDS를 대체하지 않고, 개인 장기 변화·고객 이해확인·행원 대응·공적 안전망을 하나의 흐름으로 잇는다",
    )
    callout(
        doc,
        "한 문장 정의",
        "안심리듬은 고객의 평소 금융 리듬과 급성 사기 맥락의 변화를 발견하고, 고객이 이해할 수 있는 방식으로 확인한 뒤, 행원이 동의·권한·공식 절차에 따라 안전조치를 연결하도록 돕는 금융안전 코파일럿이다.",
        fill=SKY,
        accent=BLUE,
    )
    doc.add_heading("두 개의 AI, 하나의 사건", level=2)
    table(
        doc,
        ["AI 축", "핵심 기능", "차별점"],
        [
            ["개인화 금융리듬 방패", "급성 사기위험과 장기 금융생활 변화를 별도로 탐지하고 비교근거 최대 3개를 생성", "연령·치매확률이 아니라 개인 기준선 대비 무엇이 달라졌는지 설명"],
            ["안심동행 AI + 행원 코파일럿", "쉬운 설명·음성·Teach-back·행원 질문·기록·후속조치 지원", "고객 접근성과 경보 이후 사건처리를 하나의 폐쇄형 흐름으로 연결"],
        ],
        [2200, 4140, 3020],
        compact=True,
        header_fill=TEAL,
    )
    doc.add_heading("수상 전략", level=2)
    for text in (
        "‘치매 특화 FDS’가 아니라 ‘인지취약 고객의 금융의사결정 보호 운영체계’로 포지셔닝한다.",
        "동일 alertId·동일 거래·동일 사전판정을 유지하고 생활맥락만 바꾸어 정상종결과 행원검토가 갈리는 장면을 시연한다.",
        "망분리를 약점이 아니라 ‘개인정보는 내부에, AI에는 최소 이유코드만’이라는 신뢰 설계로 보여준다.",
        "생성형 AI는 쉬운 설명·Teach-back·행원 기록에 실제 사용하되 점수·차단·상품가입은 결정하지 않는다.",
    ):
        add_bullet(doc, bullet_id, text)
    callout(doc, "성공조건", "작동하는 웹서비스 > 고객 이해 확인 > 행원 업무 완결 > 안전한 인간 최종판단", fill=MINT, accent=TEAL)


def add_competition(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "01 Competition", "대회 요구와 최종 제출 전략", "아이디어 설명보다 작동하는 고객·행원 웹 여정과 재현 가능한 검증을 우선한다")
    table(
        doc,
        ["공식 요구", "최종 대응"],
        [
            ["금융 현장 문제 + 실제 웹서비스", "고객 안심동행 화면과 행원 사건검토 화면을 같은 사건으로 연결"],
            ["기획서 PDF", "문제·차별성·AI/데이터·기대효과·확장성 중심으로 공식 양식에 재배치"],
            ["기능명세서 PDF", "구현 완료 기능만 기재하고 클릭 경로·샘플 입력·예상 출력·제약을 명시"],
            ["서비스 URL", "무로그인 데모, 합성데이터, Reset, 템플릿 fallback, 상태 모니터링"],
        ],
        [2520, 6840],
        compact=True,
    )
    doc.add_heading("제출 운영 기준", level=2)
    for text in (
        "1차 제출: 2026년 9월 7일 10:00까지 기획서 PDF·기능명세서 PDF·배포 URL.",
        "URL 의무 가용시간: 9월 7일 11:00~9월 11일 23:59. 로그인·API 키 없이 첫 화면에서 문제와 시나리오가 보여야 한다.",
        "기능명세서에는 ‘계획’과 ‘구현’을 혼합하지 않는다. 실제 작동하지 않는 신탁·보호자 발송·은행 API는 로드맵으로 분리한다.",
        "공식 페이지에 상세 배점표가 공개되지 않았으므로 비공식 평가항목을 공식 기준으로 단정하지 않는다.",
    ):
        add_bullet(doc, bullet_id, text)
    callout(doc, "팀 내부 합격선", "5초 내 문제 이해 · 90초 내 완주 · 미동의 연락 0건 · 승인 전 계좌조치 0건 · Reset 재현성 100%", fill=PALE_GOLD, accent=GOLD)
    source(doc, [("대회 공식 안내", "https://daker.ai/public/hackathons/2026-finance-ai-challenge")])


def add_problem(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "02 Problem & Evidence", "치매머니의 규모보다 중요한 ‘보호 공백’", "진단 후 자산관리와 급성 사기탐지는 존재하지만 진단 전·경도인지장애 단계의 금융생활 변화와 다음 행동은 분절돼 있다")
    table(
        doc,
        ["근거", "2026년 관점", "설계 함의"],
        [
            ["치매·경도인지장애", "2026년 치매 약 101만 명, 경도인지장애 약 314만 명 전망", "진단 치매만 대상으로 하면 늦다. 예방형·자발적 지원모드가 필요"],
            ["독거·디지털 격차", "치매가구 중 1인가구 비중이 높고, 고령층 디지털 역량이 접근수준보다 낮음", "앱 기능 추가보다 음성·한 단계 안내·직원 대체경로가 중요"],
            ["치매머니", "2023년 진단이력자 자산 154조 원은 정부 행정추계이며 정의별 변동 가능", "확정 시장규모로 과장하지 않고 금융시스템 위험의 참고수치로 사용"],
        ],
        [2100, 3680, 3580],
        compact=True,
        header_fill=TEAL,
    )
    doc.add_heading("연구 기반 신호와 한계", level=2)
    table(
        doc,
        ["신호 층", "예시", "제품 사용 원칙"],
        [
            ["급성 사기위험", "신규 수취인, 고액·반복송금, 신규기기, ATM 한도반복, 전화지시·긴박감", "즉시 구조화 질문과 안전 콜백·행원검토"],
            ["장기 금융생활 변화", "PIN 재설정, 로그인 감소·실패, 카드분실, 지점·콜 증가, 미납·중복납부, 소비범주 축소", "지속성·데이터 누락·생활사건을 함께 확인"],
            ["상호작용·상황", "같은 질문 반복, 거래목적 설명 곤란, 보호자가 답변 독점, 두려움", "치매 진단이 아니라 지원수준을 정하는 사람 검토 근거"],
        ],
        [1950, 4150, 3260],
        compact=True,
    )
    callout(doc, "연구의 안전선", "영국·미국 연구는 금융역량 상실 또는 치매진단 전 행동변화의 연관성을 보여주지만 개인의 치매를 예측하거나 인과를 증명하지 않는다.", fill=PALE_RED, accent=RED)
    source(
        doc,
        [
            ("복지부 치매역학", "https://www.mohw.go.kr/board.es?act=view&bid=0027&list_no=1484959&mid=a10503010300&nPage=16&tag="),
            ("치매머니 행정추계", "https://www.betterfuture.go.kr/front/notificationSpace/pressReleaseDetail.do?articleId=473"),
            ("NIA 디지털격차", "https://www.nia.or.kr/site/nia_kor/ex/bbs/View.do?bcIdx=27832&cbIdx=81623"),
            ("JAMA 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
            ("JAMA 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
        ],
    )


def add_landscape(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "03 Market & Gap", "이미 있는 것과 안심리듬이 맡을 연결층", "FDS·ASAP·안심차단·신탁·후견을 대체하지 않고 경보 이후 고객확인과 사건처리를 보완한다")
    table(
        doc,
        ["현재 체계", "실제 활용", "남은 공백 / 안심리듬 역할"],
        [
            ["ASAP·은행 FDS", "의심정보 공유, 확인·지연·지급정지. 2025.10~2026.4 약 31.7만 건 공유, 5,261건 지급정지, 474.6억 원 예방", "개인 장기 변화·생활맥락·행원 기록과 후속조치 표준화"],
            ["카카오뱅크 시퀀스 FDS", "앱 행동·거래 시퀀스와 XAI를 운영. 일 1,800만 회 추론, 평균 20ms 미만", "특이패턴+XAI 자체가 아닌 인지친화적 확인과 사건종결에서 차별화"],
            ["안심차단·고령자모드", "거래 사전차단, 쉬운 용어·자주 쓰는 기능·보이스피싱 신고", "큰 글씨를 넘어 맥락별 확인강도·Teach-back·행원 인계"],
            ["공적·민간 재산관리", "NPS 재산관리 2026.4.22 시범(7개 지역본부, 현금성 자산 최대 10억 원), 하나은행 치매안심 금융센터 등", "필요 고객 발견→동의 기반 상담예약→근거·상담요약 패키지"],
            ["공공후견·성년후견", "신청·심판 후 정해진 법적 범위에서 업무 지원", "실시간 탐지가 아니라 반복 위험 시 선택 가능한 공식 연결 경로"],
        ],
        [1900, 4390, 3070],
        compact=True,
        header_fill=NAVY,
    )
    doc.add_heading("국내외 벤치마크의 공통점", level=2)
    for text in (
        "미국: 고객 과거이력과 비교한 이상탐지, 신뢰연락처, 제한적 거래보류, 직원·기관 연계를 조합한다.",
        "영국: 취약고객 중앙기록, 실시간 음성분석, 직원 개입과 동등한 고객결과를 중시한다.",
        "일본: MUFG신탁은행의 AI도 인지기능 ‘진단’이 아니라 직원의 금융상품 적합성 판단 참고정보로 제한한다.",
        "가족이 착취 당사자일 수 있으므로 신뢰연락인은 송금권한이 없는 비상연락 수단이어야 하며 공유범위를 최소화한다.",
    ):
        add_bullet(doc, bullet_id, text)
    source(
        doc,
        [
            ("금융위 ASAP 성과", "https://www.fsc.go.kr/po010102/86997"),
            ("카카오뱅크 시퀀스 FDS", "https://tech.kakaobank.com/posts/2606-sequence-based-fds-model/"),
            ("NPS 재산관리 시범", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
            ("보건복지부 제5차 치매관리종합계획", "https://www.mohw.go.kr/gallery.es?act=view&b_list=12&bid=0003&cg_code=&keyField=&list_no=379875&mid=a10607030000&nPage=10&orderby=&vlist_no_npage=24"),
            ("CFPB 고령자 금융착취 대응", "https://www.consumerfinance.gov/archive/newsroom/cfpb-issues-advisory-and-report-for-financial-institutions-on-preventing-elder-financial-abuse/"),
            ("FCA 취약고객 사례", "https://www.fca.org.uk/publications/good-and-poor-practice/delivering-vulnerable-customers"),
            ("MUFG·IBM 사례", "https://jp.newsroom.ibm.com/2025-02-12-Financial-Instruments-Suitability-Check-AI-App"),
        ],
    )


def add_customer_ai(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "04 Customer AI", "고객을 위한 ‘안심동행 AI’", "자유대화형 금융챗봇이 아니라 고객이 금융업무를 이해하고 안전하게 완료하도록 돕는 폐쇄형 접근성 AI")
    callout(
        doc,
        "대표 차별점",
        "일반 FDS가 ‘거래가 이상한가’를 본다면, 안심동행 AI는 고객이 ‘누구에게 왜 보내며 어떤 결과가 생기는지’를 이해하고 있는지 Teach-back으로 확인한다.",
        fill=MINT,
        accent=TEAL,
    )
    table(
        doc,
        ["고객 기능", "구체 동작", "AI와 안전 경계"],
        [
            ["말로 찾는 업무", "‘카드를 잃어버렸어요’처럼 말하면 업무 후보 1~2개를 재확인하고 해당 화면으로 안내", "내부 업무목록만 사용. 송금·가입·해지를 자동실행하지 않음"],
            ["한 단계씩 안심 안내", "큰 글씨·고대비·한 화면 한 작업·읽어주기·다시 듣기·취소·자막", "승인 문구 범위. 금액·약관 의미를 바꾸지 않음"],
            ["거래 Teach-back", "금액·수취인·거래목적·예상결과를 설명한 뒤 고객이 자기 말로 다시 확인", "인지점수를 만들지 않고 불일치 시 재설명 또는 행원 연결"],
            ["사기 상황 점검", "수사기관·안전계좌·고수익·전화지시 등 승인 질문을 한 문항씩 제시", "자유검색 금지. AI가 강한 위험신호를 단독 해제하지 않음"],
            ["결과 기억카드", "오늘 한 일·바뀐 내용·다음 할 일·공식 문의처를 쉬운 말과 음성으로 제공", "원음 대신 확정 업무코드·결과값으로 요약"],
            ["행원 이어서 도움", "고객 동의 후 확인답변·미해결 항목·다음 질문을 최소정보 카드로 전달", "승인 전 계좌조치·가족 통보 없음"],
        ],
        [1840, 4610, 2910],
        compact=True,
        header_fill=TEAL,
    )
    doc.add_heading("접근성 기본값", level=2)
    for text in (
        "기능이 많아도 ‘몰라서 못 쓰는’ 문제를 해결하도록 고객의 말에서 업무를 찾고 한 번에 한 작업만 보여준다.",
        "음성만으로 고액이체·한도변경을 확정하지 않고 화면 재확인 또는 행원 승인으로 전환한다.",
        "원음성은 기본 미보관하며 자막·재입력·직접 수정·즉시 행원 연결을 항상 제공한다.",
    ):
        add_bullet(doc, bullet_id, text)


def add_screen_blueprint(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(
        doc,
        "05 Product Screens",
        "MVP 화면 설계도",
        "고객 변화·이해 확인·행원 검토·동의 차단을 동일 사건과 상태기계로 연결한다",
    )
    callout(
        doc,
        "화면 연결 원칙",
        "세 화면은 동일 alertId·원거래·evidence·preDecision을 공유한다. 고객 답변은 위험점수를 덮어쓰지 않고 다음 확인과 사람 검토를 결정하는 맥락으로만 사용한다.",
        fill=SKY,
        accent=BLUE,
    )
    table(
        doc,
        ["화면", "고객·행원이 보는 핵심 문구", "주요 조작", "상태·안전 증거"],
        [
            ["① 고객 변화 확인", "‘평소와 다른 변화 3가지가 있어요’ — 금액·신규 수취인·신규기기만 쉬운 말로 설명", "자세히 보기·다시 듣기·모름·직원 도움", "질병 라벨 0개, 이유코드 최대 3개, 자동조치 없음"],
            ["② 거래 Teach-back", "‘누구에게, 왜 보내시나요?’ — 한 문항씩 듣고 고객이 자기 말로 거래목적과 결과를 확인", "말하기·텍스트 입력·수정·다시 설명·취소", "정상 맥락은 CLOSED_NORMAL, 강한 사기 맥락은 BANK_REVIEW"],
            ["③ 행원 사건검토", "급성 사기위험과 장기 금융생활 변화를 분리하고 근거·질문·권한·보호수단을 한 화면에 표시", "질문 수정·고객 확인·승인·종결·재연락", "직원 결정 전 송금차단·가족통보 0건, 모든 수정·결정 감사"],
            ["④ 동의·연락 차단", "신뢰연락인 공유항목·연락조건·기간·철회상태와 의심 연락인 제외 여부를 확인", "동의 설정·철회·공유 미리보기·행원 승인", "동의 없으면 CONSENT_REQUIRED, 서버도 발송 거부; MVP 실제 발송 없음"],
        ],
        [1600, 3500, 2250, 2010],
        compact=True,
        header_fill=NAVY,
    )
    doc.add_heading("접근성·시연 고정값", level=2)
    for text in (
        "핵심 문구 18pt 이상·고대비·한 화면 한 질문·모든 음성에 자막과 직접 수정 경로를 제공한다.",
        "A/B 시나리오는 Reset 뒤에도 같은 alertId와 사전판정을 사용해 생활맥락만으로 결과가 갈리는 것을 증명한다.",
        "음성 실패·키 없음·timeout에도 버튼/텍스트 기본경로와 승인된 설명 템플릿으로 90초 시연을 완주한다.",
    ):
        add_bullet(doc, bullet_id, text)


def add_staff_ai(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "06 Staff Copilot", "행원 보호업무 코파일럿", "음성 자문의 핵심인 업무 태스크 분해·데이터 정비·인간 최종판단을 실제 사건 흐름으로 구현한다")
    master.base.add_process(
        doc,
        [
            ("1 동의", "목적·범위·철회"),
            ("2 탐지", "규칙·통계·시계열"),
            ("3 근거 봉인", "이유코드·버전"),
            ("4 고객 확인", "쉬운 설명·Teach-back"),
            ("5 행원 검토", "질문·권한·조치"),
            ("6 기록·후속", "요약·재연락·감사"),
        ],
        colors=[LIGHT_GRAY, SKY, PALE_GOLD, MINT, SKY, LIGHT_GRAY],
    )
    table(
        doc,
        ["행원에게 제공", "AI가 하는 일", "사람이 결정"],
        [
            ["사건 타임라인", "위험근거 3~5개, 개인 기준선, 급성/장기 신호 분리", "경보가 실제 보호업무 대상인지"],
            ["중립 질문", "승인 문서와 사건근거로 확인 질문·표준스크립트 작성", "질문 수정·생략·추가"],
            ["권한·동의 카드", "보호자·신뢰연락인·대리권·철회상태 표시", "연락 또는 외부기관 연결 여부"],
            ["보호수단 안내", "안심차단·112·1332·신탁·치매안심센터 등 공식 경로 검색", "적합한 설명과 상담 연결"],
            ["상담기록 초안", "고객 확인답변·근거·직원 조치를 구조화 요약", "최종 문구와 사건 종결"],
        ],
        [2010, 4300, 3050],
        compact=True,
    )
    doc.add_heading("현업 효과를 측정하는 방법", level=2)
    for text in (
        "사건당 정보검색시간·화면 전환 수·질문 작성시간·상담기록시간·재연락률을 As-is와 비교한다.",
        "AI 출력의 수정률·근거문서 적합률·직원 간 판단 일치도·override 사유를 감사한다.",
        "전사 플랫폼을 새로 만드는 것이 아니라 금융사의 기존 Agent 플랫폼에 탑재 가능한 소비자보호 업무 Agent로 설계한다.",
    ):
        add_bullet(doc, bullet_id, text)


def add_journey_demo(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "07 Journey & Demo", "같은 경보, 다른 이해, 다른 다음 행동", "고객 안심동행 화면과 행원 사건검토 화면을 동일 alertId로 연결하는 90초 핵심 시연")
    table(
        doc,
        ["시간", "시연", "화면에 남는 증거"],
        [
            ["0~15초", "무로그인 시작→신규 수취인·평소 대비 4.2배·신규기기 경보 선택", "동일 alertId·evidence·preDecision"],
            ["15~35초", "고객 AI가 쉬운 문장·음성으로 변화 3개를 안내", "이유코드만 전달, 질병 라벨 0개"],
            ["35~55초", "A: ‘이사 보증금, 계약서 확인 가능’ Teach-back", "생활맥락 검증 후 CLOSED_NORMAL"],
            ["55~65초", "Reset 후 동일 사건 B 재실행", "원거래·알고리즘·사전판정 동일"],
            ["65~78초", "B: ‘검찰이 안전계좌로 보내라고 했다’ 응답", "수사기관 사칭 맥락, BANK_REVIEW"],
            ["78~88초", "행원 화면에서 근거·질문·권한·기록초안 확인", "승인 전 차단·가족통보 없음"],
            ["88~90초", "안전원칙 요약", "진단 없음·자동조치 없음·미동의 공유 없음"],
        ],
        [1250, 5010, 3100],
        compact=True,
        header_fill=NAVY,
    )
    doc.add_heading("세 번째 장기변화 카드", level=2)
    paragraph(doc, "급성 송금사건과 별도로 PIN 재설정·로그인 감소·카드분실·미납 증가가 수개월 누적된 사례를 보여준다. 출력은 ‘장기 금융생활 변화 확인 필요’이며 ‘치매 의심’으로 표시하지 않는다.")
    callout(doc, "하드 규칙", "고객의 단순 ‘괜찮아요’나 Teach-back 1회만으로 고액 신규수취인·수사기관 사칭 같은 강한 신호를 자동 해제하지 않는다.", fill=PALE_RED, accent=RED)


def add_detection(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "08 Detection & Evaluation", "이중 탐지엔진과 검증 설계", "급성 사기위험과 장기 금융생활 변화는 별도 점수·별도 정책으로 다루며 어느 쪽도 치매를 의미하지 않는다")
    table(
        doc,
        ["엔진", "핵심 신호", "출력·다음 행동"],
        [
            ["급성 사기위험", "신규 수취인·비정상 고액/반복·신규기기·해외/가상자산·전화지시·긴박감", "즉시 고객확인, 안전콜백, 행원 검토"],
            ["장기 금융생활 변화", "MAD·추세·변화점, 로그인/PIN/카드분실, 미납·중복납부, 소비범주·채널 변화", "지원필요 확인, 상담 제안, 기준선 검토"],
        ],
        [1840, 4900, 2620],
        compact=True,
        header_fill=TEAL,
    )
    doc.add_heading("합성 골든셋 A/B/C", level=2)
    table(
        doc,
        ["비교군", "구성", "검증 질문"],
        [
            ["A 전역 규칙", "모든 고객에 공통 금액·횟수 임계값", "개인차 때문에 정상생활이 얼마나 과다 경보되는가"],
            ["B 개인 기준선", "중앙값/MAD·추세·준비상태", "동일 경보예산에서 오탐과 선행시간이 개선되는가"],
            ["C 기준선+맥락", "B + 검증된 정상 생활사건", "정상경보를 줄이면서 위험사건 검토를 유지하는가"],
        ],
        [1720, 3480, 4160],
        compact=True,
    )
    doc.add_heading("제출 전 실제로 채울 지표", level=2)
    for text in (
        "탐지: PR-AUC, 고정 심사인력에서 recall, 고객 1,000명·월당 오경보, 평균 선행시간.",
        "고객: 시나리오 완주율, Teach-back 성공률, 뒤로가기·도움요청, 이해도.",
        "행원: 검토·질문·기록 시간, 재연락 누락률, AI 수정률.",
        "안전·공정성: 잘못된 보류·해제시간, 필수지출 중단, 미동의 공유, 연령·지역·디지털수준별 TPR/FPR 격차.",
    ):
        add_bullet(doc, bullet_id, text)
    callout(doc, "현재 상태", "비교 설계와 목표는 제안이며 측정결과가 아니다. 합성 골든셋 결과를 실행해 실제 수치와 그래프를 채우기 전까지 성능 성과로 표현하지 않는다.", fill=PALE_GOLD, accent=GOLD)


def add_network_architecture(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "09 Network Separation & AI", "망분리를 전제로 한 AI 이원구조", "망분리가 고객 AI를 불가능하게 만드는 것이 아니라 데이터 등급과 모델 위치를 분리하게 만든다")
    master.base.add_process(
        doc,
        [
            ("원거래·기준선", "내부망 FDS"),
            ("이유코드", "최소·구조화"),
            ("내부 AI", "STT/TTS·폐쇄형 RAG"),
            ("고객 확인", "쉬운 설명·Teach-back"),
            ("행원 승인", "조치·감사로그"),
        ],
        colors=[PALE_RED, PALE_GOLD, SKY, MINT, LIGHT_GRAY],
    )
    table(
        doc,
        ["영역", "허용 데이터·기능", "금지·통제"],
        [
            ["내부망", "원거래, 개인 기준선, 고객 답변, 내부 소형 LLM, 정책엔진, 감사로그", "최소권한·모델 registry·골든셋·kill switch"],
            ["단말/내부 처리", "STT/TTS, 의도분류, 업무코드, 승인 문구", "원음 기본 미보관, 자막·수정·인간 인계"],
            ["외부 상용 AI", "PUBLIC·SYNTHETIC, 공개문서 요약·표현 생성", "주민번호·계좌번호·거래내역·고객음성·개인신용정보·고객연결 가능 이유코드 차단"],
            ["규칙·정책엔진", "송금 진행 여부, 보류 가능성, 연락권한, 상태전환", "LLM 출력으로 정책·reasonCode를 덮어쓰지 않음"],
        ],
        [1840, 4470, 3050],
        compact=True,
        header_fill=NAVY,
    )
    doc.add_heading("MVP와 은행 실서비스", level=2)
    table(
        doc,
        ["구분", "대회 MVP", "금융회사 도입"],
        [
            ["데이터", "합성 거래·합성 음성/텍스트", "고객 동의·목적제한·내부 정보등급"],
            ["모델", "외부 LLM 또는 로컬 모델 + 결정론적 템플릿 fallback", "내부 소형 LLM·전용 중계·승인된 상용 AI"],
            ["연동", "가짜 발송·가짜 조치, 실제 은행 API 없음", "샌드박스·보안성평가·기존 업무시스템과 단계 연동"],
            ["표시", "‘합성데이터’와 ‘실서비스 배치 차이’를 화면에 명시", "모델·문서·프롬프트·직원결정 버전 감사"],
        ],
        [1500, 3900, 3960],
        compact=True,
        header_fill=TEAL,
    )
    callout(doc, "음성 주의", "브라우저 음성 API는 공급자 서버로 원음이 전달될 수 있다. 대회에서는 합성 시연·명시적 입력으로 제한하고, 실서비스는 단말 또는 내부환경 처리·원음 미보관·전사문 최소화를 전제로 한다.", fill=PALE_GOLD, accent=GOLD)
    source(
        doc,
        [
            ("금융권 AI 이원 활용체계", "https://www.fsc.go.kr/no010101/83594"),
            ("생성형 AI 절차 개선", "https://fsc.go.kr/no010101/86712"),
            ("SaaS 망분리 개선", "https://www.fsc.go.kr/no010101/86745"),
            ("금융분야 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142"),
            ("금융보안원 AI 보안 안내서", "https://www.fsec.or.kr/bbs/detail?bbsNo=11977&menuNo=222"),
        ],
    )


def add_safety(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "10 Policy & Safety", "동의·단계적 개입·되돌릴 수 있는 보호", "누가 어떤 정보를 보고 어떤 행동을 할 수 있는지까지 제품으로 구현한다")
    table(
        doc,
        ["단계", "조건", "출력", "결정권자"],
        [
            ["L0 기록", "낮은 단일 변화", "대시보드 기록", "고객"],
            ["L1 확인", "설명이 필요한 사건", "근거·쉬운 설명·Teach-back", "고객"],
            ["L2 지원", "반복·지속 또는 다중 신호", "안전계획·상담 권유", "고객"],
            ["L3 검토", "본인 확인 불가·강한 사기신호", "행원 사건 큐·기록초안", "행원"],
            ["L4 외부연계", "사기 임박·별도 상담의사", "112·1332·안심차단·공공/민간 상담", "권한기관"],
        ],
        [1190, 2550, 3730, 1890],
        compact=True,
        header_fill=TEAL,
    )
    doc.add_heading("신뢰연락인 안전설계", level=2)
    for text in (
        "고객이 판단능력이 있을 때 연락조건·공유항목·기간·철회방법·1·2순위 연락인을 별도로 설정한다.",
        "가족에게 원거래 전체를 보내지 않고 ‘본인 확인 필요’ 등 최소정보만 제공하며 거래권한은 부여하지 않는다.",
        "가족·보호자가 착취 당사자일 수 있으므로 의심 연락인은 제외하고 복수 연락처·직원전용 검토·철회·열람로그를 둔다.",
        "전면 계좌동결로 임대료·약값·돌봄비가 막히지 않도록 의심거래와 필수지출을 분리한다.",
    ):
        add_bullet(doc, bullet_id, text)
    callout(doc, "출시 가드레일", "미동의 제3자 공유 0건 · 직원 승인 전 계좌조치 0건 · 필수 생활비 중단 0건 · 질병/인지점수 0개 · 근거 없는 안내 0건", fill=PALE_RED, accent=RED)


def add_data_audit(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "11 Data & Audit", "결정론적 기본경로와 감사 가능성", "탐지·정책결정과 생성형 설명을 분리하고 AI 장애가 고객의 권리와 사건 상태를 바꾸지 않게 한다")
    table(
        doc,
        ["모듈", "책임", "감사 필드"],
        [
            ["ledger/baseline", "합성거래 정규화·개인 기준선·준비상태", "data_cutoff, feature_code, median, MAD, algorithm_version"],
            ["detection", "급성/장기 신호·reasonCode·incident grouping", "alert_id, evidence_ids, pre_decision, feature_hash"],
            ["case/policy", "맥락·동의·상태기계·연락권한", "context_source, consent_snapshot, post_decision, expiry"],
            ["explanation", "템플릿·내부/외부 LLM·승인문서 RAG", "model/prompt/schema/document version, citation, fallback"],
            ["staff/audit", "직원수정·승인·종결·재연락", "staff_override, action_code, timestamps, retention"],
        ],
        [1770, 3660, 3930],
        compact=True,
        header_fill=NAVY,
    )
    doc.add_heading("LLM 실패와 데이터 통제", level=2)
    for text in (
        "키 없음·timeout·429·5xx·스키마 오류 시 승인 템플릿으로 fallback하고 reasonCode·actionCode·연락권한은 보존한다.",
        "검색 문서가 입력보다 민감하면 effectiveDataClass를 상향해 외부 라우팅을 차단한다.",
        "원문 거래·음성·상담·prompt/completion은 일반 로그에 저장하지 않고 사건근거·버전·직원결정만 최소범위로 감사한다.",
        "질병 위험점수나 인지상태 라벨은 데이터 필드·화면·로그 어디에도 만들지 않는다.",
    ):
        add_bullet(doc, bullet_id, text)


def add_roadmap(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "12 Roadmap & Submission", "대회 MVP부터 제한 실증까지", "이번 제출은 고객 AI 세 기능과 행원 코파일럿의 완결된 한 사건에 집중한다")
    table(
        doc,
        ["우선순위", "구현 범위", "완료조건"],
        [
            ["P0 고객 MVP", "말로 찾는 업무·쉬운 설명·Teach-back·행원 연결", "버튼/텍스트 기본경로 완주, 음성 선택, 고위험 자동실행 0건"],
            ["P0 행원 MVP", "사건 타임라인·질문·권한·보호수단·기록초안", "직원 승인·수정·종결·감사로그 작동"],
            ["P0 검증", "전역/개인/맥락 골든셋, 고객·행원 사용성", "실제 측정수치·그래프·한계 명시"],
            ["P0 제출", "공식 기획서·기능명세·URL·smoke test", "9/7 10:00 제출, 의무 가용시간 모니터링"],
            ["P1 확장", "금융안전계획·신뢰연락인 동의 게이트", "실제 발송 없이 권한·철회·최소정보 시연"],
            ["P2 실증", "비식별 shadow mode→동의 고객·단일기관", "전문검토자 판정·재피해·필수지출·공정성 평가"],
        ],
        [1650, 4430, 3280],
        compact=True,
        header_fill=TEAL,
    )
    doc.add_heading("기능 상태 표기 원칙", level=2)
    table(
        doc,
        ["ID·기능", "이번 제출 상태", "기능명세서 완료조건"],
        [
            ["F-01 말로 찾는 업무", "P0 구현·검증 대상", "실제 URL 클릭경로·합성 입력·업무 후보·텍스트 대체경로 확인"],
            ["F-02 쉬운 설명·읽어주기", "P0 구현·검증 대상", "이유코드 기반 설명·자막·수정·템플릿 fallback 확인"],
            ["F-03 Teach-back·행원 연결", "P0 구현·검증 대상", "CLOSED_NORMAL/BANK_REVIEW·같은 alertId·감사로그 확인"],
            ["F-04 신뢰연락인·공공연계", "P1 향후계획", "명세 제외; 실제 발송 없이 CONSENT_REQUIRED·철회·권한만 시연"],
        ],
        [2600, 2080, 4680],
        compact=True,
        header_fill=NAVY,
    )
    callout(
        doc,
        "제출 원칙",
        "기획서는 목표·로드맵을, 기능명세서는 실제 URL의 완료·부분구현·미구현만 적는다. 작동하지 않는 기능은 완료로 표시하지 않는다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def add_appendix(doc: Document, bullet_id: int) -> None:
    page_break(doc)
    section_title(doc, "Appendix", "표현 금지선·검증 체크리스트·주요 출처", "과장과 낙인을 피하고 문서의 주장과 실제 구현 상태를 일치시킨다")
    table(
        doc,
        ["피해야 할 주장·기능", "최종 표현·대안"],
        [
            ["치매 가능성·인지취약 점수", "개인 기준선 대비 금융생활 변화와 본인 확인 필요 사유"],
            ["국내 최초·유일 FDS", "기존 FDS 이후 장기 맥락확인·고객 이해·사건처리 연결층"],
            ["AI가 송금·차단·가족통보·상품가입", "규칙·정책·본인확인·직원승인·법적 권한 후 별도 절차"],
            ["음성·표정으로 치매 추정", "단말/내부 STT·TTS와 거래 이해 Teach-back"],
            ["마이데이터로 즉시 전 금융사 연결", "MVP는 합성데이터, 상용화는 금융회사·허가사업자 제휴"],
            ["합성 정확도=실제 금융권 성능", "코드 회귀·설계비교용 시뮬레이션 결과와 한계"],
        ],
        [3630, 5730],
        compact=True,
        header_fill=NAVY,
    )
    doc.add_heading("출시·제출 전 체크", level=2)
    checks = (
        "동일 alertId·원거래·evidence·preDecision 고정",
        "MAD=0·cold-start·결측·취소·환불·데이터 단절 처리",
        "미동의 연락 0건·승인 전 조치 0건·음성 원본 기본 미보관",
        "고객 문구에 질병·인지 라벨 0개, 언제든 취소·행원 연결 제공",
        "no-key/timeout fallback, 익명 세션 격리, Reset 멱등성",
        "합성 골든셋 실제 실행값, 접근성·행원효율·공정성 지표 표시",
        "공식 기능명세서와 실제 URL의 구현상태 일치",
        "PDF 폰트 포함·링크·페이지번호·URL 가용성 최종 확인",
    )
    for text in checks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3.5)
        set_run_font(p.add_run("□ "), size=10, bold=True, color=TEAL)
        set_run_font(p.add_run(text), size=10)
    doc.add_heading("핵심 공식 출처", level=2)
    source(
        doc,
        [
            ("대회 안내", "https://daker.ai/public/hackathons/2026-finance-ai-challenge"),
            ("금융위 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142"),
            ("금융보안원 AI 보안 안내서", "https://www.fsec.or.kr/bbs/detail?bbsNo=11977&menuNo=222"),
            ("금융위 ASAP", "https://www.fsc.go.kr/po010102/86997"),
            ("보건복지부 치매역학조사", "https://www.mohw.go.kr/board.es?act=view&bid=0027&list_no=1484959&mid=a10503010300&nPage=16&tag="),
            ("NPS 재산관리", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
            ("카카오뱅크 FDS", "https://tech.kakaobank.com/posts/2606-sequence-based-fds-model/"),
            ("FINRA Rule 2165", "https://www.finra.org/rules-guidance/rulebooks/finra-rules/2165"),
        ],
    )
    paragraph(
        doc,
        "본 문서는 팀 내부 의사결정과 공모전 제출 준비를 위한 통합 최종본이다. 의료진단·법률자문·금융상품 추천을 제공하지 않으며, 제도·법령·상품·대회 일정은 제출 직전 공식 원문으로 다시 확인한다.",
        style="Small",
    )


def set_all_run_fonts(doc: Document) -> None:
    def apply_paragraph(p):
        for run in p.runs:
            current = run.font.size.pt if run.font.size else None
            set_run_font(run, size=current)

    for p in doc.paragraphs:
        apply_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_paragraph(p)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                apply_paragraph(p)


def normalize_table_geometry(doc: Document) -> None:
    for t in doc.tables:
        grid = [int(c.get(qn("w:w"))) for c in t._tbl.tblGrid]
        master.set_table_geometry(t, grid, indent=120)
        for row in t.rows:
            for cell in row.cells:
                master.set_cell_margins(cell, top=80, start=120, bottom=80, end=120)


def audit(path: Path) -> None:
    doc = Document(path)
    assert len(doc.sections) == 1
    section = doc.sections[0]
    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.left_margin == Inches(1.0)
    assert section.right_margin == Inches(1.0)
    assert section.top_margin == Inches(1.0)
    assert section.bottom_margin == Inches(1.0)
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Source", "Table Text"):
        assert doc.styles[style_name].font.name == FONT
    for t in doc.tables:
        tbl_pr = t._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa"
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == "120"
        grid = [int(c.get(qn("w:w"))) for c in t._tbl.tblGrid]
        assert sum(grid) == int(tbl_w.get(qn("w:w"))) == 9360
        for row in t.rows:
            assert row.height is None
            for idx, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tcw is not None and int(tcw.get(qn("w:w"))) == grid[idx]
    text = "\n".join(p.text for p in doc.paragraphs)
    forbidden = ("치매 가능성 78%", "국내 최초", "자동으로 가족에게")
    assert not any(term in text for term in forbidden)
    with ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        sects = root.findall(".//w:sectPr", ns)
        assert len(sects) == 1
    print(f"AUDIT OK: paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")


def build() -> None:
    patch_tokens()
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    bullet_id, number_id = configure_styles(doc)
    setup_page(doc)
    doc.core_properties.title = "안심리듬 2026 금융 AI Challenge 통합 최종 보고서"
    doc.core_properties.subject = "인지취약 고객 금융의사결정 보호 AI 금융안전 코파일럿"
    doc.core_properties.author = "Ansim Rhythm Project Team"
    doc.core_properties.keywords = "안심리듬, 치매머니, 금융 AI, Teach-back, 망분리, FDS, 행원 코파일럿"

    add_cover(doc)
    add_exec(doc, bullet_id)
    add_competition(doc, bullet_id)
    add_problem(doc, bullet_id)
    add_landscape(doc, bullet_id)
    add_customer_ai(doc, bullet_id)
    add_screen_blueprint(doc, bullet_id)
    add_staff_ai(doc, bullet_id)
    add_journey_demo(doc, bullet_id)
    add_detection(doc, bullet_id)
    add_network_architecture(doc, bullet_id)
    add_safety(doc, bullet_id)
    add_data_audit(doc, bullet_id)
    add_roadmap(doc, bullet_id)
    add_appendix(doc, bullet_id)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    set_all_run_fonts(doc)
    normalize_table_geometry(doc)
    doc.save(OUT)
    audit(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
