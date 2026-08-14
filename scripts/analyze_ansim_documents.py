from __future__ import annotations

import json
import re
import sys
from collections import Counter
from pathlib import Path
from zipfile import ZipFile

from docx import Document


def summarize(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    headings = [
        {"style": p.style.name, "text": p.text.strip()}
        for p in paragraphs
        if p.style and p.style.name.startswith("Heading")
    ]
    table_rows = []
    for table_index, table in enumerate(doc.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            table_rows.append(
                {
                    "table": table_index,
                    "row": row_index,
                    "cells": [cell.text.strip() for cell in row.cells],
                }
            )
    paragraph_text = "\n".join(p.text.strip() for p in paragraphs)
    table_text = "\n".join(" | ".join(row["cells"]) for row in table_rows)
    all_text = paragraph_text + "\n" + table_text
    normalized = re.sub(r"\s+", " ", all_text).strip()
    with ZipFile(path) as archive:
        hyperlinks = 0
        rels_name = "word/_rels/document.xml.rels"
        if rels_name in archive.namelist():
            rels = archive.read(rels_name).decode("utf-8", errors="replace")
            hyperlinks = rels.count("relationships/hyperlink")
        styles = archive.read("word/styles.xml").decode("utf-8", errors="replace")
    return {
        "path": str(path),
        "paragraphs": len(paragraphs),
        "headings": headings,
        "tables": len(doc.tables),
        "table_rows": len(table_rows),
        "hyperlinks": hyperlinks,
        "characters": len(normalized),
        "style_counts": Counter(p.style.name if p.style else "" for p in paragraphs),
        "full_text": all_text,
        "table_rows_detail": table_rows,
        "legacy_terms": {
            term: all_text.count(term)
            for term in (
                "치매머니",
                "Dementia Money",
                "치매 위험",
                "치매 의심",
                "인지저하",
                "Isolation Forest",
                "MAD",
                "Page-Hinkley",
                "Theil-Sen",
                "alertId",
                "신뢰연락인",
            )
        },
        "has_styles_xml": bool(styles),
    }


if __name__ == "__main__":
    paths = [Path(arg) for arg in sys.argv[1:]]
    results = [summarize(path) for path in paths]
    print(json.dumps(results, ensure_ascii=False, indent=2, default=dict))
