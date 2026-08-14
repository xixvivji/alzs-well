from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_dementia_money_report import (
    BLUE,
    GOLD,
    GRAY,
    INK,
    LIGHT_GRAY,
    MINT,
    NAVY,
    PALE_GOLD,
    PALE_RED,
    RED,
    SKY,
    TEAL,
    WHITE,
    W_NS,
    add_bullet,
    add_number,
    add_page_field,
    add_process,
    audit,
    callout,
    configure_styles,
    hyperlink,
    new_page,
    set_cell_border,
    set_font,
    set_repeat_table_header,
    set_table_geometry,
    shade_cell,
    source_paragraph,
    table,
    title_block,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "Ansim_Rhythm_AI_Financial_Safety_Proposal_2026_Revised.docx"


def configure_document(doc: Document) -> None:
    configure_styles(doc)

    # The proposal keeps the visual language of the original report while using
    # a slightly more compact proposal-centered hierarchy.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.24

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(13.2)
    doc.styles["Heading 3"].font.size = Pt(11.3)
    doc.styles["Lead"].font.size = Pt(12)
    doc.styles["Table Text"].font.size = Pt(8.8)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.76)
        section.bottom_margin = Inches(0.70)
        section.left_margin = Inches(0.88)
        section.right_margin = Inches(0.88)
        section.header_distance = Inches(0.30)
        section.footer_distance = Inches(0.30)

        header = section.header
        p = header.paragraphs[0]
        p.text = "안심리듬 | AI 금융안전 제안서"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        set_font(p.runs[0], size=8, color=GRAY)

        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("2026.08.13  |  ")
        set_font(r, size=8.2, color=GRAY)
        add_page_field(p)


def metric_strip(doc: Document, items: list[tuple[str, str]], *, fill: str = SKY) -> None:
    widths = [3120, 3120, 3120]
    t = doc.add_table(rows=1, cols=3)
    set_table_geometry(t, widths, indent=140)
    set_repeat_table_header(t.rows[0])
    for i, (value, label) in enumerate(items):
        cell = t.cell(0, i)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, fill if i != 1 else MINT)
        set_cell_border(
            cell,
            top={"val": "single", "sz": 8, "color": "D9E0E5"},
            left={"val": "single", "sz": 8, "color": "D9E0E5"},
            bottom={"val": "single", "sz": 8, "color": "D9E0E5"},
            right={"val": "single", "sz": 8, "color": "D9E0E5"},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(value)
        set_font(r, size=17, bold=True, color=NAVY)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=8.4, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def small_note(doc: Document, text: str, *, color: str = GRAY) -> None:
    p = doc.add_paragraph(style="Small")
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(text)
    set_font(r, size=8.2, color=color)


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("2026 금융 AI Challenge · 개정 제안서")
    set_font(r, size=11, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("안심리듬")
    set_font(r, size=42, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("ANSIM RHYTHM")
    set_font(r, size=15, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run("개인별 금융생활 변화를 발견하고\n고객과 금융회사의 선제 대응을 돕는 AI 금융안전 플랫폼")
    set_font(r, size=15, bold=True, color=INK)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(16)
    r = line.add_run("━" * 24)
    set_font(r, size=9, color=GOLD)

    callout(
        doc,
        "CORE PROMISE",
        "같은 경보라도 생활맥락에 따라 정상 변화는 종결하고, 설명이 필요한 사건만 행원 보호업무로 연결한다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("프로젝트 코드: 치매머니")
    set_font(r, size=9.4, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("조사·기획 기준일  2026. 8. 13.")
    set_font(r, size=9.2, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("질병 진단·인지상태 추론·자동 거래통제를 수행하지 않는 인간 중심 설계")
    set_font(r, size=8.5, color=GRAY)


def executive_summary(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "Executive Summary", "안심리듬 한눈에 보기", "조기경보와 보호업무를 하나의 닫힌 흐름으로 연결한다")
    callout(
        doc,
        "한 문장 정의",
        "안심리듬은 치매를 예측하지 않는다. 개인별 금융생활 변화를 발견하고, 생활맥락을 확인해 정상 변화는 해제하며, 설명이 필요한 사건만 행원 보호업무로 연결한다.",
    )

    metric_strip(doc, [("3개", "출품 MVP 화면"), ("90초", "핵심 데모"), ("25일", "Spring 중심 구현")])

    doc.add_heading("문제와 해결", level=2)
    table(
        doc,
        ["대상", "현재의 공백", "안심리듬의 가치"],
        [
            ["고객", "변화를 스스로 알아차리고 적절한 대책을 고르기 어려움", "개인 기준선 경보, 쉬운 설명, 생활맥락 확인, 안전계획"],
            ["행원", "FDS 경보 이후 연락·질문·제도검색·기록·후속관리가 반복됨", "사건 묶음, 우선순위, 확인질문, 근거 기반 조치·기록 초안"],
            ["금융회사", "취약고객 보호와 업무효율을 동시에 입증해야 함", "인간 최종판단을 보존하는 설명 가능 보호 프로세스"],
        ],
        [1350, 3850, 4160],
        compact=True,
    )

    doc.add_heading("출품 의사결정", level=2)
    table(
        doc,
        ["항목", "확정안"],
        [
            ["주 구매부서", "은행 소비자보호·FDS 후속업무팀"],
            ["AI의 역할", "탐지·설명·질문·근거검색·기록 초안"],
            ["판정 모델", "설명 가능한 통계·규칙이 본선, Isolation Forest는 오프라인 비교"],
            ["데이터", "20~30개 합성 페르소나 × 12개월; 실제 계좌·의료정보 미사용"],
            ["실행 경계", "가족 알림·한도변경·거래보류·지급정지는 자동 실행하지 않음"],
        ],
        [1900, 7460],
        compact=True,
        header_fill=TEAL,
    )
    small_note(doc, "현업 대화 반영: 모델 성능만이 아니라 망분리·감사·사람의 승인·기존 업무시스템 연결 가능성을 제안서의 중심으로 이동했다.")


def problem_and_evidence(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "01 Problem", "보호수단은 있지만, 발견과 행동은 끊겨 있다", "기존 보호수단을 부정하지 않고 정확한 연결 공백을 정의한다")
    p = doc.add_paragraph(style="Lead")
    r = p.add_run("정확한 문제는 ‘모든 제도가 사후 대응뿐’이 아니라, 개인의 장기 금융생활 변화를 발견하고 정상 맥락을 확인한 뒤 적합한 대책으로 연결하는 체계가 분절돼 있다는 점이다.")
    set_font(r, size=12, bold=True, color=NAVY)

    add_process(
        doc,
        [
            ("평소 생활", "변화가 작아 스스로 알아채기 어려움"),
            ("변화 누적", "미납·중복·반복송금이 흩어짐"),
            ("위험 인지", "본인·가족이 뒤늦게 확인"),
            ("대책 탐색", "차단·신탁·후견을 직접 선택"),
            ("후속 관리", "상담·기록·재연락이 분절"),
        ],
        colors=[LIGHT_GRAY, PALE_GOLD, PALE_RED, SKY, MINT],
    )

    doc.add_heading("두 사용자의 미충족 업무", level=2)
    table(
        doc,
        ["고객 측", "금융회사 측"],
        [
            ["이상 변화를 본인·가족이 먼저 인지해야 제도를 찾기 시작함", "FDS가 경보를 내도 실제 사기인지 생활변화인지 사람이 확인"],
            ["여러 계좌·카드의 작은 미납과 반복행동을 장기 추세로 보기 어려움", "고객 연락, 질문 구성, 보호서비스 검색, 기록·후속관리 반복"],
            ["어떤 보호수단을 언제, 어떤 순서로 신청할지 판단하기 어려움", "모든 고령 고객을 사람이 장기간 전수검토하는 것은 불가능"],
        ],
        [4680, 4680],
        compact=True,
    )

    doc.add_heading("연구 근거는 ‘연관성’, 서비스 출력은 ‘확인 필요’", level=2)
    p = doc.add_paragraph()
    p.add_run("연구에서는 인지저하·치매 전후에 연체, 계좌관리 실수, 지출통제 저하, 사기 취약성이 집단 수준에서 증가하는 경향이 관찰됐다. 그러나 입원·이사·여행·사별·소득변화도 유사한 패턴을 만들 수 있다. 따라서 거래데이터를 개인의 질병 확률로 변환하지 않는다.")
    callout(
        doc,
        "표현 가드레일",
        "금지: ‘치매 위험 78%’  →  사용: ‘최근 30일 금융패턴이 평소와 달라 본인 확인이 필요할 수 있습니다.’",
        fill=PALE_RED,
        accent=RED,
    )
    source_paragraph(
        doc,
        [
            ("JAMA Internal Medicine 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
            ("Neurology 2009", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
            ("JAMA Network Open 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
        ],
    )


def landscape_and_positioning(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "02 Landscape", "기존 FDS를 대체하지 않는 보완 계층", "차별점은 다기관 자체가 아니라 기준선·맥락·후속업무의 결합이다")

    doc.add_heading("현재 국내 보호수단의 역할", level=2)
    table(
        doc,
        ["체계", "이미 잘하는 일", "안심리듬이 보완하는 지점"],
        [
            ["FDS·ASAP", "접속·거래·공유 의심정보로 보이스피싱 위험을 탐지하고 조치 지원", "고객 개인의 장기 생활변화와 경보 이후 상담·기록·후속관리"],
            ["안심차단·지연이체", "고객이 사전에 신청한 계좌개설·여신·오픈뱅킹 차단과 거래 지연", "언제 어떤 조치를 검토할지 맥락 기반 체크리스트 제공"],
            ["공공·민간 신탁", "맡긴 자산을 사전계획에 따라 생활·의료·요양비로 관리", "대상 제도로 자동 확정하지 않고 상담 가능한 옵션으로 연결"],
            ["후견", "법원 심판 범위에서 재산·복지·의료 등 의사결정 지원", "신뢰연락인과 법적 대리권을 분리하고 조기상담 연결"],
        ],
        [1750, 3700, 3910],
        compact=True,
    )
    source_paragraph(
        doc,
        [
            ("금융위 FDS·ASAP 성과", "https://www.fsc.go.kr/po010102/86997"),
            ("금융거래 안심차단", "https://www.fsc.go.kr/no010101/85644"),
            ("NPS 재산관리 시범사업", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
            ("법원 성년후견", "https://www.scourt.go.kr/nm/min_3/min_3_12/index.html"),
        ],
    )

    doc.add_heading("비교 기준을 정밀하게 바꾼다", level=2)
    table(
        doc,
        ["비교", "FDS·ASAP", "안심리듬"],
        [
            ["주 목적", "사기·계정탈취·의심계좌 탐지와 조치", "금융생활 관리변화 확인과 보호업무 지원"],
            ["시간축", "실시간·사건 중심", "7·30·90일과 6~12개월 개인 추세"],
            ["핵심 입력", "거래·접속·기관 공유 사기정보", "동의된 거래·정기납부·생활맥락"],
            ["후속업무", "회사별 확인·차단·신고 프로세스", "질문·조치조건·기록·재연락 초안까지 닫힌 루프"],
        ],
        [1550, 3650, 4160],
        compact=True,
        header_fill=TEAL,
    )
    callout(
        doc,
        "진짜 차별축",
        "개인 기준선 + 생활맥락 재평가 + 국내 보호수단 적용조건/공식근거 + 행원 질문·기록·후속관리의 결합. EverSafe·Carefull도 다계좌 기준선·가족알림을 제공하므로 ‘국내 최초·유일’은 사용하지 않는다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def product_journey(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "03 Product", "고객 조기경보와 행원 코파일럿을 하나로", "AI는 발견·정리·추천을 맡고 고객과 행원이 최종 판단한다")
    add_process(
        doc,
        [
            ("1. 동의", "분석범위·연락인 권한을 분리 설정"),
            ("2. 기준선", "개인별 거래·납부·현금 리듬 계산"),
            ("3. 변화", "복수 신호와 기여 사유 생성"),
            ("4. 맥락", "입원·이사·여행·본인거래 확인"),
            ("5. 사람", "정상종결 또는 행원 보호업무 연결"),
        ],
        colors=[LIGHT_GRAY, SKY, SKY, MINT, PALE_GOLD],
    )

    doc.add_heading("세 화면으로 증명하는 사용자 여정", level=2)
    table(
        doc,
        ["화면", "고객·행원이 보는 것", "핵심 상태"],
        [
            ["1. 금융생활 변화", "평소값과 현재값, 신규 수취인·반복송금·납부누락의 기여사유", "CONTEXT_REQUIRED"],
            ["2. 생활맥락 확인", "본인거래 여부와 입원·이사·여행·가족지원 등 정상사유 질문", "CLOSED_NORMAL 또는 REVIEW_REQUIRED"],
            ["3. 행원 사건큐", "사건 타임라인, 확인질문, 보호조치 적용조건, 기록·재연락 초안", "PENDING_BANK_REVIEW"],
        ],
        [1900, 5150, 2310],
        compact=True,
    )

    doc.add_heading("핵심 골든 시나리오", level=2)
    table(
        doc,
        ["시나리오", "맥락·동의", "기대 결과"],
        [
            ["정상 입원비", "본인결제·입원 확인", "정상종결, 공과금 납부 체크리스트만 제공"],
            ["신규 수취인 반복고액 + 미납", "거래를 기억하지 못하고 정상사유 없음", "행원 검토큐 등록, 근거·질문·조치 초안"],
            ["신뢰연락인 미동의", "trustedContactConsent=false", "알림 미리보기만 표시, 발송은 BLOCKED_BY_CONSENT"],
        ],
        [2500, 3110, 3750],
        compact=True,
        header_fill=TEAL,
    )

    callout(
        doc,
        "권한 경계",
        "AI는 거래를 차단하지 않고, 가족에게 자동 연락하지 않으며, 치매·판단능력·사기 여부를 확정하지 않는다. 중요 조치는 고객 재확인 또는 권한 있는 직원의 실질 검토 이후에만 가능하다.",
        fill=PALE_RED,
        accent=RED,
    )


def ai_design(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "04 AI Design", "설명 가능한 탐지와 제한된 생성형 AI", "판정은 통계·규칙, 생성형 AI는 쉬운 설명과 업무 초안만 맡는다")

    doc.add_heading("탐지 파이프라인", level=2)
    add_process(
        doc,
        [
            ("거래 정규화", "합성 마이데이터형 스키마"),
            ("7·30·90일 특징", "금액·빈도·신규성·누락·추세"),
            ("개인 기준선", "중앙값·MAD·정기주기"),
            ("사유별 기여도", "reasonCode와 contribution"),
            ("맥락 재점수", "정상사유 감점, 사람 검토 전환"),
        ],
        colors=[LIGHT_GRAY, SKY, SKY, MINT, PALE_GOLD],
    )

    table(
        doc,
        ["분석 요소", "MVP 구현", "설명 예시"],
        [
            ["신규 수취인", "과거 수취인 집합과 최초 등장 여부", "처음 보는 수취인에게 2회 송금"],
            ["강건 이상도", "중앙값과 MAD 기반 robust z-score", "평소 32만원 대비 오늘 총 180만원"],
            ["정기납부 누락", "예상 납부일 ± 허용구간 규칙", "매달 납부하던 공과금 1건 누락"],
            ["점진적 변화", "4·8주 빈도와 금액 기울기", "현금인출 빈도가 8주간 증가"],
            ["복합 사건", "수취인·반복·누락·자산대비금액 합성", "단일 거래가 아닌 고객 사건으로 묶음"],
        ],
        [1900, 3350, 4110],
        compact=True,
    )

    doc.add_heading("생성형 AI의 허용·금지 범위", level=2)
    table(
        doc,
        ["허용", "금지"],
        [
            ["구조화된 사유를 고객용 쉬운 문장으로 변환", "거래 원문·식별자를 자유 프롬프트로 전송"],
            ["중립적 확인질문과 상담기록 초안 생성", "치매·사기·법 위반·적합성 최종판정"],
            ["공식 보호수단 문서의 조건·절차 요약", "근거 없는 상품추천·거래차단·가족연락 실행"],
        ],
        [4680, 4680],
        compact=True,
        header_fill=TEAL,
    )
    callout(
        doc,
        "고객 설명 예시",
        "평소 처음 보내는 사람에게 월 1회 이하, 평균 32만원을 송금했습니다. 오늘은 처음 보는 수취인에게 같은 금액을 두 번 보내고 공과금 1건이 누락돼 본인 거래인지 확인이 필요합니다.",
        fill=MINT,
        accent=TEAL,
    )


def demo_design(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "05 Demo", "같은 거래, 다른 맥락, 다른 다음 행동", "모델 과시보다 안전한 의사결정 흐름을 90초에 증명한다")
    callout(
        doc,
        "고정 사건 AR-2026-0813-001",
        "가상 고객 김안심(72세). 평소 신규 수취인 송금 월 0~1회·평균 32만원. 오늘 ‘에이치케어’에 90만원씩 2회 송금했고 같은 달 공과금 1건이 누락됐다.",
        fill=SKY,
        accent=BLUE,
    )

    doc.add_heading("Paired counterfactual", level=2)
    table(
        doc,
        ["항목", "A. 정상 생활변화", "B. 설명이 필요한 사건"],
        [
            ["고객 답변", "입원 중 병원비 결제대행처에 본인이 두 번 결제", "수취인과 송금을 기억하지 못하며 입원·여행도 없음"],
            ["상태", "CLOSED_NORMAL", "PENDING_BANK_REVIEW"],
            ["다음 행동", "공과금 납부 확인만 제안, 자동 전달 없음", "사건·근거·질문 초안을 행원 큐에 등록"],
            ["연락인", "불필요", "미동의이므로 BLOCKED_BY_CONSENT"],
            ["거래조치", "없음", "행원 승인 전 실행 0건"],
        ],
        [1600, 3880, 3880],
        compact=True,
        header_fill=TEAL,
    )

    doc.add_heading("90초 시연 흐름", level=2)
    table(
        doc,
        ["시간", "화면", "보여줄 증거"],
        [
            ["0~20초", "고객 홈", "개인 기준선과 사유별 기여도"],
            ["20~43초", "맥락 확인 A", "같은 alertId가 정상종결"],
            ["43~65초", "맥락 확인 B", "기억하지 못함 → 행원 검토"],
            ["65~77초", "동의 가드", "가족 알림 미발송과 감사로그"],
            ["77~90초", "행원 사건큐", "근거·질문·조치·기록 초안, 승인 전 무조치"],
        ],
        [1350, 2100, 5910],
        compact=True,
    )
    small_note(doc, "데모는 로그인·실제 계좌연결·외부 API 키 없이 즉시 실행한다. LLM 장애 시에도 템플릿 설명으로 동일 흐름을 완주한다.")


def spring_architecture(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "06 Architecture", "Java Spring 기반 모듈형 모놀리스", "25일 MVP에서는 분산시스템보다 경계가 선명한 한 애플리케이션이 유리하다")

    callout(
        doc,
        "권장 스택",
        "Java 21 · Spring Boot · Spring Security · Spring Data JPA · Scheduler · PostgreSQL · Spring AI(선택) · Thymeleaf/HTMX 또는 익숙한 프론트엔드",
        fill=MINT,
        accent=TEAL,
    )

    doc.add_heading("모듈 경계", level=2)
    table(
        doc,
        ["모듈", "책임", "핵심 산출물"],
        [
            ["ingest", "합성 거래 입력·검증·멱등처리", "transaction, recurring_obligation"],
            ["feature", "7·30·90일 특징과 개인 기준선", "baseline_feature, feature_snapshot"],
            ["detection", "MAD·신규성·누락·추세·복합규칙", "anomaly_signal(reason_code, contribution)"],
            ["case", "신호 병합·맥락 재평가·상태전이", "protection_case, risk_assessment(version)"],
            ["consent", "분석·연락인 동의와 철회 가드", "consent_policy, trusted_contact_policy"],
            ["action", "공식 보호수단 조건·버전 관리", "action_catalog, action_plan"],
            ["copilot", "쉬운 설명·질문·기록 초안·RAG", "draft, evidence_match"],
            ["audit", "동의·모델·규칙·조회·승인 이력", "audit_event, staff_review"],
        ],
        [1550, 3450, 4360],
        compact=True,
    )

    doc.add_heading("핵심 API", level=2)
    table(
        doc,
        ["고객·데모", "행원·검토"],
        [
            ["POST /simulations/{scenario}/ingest", "GET /staff/cases"],
            ["GET /customers/{id}/alerts", "GET /cases/{id}"],
            ["POST /alerts/{id}/context", "POST /cases/{id}/action-plan"],
            ["GET /alerts/{id}/audit", "POST /cases/{id}/review"],
        ],
        [4680, 4680],
        compact=True,
        header_fill=TEAL,
    )
    small_note(doc, "MVP에서 제외: Kafka, Kubernetes, 별도 Python ML 서비스, 멀티에이전트, 운영 중 자동학습. pgvector는 공식문서 RAG가 실제로 필요할 때만 추가한다.")


def deployment_and_network(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "07 Deployment", "망분리를 전제로 한 생산 아키텍처", "거래 원문과 식별자는 내부에 남기고, AI에는 최소화된 사건요약만 전달한다")

    doc.add_heading("권장 데이터 흐름", level=2)
    add_process(
        doc,
        [
            ("고객·행원 채널", "인증·RBAC·동의 확인"),
            ("내부 API", "Spring 앱과 사건 오케스트레이션"),
            ("내부 데이터", "원거래·식별자·특징·감사로그"),
            ("승인 AI 중계", "비식별 요약·허용목록·출력필터"),
            ("내부/사설 LLM", "질문·설명·기록 초안만"),
        ],
        colors=[LIGHT_GRAY, SKY, MINT, PALE_GOLD, SKY],
    )

    table(
        doc,
        ["영역", "허용 데이터", "통제"],
        [
            ["내부 보호영역", "거래 원문, 고객ID, 동의, 특징, 사건, 감사로그", "암호화, 최소권한, 직무분리, 보존·삭제 정책"],
            ["AI 중계영역", "가명 사건ID, 구조화 reasonCode, 평소값·현재값의 최소 요약", "PII 차단, 프롬프트 템플릿, 호출 허용목록, 타임아웃"],
            ["LLM", "자유텍스트가 아닌 최소 구조화 요약과 승인된 정책 문서", "학습·로그 저장 금지 조건, 출력근거 검사, 실패 시 템플릿"],
            ["행원망", "AI 초안과 공식근거, 수정·승인·종결 UI", "실질 검토, override, 이의제기, kill switch"],
        ],
        [1750, 3860, 3750],
        compact=True,
    )

    doc.add_heading("MVP와 상용화의 배치 차이", level=2)
    table(
        doc,
        ["단계", "배치", "데이터"],
        [
            ["공모전", "단일 Spring 앱 + PostgreSQL, 공개 URL", "합성데이터만; 문자·코어뱅킹 실행 커넥터 없음"],
            ["은행 PoC", "은행 전용 VPC/내부망, 승인된 AI 중계 또는 사설 모델", "가상·비식별 사건부터 검증"],
            ["운영", "금융회사 내부/전용 배치, 보안·준법 승인 후 제한 연계", "목적별 최소 데이터, 실행은 기존 업무시스템과 권한자가 담당"],
        ],
        [1500, 3830, 4030],
        compact=True,
        header_fill=TEAL,
    )
    source_paragraph(
        doc,
        [
            ("금융위 망분리 정책 2026", "https://www.fsc.go.kr/no010101/86972"),
            ("금융보안원 AI 보안 안내서", "https://www.fsec.or.kr/bbs/detail?bbsNo=11977&menuNo=222"),
            ("개인정보위 생성형 AI 안내", "https://www.pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=11410"),
        ],
    )


def data_and_evaluation(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "08 Evaluation", "치매 진단 정확도가 아니라 보호업무 성과를 검증", "합성 골든셋으로 개인 기준선과 맥락 확인의 추가가치를 비교한다")

    doc.add_heading("합성 골든셋", level=2)
    add_bullet(doc, "20~30개 가상 페르소나에 은퇴소득, 정기납부, 송금, 현금, 구독 패턴을 부여하고 12개월 거래를 생성한다.")
    add_bullet(doc, "정상 입원비, 신규 수취인 반복고액+미납, 연락인 미동의의 3개 핵심 시나리오를 고정한다.")
    add_bullet(doc, "90일 데이터에서는 임시 기준선과 낮은 신뢰도를 표시하고, 12개월 이력이 연결되면 신뢰도를 높인다.")
    add_bullet(doc, "생활맥락 답변은 진실로 간주하지 않고 우선순위 감점 또는 추가검토 신호로만 사용한다.")

    doc.add_heading("A/B/C 비교", level=2)
    table(
        doc,
        ["실험군", "구성", "검증 질문"],
        [
            ["A. 규칙만", "전 고객 공통 금액·횟수 임계값", "개인차 때문에 얼마나 많은 오탐이 생기는가"],
            ["B. 개인 기준선", "규칙 + 중앙값/MAD + 신규성 + 정기납부 주기", "개인화가 오탐을 얼마나 줄이는가"],
            ["C. 기준선+맥락", "B + 고객의 정상 생활변화 확인", "맥락 확인이 정상사건을 얼마나 안전하게 종결하는가"],
        ],
        [1800, 3700, 3860],
        compact=True,
    )

    doc.add_heading("내부 MVP 통과 기준", level=2)
    table(
        doc,
        ["범주", "지표", "통과 기준"],
        [
            ["작동", "골든 시나리오 재현", "3개 시나리오 상태전이 100% 재현"],
            ["안전", "미동의 연락·승인 전 실행", "외부 발송·거래 실행 0건"],
            ["설명", "평소값·현재값·기여사유", "모든 경보에 세 요소 표시"],
            ["탐지", "scenario recall·false alerts/person-month", "A/B/C를 동일 데이터에서 비교해 수치 공개"],
            ["업무", "사건검토·기록 작성시간", "수동 대비 변화와 AI 초안 수정률 측정"],
            ["RAG", "공식근거 정확성·금지행동률", "근거 없는 답변은 거절, 금지행동 0%"],
        ],
        [1500, 3360, 4500],
        compact=True,
        header_fill=TEAL,
    )
    callout(doc, "측정하지 않는 것", "치매 진단 정확도, 인지저하 확률, 연령만으로 계산한 위험점수는 만들지도 평가하지도 않는다.", fill=PALE_RED, accent=RED)


def compliance_and_rights(doc: Document) -> None:
    # The evaluation page naturally fills the preceding page. Omitting an
    # explicit break here avoids a blank page in LibreOffice pagination.
    title_block(doc, "09 Governance", "법·윤리·소비자보호를 기능으로 구현", "동의, 설명, 사람의 검토, 감사로그를 화면과 API 상태로 증명한다")

    table(
        doc,
        ["기준", "안심리듬 설계 대응"],
        [
            ["개인정보보호법 민감정보", "거래로 치매·인지저하 라벨이나 확률을 추론·저장하지 않음. 건강정보는 MVP에서 수집하지 않음"],
            ["자동화된 결정", "알림·추천만 제공하고 불리한 자동조치 배제. 설명·이의제기·사람 재검토 경로와 override 기록"],
            ["신용정보법·마이데이터", "MVP는 합성데이터. 상용은 금융회사 또는 허가 마이데이터사 제휴와 적법한 전송요구 구조"],
            ["AI기본법", "AI 사용 사전고지. 거래제한·대출 등으로 확장 시 고영향 여부 사전검토와 위험관리·문서화"],
            ["2026 금융 AI 가이드라인", "AI는 보조수단, 최종 의사결정과 책임은 임직원. 위험등급·전 생애주기·중단수단 운영"],
            ["금융보안원 AI 보안", "자산목록, 입력·출력필터, 프롬프트 인젝션·오염·추출·공급망 시험, 모니터링"],
        ],
        [2600, 6760],
        compact=True,
    )

    doc.add_heading("신뢰연락인 안전설계", level=2)
    add_bullet(doc, "제공받는 자·목적·항목·보유기간·거부효과를 알린 별도·구체적·철회 가능한 동의를 받는다.")
    add_bullet(doc, "기본값은 본인만. 연락인에게 전체 거래내역·상호·점수·‘치매 의심’을 보내지 않고 최소 확인신호만 제공한다.")
    add_bullet(doc, "연락인이 가해자일 가능성에 대비해 차단·교체·복수지정·직원전용 모드·조회 및 통보 감사로그를 둔다.")
    add_bullet(doc, "신뢰연락인은 후견인·법정대리인이 아니며 송금·해지·동결·전체내역 조회 권한을 자동 취득하지 않는다.")

    source_paragraph(
        doc,
        [
            ("개인정보보호법 제23조", "https://www.law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1027416043"),
            ("자동화된 결정 안내", "https://m.pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=10611"),
            ("AI기본법", "https://www.law.go.kr/LSW/lsInfoP.do?lsId=014820"),
            ("금융분야 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142"),
            ("신용정보법", "https://www.law.go.kr/LSW/lsInfoP.do?ancYnChk=0&lsId=001540"),
        ],
    )
    small_note(doc, "본 문서는 공모전 기획안이며 법률자문이 아니다. 실제 적용 시 금융회사 준법·보안·개인정보보호 부서와 최신 원문을 재검토한다.")


def build_plan(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "10 Build Plan", "25일 안에 완성하는 범위", "기능을 더하는 대신 세 화면의 재현성·안전성·설명력을 끝까지 만든다")

    doc.add_heading("개발 순서", level=2)
    table(
        doc,
        ["기간", "구현", "완료 기준"],
        [
            ["D1~3", "데이터 계약·상태기계·합성 페르소나·화면 와이어프레임", "고정 alertId와 A/B 기대상태 정의"],
            ["D4~8", "ingest·feature·MAD·신규성·누락·복합규칙", "개인 기준선과 reasonCode API"],
            ["D9~13", "고객 홈·맥락 확인·동의가드", "A 정상종결, B 행원검토, 연락인 차단"],
            ["D14~17", "사건큐·상세·조치카탈로그·감사로그", "행원 승인 전 실행 이벤트 0건"],
            ["D18~20", "템플릿 설명·선택형 LLM/RAG·출력가드", "API 키 제거·오류 시 fallback"],
            ["D21~23", "골든셋 회귀시험·접근성·모바일·성능", "세 화면 90초 완주와 테스트 자동화"],
            ["D24~25", "배포·기능명세·발표영상·URL 점검", "새 브라우저에서 로그인 없이 즉시 실행"],
        ],
        [1300, 4400, 3660],
        compact=True,
    )

    doc.add_heading("범위 안 / 범위 밖", level=2)
    table(
        doc,
        ["MVP에서 반드시 구현", "향후 실증으로 넘김"],
        [
            ["3화면, 3시나리오, 개인 기준선, 사유 기여도", "실제 은행·마이데이터·코어뱅킹 연동"],
            ["맥락 재평가, 연락인 동의가드, 행원 승인 상태", "실제 문자·가족연락·한도변경·거래보류"],
            ["공식 조치카탈로그, 템플릿/LLM fallback, 감사로그", "자체 LLM 학습, 온프레미스 GPU, 대규모 음성·OCR"],
            ["rule-only vs baseline vs baseline+context 평가", "멀티에이전트·Kafka·K8s·별도 Python 서비스"],
        ],
        [4680, 4680],
        compact=True,
        header_fill=TEAL,
    )

    doc.add_heading("완성도를 지키는 기술 선택", level=2)
    add_bullet(doc, "로그인 없는 데모 세션은 서버에 실제 개인정보를 남기지 않고 고정 fixture로 재시작한다.")
    add_bullet(doc, "모든 상태전이는 alertId·ruleVersion·contextVersion·actor·timestamp를 감사로그에 남긴다.")
    add_bullet(doc, "LLM이 없어도 탐지·맥락분기·정책추천·행원큐가 동작하도록 결정과 문장생성을 분리한다.")
    add_bullet(doc, "화면명·상태코드·수치가 기획서, 기능명세, 실제 코드에서 동일하도록 자동 회귀시험한다.")


def business_and_roadmap(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "11 Adoption", "주 구매자를 선명하게 하고 실증으로 확장", "은행 소비자보호·FDS 후속업무팀이 첫 번째 고객이다")

    doc.add_heading("도입 가치", level=2)
    table(
        doc,
        ["이해관계자", "구매·사용 이유", "실증 KPI"],
        [
            ["소비자보호/FDS팀", "경보 후 확인·설명·기록·재연락 시간을 줄이고 품질을 표준화", "사건 검토시간, 기록 작성시간, 재연락 누락률"],
            ["고객", "정상 생활변화를 직접 설명하고 불필요한 통제를 피하면서 필요한 도움을 선택", "경보 이해도, 정상종결률, 이의제기 처리시간"],
            ["준법·보안", "동의·근거·모델버전·직원 승인·종결사유를 감사 가능하게 관리", "미동의 제공 0건, 승인 전 실행 0건, 로그 완전성"],
        ],
        [1900, 4550, 2910],
        compact=True,
    )

    doc.add_heading("단계별 확장", level=2)
    add_process(
        doc,
        [
            ("1. 공모전", "합성데이터·3화면·90초 데모"),
            ("2. 내부 PoC", "비식별 사건과 직원 사용성"),
            ("3. 제한 실증", "동의 고객·단일 금융회사·읽기 전용"),
            ("4. 제휴", "허가 마이데이터사/금융회사 연계"),
            ("5. 확장", "신탁·후견·공공상담 연결"),
        ],
        colors=[SKY, MINT, PALE_GOLD, SKY, LIGHT_GRAY],
    )

    doc.add_heading("핵심 리스크와 대응", level=2)
    table(
        doc,
        ["리스크", "대응"],
        [
            ["치매 낙인·건강정보 추론", "중립 브랜드와 금융변화 용어; 질병 라벨·점수 미생성"],
            ["오탐과 고객 불안", "개인 기준선, 생활맥락, 이의제기, 정상종결 피드백"],
            ["가족에 의한 금융착취", "최소정보·철회·차단·복수지정·직원전용 검토"],
            ["근거 없는 AI 답변", "공식문서 허용목록, 인용, 출력검사, 템플릿 fallback"],
            ["금융권 배치 제약", "Spring 모듈형 모놀리스와 내부/전용 배치, AI 중계 분리"],
        ],
        [2600, 6760],
        compact=True,
        header_fill=TEAL,
    )

    callout(
        doc,
        "최종 제안",
        "안심리듬은 고객마다 다른 금융생활의 리듬을 이해하고, 같은 경보라도 생활맥락에 따라 정상 변화는 종결하며, 설명이 필요한 사건만 행원에게 연결하는 인간 중심 AI 금융안전 플랫폼이다.",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def sources(doc: Document) -> None:
    new_page(doc)
    title_block(doc, "Sources", "공식 근거와 최신 정책", "공식자료를 우선하고 조사 사실과 서비스 제안을 구분했다")

    groups = [
        (
            "연구·문제 배경",
            [
                ("Nicholas et al., JAMA Internal Medicine, 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
                ("Triebel et al., Neurology, 2009", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
                ("UK bank data study, JAMA Network Open, 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
            ],
        ),
        (
            "대한민국 금융·공공 대응",
            [
                ("금융위원회 FDS·ASAP 운영자료", "https://www.fsc.go.kr/po010102/86997"),
                ("금융위원회 금융거래 안심차단", "https://www.fsc.go.kr/no010101/85644"),
                ("금융위원회 마이데이터 2.0", "https://fsc.go.kr/po010101/84780"),
                ("국민연금공단 치매안심 재산관리 시범사업", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
                ("대한민국 법원 성년후견 안내", "https://www.scourt.go.kr/nm/min_3/min_3_12/index.html"),
            ],
        ),
        (
            "AI·개인정보·보안",
            [
                ("금융위원회 2026 금융분야 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142"),
                ("금융위원회 2026 망분리 정책", "https://www.fsc.go.kr/no010101/86972"),
                ("금융보안원 금융분야 인공지능 보안 안내서", "https://www.fsec.or.kr/bbs/detail?bbsNo=11977&menuNo=222"),
                ("개인정보보호법 제23조", "https://www.law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1027416043"),
                ("개인정보위 자동화된 결정 안내서", "https://m.pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=10611"),
                ("개인정보위 생성형 AI 개인정보 처리 안내", "https://www.pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=11410"),
                ("AI기본법", "https://www.law.go.kr/LSW/lsInfoP.do?lsId=014820"),
                ("신용정보법", "https://www.law.go.kr/LSW/lsInfoP.do?ancYnChk=0&lsId=001540"),
            ],
        ),
        (
            "해외 원칙·유사 서비스",
            [
                ("FINRA Trusted Contact", "https://www.finra.org/rules-guidance/notices/22-31"),
                ("FINRA Rule 2165", "https://www.finra.org/rules-guidance/rulebooks/finra-rules/2165"),
                ("FCA Vulnerable Customers FG21/1", "https://www.fca.org.uk/publications/finalised-guidance/guidance-firms-fair-treatment-vulnerable-customers"),
                ("EverSafe", "https://www.eversafe.com/for-families/"),
                ("Carefull", "https://getcarefull.com/"),
            ],
        ),
        (
            "공모전",
            [("2026 Finance AI Challenge", "https://daker.ai/public/hackathons/2026-finance-ai-challenge")],
        ),
    ]

    for group, links in groups:
        doc.add_heading(group, level=2)
        for label, url in links:
            p = doc.add_paragraph(style="Small")
            p.paragraph_format.left_indent = Inches(0.15)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run("• ")
            set_font(r, size=8.2, color=TEAL)
            hyperlink(p, label, url, size=8.2)

    p = doc.add_paragraph(style="Small")
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("주의: 본 제안서는 공모전 기획을 위한 조사·서비스 설계 문서이며 의료·법률·투자자문을 제공하지 않는다. 실제 사업화 전에는 최신 법령·가이드라인·금융회사 내부정책과 전문가 검토가 필요하다.")
    set_font(r, size=8.3, italic=True, color=GRAY)


def update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def semantic_audit(doc_path: Path) -> None:
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    with ZipFile(doc_path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
        ns = {"w": W_NS}
        text = "".join(node.text or "" for node in root.findall(".//w:t", ns))

    required = [
        "안심리듬은 치매를 예측하지 않는다",
        "BLOCKED_BY_CONSENT",
        "Spring Boot",
        "금융분야 AI 가이드라인",
        "CLOSED_NORMAL",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError("missing required content: " + ", ".join(missing))

    forbidden = ["치매 위험 78%로 판정", "AI가 자동 지급정지", "국내 최초·유일한"]
    present = [item for item in forbidden if item in text]
    if present:
        raise RuntimeError("forbidden phrasing present: " + ", ".join(present))


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "안심리듬 AI 금융안전 플랫폼 개정 제안서"
    doc.core_properties.subject = "2026 금융 AI Challenge 출품용 개정안"
    doc.core_properties.author = "Ansim Rhythm Project Team"
    doc.core_properties.keywords = "안심리듬, 금융 AI, 금융생활 변화, FDS, 행원 코파일럿, Spring Boot, 망분리"

    cover(doc)
    executive_summary(doc)
    problem_and_evidence(doc)
    landscape_and_positioning(doc)
    product_journey(doc)
    ai_design(doc)
    demo_design(doc)
    spring_architecture(doc)
    deployment_and_network(doc)
    data_and_evaluation(doc)
    compliance_and_rights(doc)
    build_plan(doc)
    business_and_roadmap(doc)
    sources(doc)

    update_fields_on_open(doc)
    doc.save(OUT)
    audit(OUT)
    semantic_audit(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
