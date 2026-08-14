from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "Dementia_Money_AI_Financial_Safety_Report_2026.docx"

FONT = "AppleGothic"
NAVY = "17324D"
BLUE = "2B6F9F"
TEAL = "2C8C82"
SKY = "EAF4F8"
MINT = "EAF6F3"
GOLD = "D69B2D"
PALE_GOLD = "FFF5DF"
RED = "A64242"
PALE_RED = "FAECEC"
INK = "202B35"
GRAY = "5F6B75"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D9E0E5"
WHITE = "FFFFFF"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def set_font(run, *, size=None, bold=None, italic=None, color=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in kwargs[edge]:
                element.set(qn("w:" + key), str(kwargs[edge][key]))


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=140):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height = None
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, cell in enumerate(row.cells):
            width = widths[i]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=GRAY)


def hyperlink(paragraph, text, url, *, color=BLUE, size=8.5):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        r_fonts.set(qn("w:" + attr), FONT)
    r_pr.append(r_fonts)
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), color)
    r_pr.append(clr)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_repeatable_header_footer(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.78)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.88)
        section.right_margin = Inches(0.88)
        section.header_distance = Inches(0.32)
        section.footer_distance = Inches(0.32)
        header = section.header
        p = header.paragraphs[0]
        p.text = "치매머니 | AI 금융안전 플랫폼 조사·서비스 기획"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        if p.runs:
            set_font(p.runs[0], size=8, color=GRAY)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("2026.08.13  |  ")
        set_font(r, size=8.5, color=GRAY)
        add_page_field(p)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.4)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.27
    pf.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 14, 7),
        ("Heading 2", 14, BLUE, 10, 5),
        ("Heading 3", 11.5, TEAL, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Inches(0.34)
        style.paragraph_format.first_line_indent = Inches(-0.17)
        style.paragraph_format.space_after = Pt(3.5)
        style.paragraph_format.line_spacing = 1.2

    for name in ("Source", "Lead", "Small", "Table Text"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    source = styles["Source"]
    source.font.name = FONT
    source._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    source.font.size = Pt(8.2)
    source.font.color.rgb = RGBColor.from_string(GRAY)
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.13
    lead = styles["Lead"]
    lead.font.name = FONT
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    lead.font.size = Pt(12.5)
    lead.font.bold = True
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.25
    small = styles["Small"]
    small.font.name = FONT
    small._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    small.font.size = Pt(8.7)
    small.font.color.rgb = RGBColor.from_string(GRAY)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.15
    table_text = styles["Table Text"]
    table_text.font.name = FONT
    table_text._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    table_text.font.size = Pt(9.1)
    table_text.font.color.rgb = RGBColor.from_string(INK)
    table_text.paragraph_format.space_after = Pt(1)
    table_text.paragraph_format.line_spacing = 1.13


def new_page(doc):
    doc.add_page_break()


def title_block(doc, kicker, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(kicker.upper())
    set_font(r, size=9, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_font(r, size=24, bold=True, color=NAVY)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        r = p.add_run(subtitle)
        set_font(r, size=11.7, color=GRAY)


def callout(doc, label, text, *, fill=SKY, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=140)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell,
                    top={"val": "single", "sz": 12, "color": accent},
                    left={"val": "single", "sz": 12, "color": accent},
                    bottom={"val": "single", "sz": 12, "color": accent},
                    right={"val": "single", "sz": 12, "color": accent})
    p = cell.paragraphs[0]
    p.style = doc.styles["Table Text"]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "  ")
    set_font(r, size=9.5, bold=True, color=accent)
    r = p.add_run(text)
    set_font(r, size=10.1, bold=True, color=NAVY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_bullet(doc, text, *, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)
    return p


def table(doc, headers, rows, widths, *, header_fill=NAVY, compact=False, alignments=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    set_table_geometry(t, widths, indent=140)
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.style = doc.styles["Table Text"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_font(r, size=8.9 if compact else 9.2, bold=True, color=WHITE)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            if alignments:
                p.alignment = alignments[i]
            r = p.add_run(str(value))
            set_font(r, size=8.5 if compact else 9.1)
    set_table_geometry(t, widths, indent=140)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def source_paragraph(doc, items):
    p = doc.add_paragraph(style="Source")
    r = p.add_run("출처: ")
    set_font(r, size=8.2, bold=True, color=GRAY)
    for i, (label, url) in enumerate(items):
        if i:
            r = p.add_run(" · ")
            set_font(r, size=8.2, color=GRAY)
        hyperlink(p, label, url)
    return p


def add_process(doc, steps, *, colors=None):
    colors = colors or [SKY] * len(steps)
    widths = [int(9360 / len(steps))] * len(steps)
    widths[-1] += 9360 - sum(widths)
    t = doc.add_table(rows=1, cols=len(steps))
    set_table_geometry(t, widths, indent=140)
    set_repeat_table_header(t.rows[0])
    for i, (title, body) in enumerate(steps):
        c = t.cell(0, i)
        shade_cell(c, colors[i])
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        border_color = MID_GRAY
        set_cell_border(c,
                        top={"val": "single", "sz": 8, "color": border_color},
                        left={"val": "single", "sz": 8, "color": border_color},
                        bottom={"val": "single", "sz": 8, "color": border_color},
                        right={"val": "single", "sz": 8, "color": border_color})
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        set_font(r, size=9.1, bold=True, color=NAVY)
        p = c.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(body)
        set_font(r, size=8.2, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2026 금융 AI Challenge")
    set_font(r, size=11, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(56)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("치매머니")
    set_font(r, size=39, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Dementia Money")
    set_font(r, size=18, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("개인별 금융행동 변화 조기탐지와 선제 대응을 위한\nAI 금융안전 플랫폼 조사·서비스 기획 보고서")
    set_font(r, size=15.2, bold=True, color=INK)
    p.paragraph_format.line_spacing = 1.35
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_after = Pt(20)
    r = line.add_run("━" * 25)
    set_font(r, size=9, color=GOLD)
    callout(doc, "CORE PROMISE", "고객에게는 금융안전 조기경보를, 금융회사에는 보호업무 코파일럿을", fill=PALE_GOLD, accent=GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("조사 기준일  2026. 8. 13.")
    set_font(r, size=10, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("공유용 통합본 | 조사 사실·해석·서비스 제안을 구분해 작성")
    set_font(r, size=8.7, color=GRAY)


def add_exec_summary(doc):
    new_page(doc)
    title_block(doc, "Executive Summary", "한눈에 보는 치매머니", "진단이 아니라 조기경보, 차단이 아니라 단계적 보호")
    callout(doc, "한 문장 정의", "치매머니는 개인의 평소 금융생활과 다른 변화를 조기에 발견하고, 본인·신뢰연락인·금융회사가 피해 전에 함께 대응하도록 지원하는 AI 금융안전 플랫폼이다.")

    doc.add_heading("왜 필요한가", level=2)
    add_bullet(doc, "보이스피싱 FDS, 후견, 공공·민간 신탁 등 보호수단은 존재하지만 서로 분절돼 있다.")
    add_bullet(doc, "고객은 스스로 위험을 인지하고 적절한 제도를 찾아 신청해야 하며, 인지저하 이후에는 오히려 계약과 의사결정이 더 어려워질 수 있다.")
    add_bullet(doc, "행원은 FDS 경보 뒤 확인 전화, 맥락 파악, 보호수단 안내, 상담기록과 사후관리를 수행해야 하므로 전수점검이 어렵다.")

    doc.add_heading("무엇을 해결하는가", level=2)
    table(doc,
          ["대상", "현재 어려움", "치매머니의 가치"],
          [
              ["고객", "이상 변화와 필요한 대응책을 스스로 판단하기 어려움", "개인 기준선 기반 조기경보와 맞춤 안전계획"],
              ["가족·조력자", "과도한 통제 없이 필요한 순간만 돕기 어려움", "동의 기반 최소권한 알림·공동확인"],
              ["금융회사", "경보 이후 반복 확인·상담·기록 업무 부담", "고위험 사건 선별과 보호업무 코파일럿"],
              ["공공기관", "후견·신탁이 필요한 대상을 뒤늦게 만남", "적합한 공공 보호체계로 조기 연결"],
          ], [1450, 3600, 4310])

    doc.add_heading("서비스의 경계", level=2)
    table(doc,
          ["하는 일", "하지 않는 일"],
          [["금융행동의 변화와 피해위험 신호를 설명한다", "거래내역만으로 치매를 진단하지 않는다"],
           ["고객과 행원의 대응 선택을 지원한다", "AI가 단독으로 계좌를 동결하거나 고객 능력을 판정하지 않는다"],
           ["보호수단과 상담기관을 연결한다", "법적 권한 없는 가족에게 계좌 통제권을 주지 않는다"]],
          [4680, 4680], header_fill=TEAL)


def add_problem_and_evidence(doc):
    new_page(doc)
    title_block(doc, "01 Problem", "문제 정의", "보호수단은 있지만, 위험을 알아차리고 행동으로 옮기는 연결층이 약하다")
    p = doc.add_paragraph(style="Lead")
    p.add_run("핵심 문제는 ‘제도가 전혀 없다’가 아니라, 고객의 위험 인지부터 적합한 대책 실행까지의 여정이 끊겨 있다는 점이다.")

    add_process(doc, [
        ("평소 금융생활", "본인도 변화를 잘 인식하지 못함"),
        ("이상 변화 발생", "중복·미납·낯선 송금이 흩어져 보임"),
        ("피해 또는 진단", "가족·은행이 뒤늦게 문제를 인지"),
        ("보호제도 탐색", "신탁·후견·차단제도를 직접 찾아 신청"),
        ("사후 관리", "기관별로 분절된 상담과 기록"),
    ], colors=[LIGHT_GRAY, PALE_GOLD, PALE_RED, SKY, MINT])

    doc.add_heading("고객 관점의 공백", level=2)
    add_bullet(doc, "예방서비스 대부분은 본인 또는 가족이 필요성을 먼저 인식하고 신청해야 한다.")
    add_bullet(doc, "여러 은행·카드·자동이체에 흩어진 작은 실수를 하나의 변화 추세로 보기 어렵다.")
    add_bullet(doc, "위험을 알아차려도 어떤 순서로 이체한도, 안심차단, 두낫콜, 신탁 또는 후견을 활용할지 판단하기 어렵다.")

    doc.add_heading("금융회사 관점의 공백", level=2)
    add_bullet(doc, "기존 FDS가 의심거래를 자동 선별하더라도 실제 사기인지 생활변화인지 확인하는 업무는 사람에게 남는다.")
    add_bullet(doc, "행원은 고객 연락, 확인 질문, 보호서비스 설명, 상담기록 작성, 후속연락을 반복해야 한다.")
    add_bullet(doc, "모든 고령 고객을 수작업으로 장기간 전수 모니터링하는 것은 현실적으로 어렵다.")

    callout(doc, "정확한 문제 문장", "대한민국에는 다양한 보호수단이 존재하지만, 개인별 금융생활 변화를 조기에 발견하고 고객·행원에게 적합한 대응을 연결하는 통합 지원체계가 부족하다.", fill=PALE_GOLD, accent=GOLD)

    new_page(doc)
    title_block(doc, "02 Evidence", "금융행동 변화에 대한 근거", "연관성은 존재하지만, 개인 한 명의 치매 진단 지표로 사용해서는 안 된다")
    p = doc.add_paragraph()
    p.add_run("연구에서는 인지저하·치매 전후에 연체, 계좌관리 실수, 지출통제 저하, 금융사기 취약성이 집단 수준에서 증가하는 경향이 관찰된다. 다만 은퇴, 입원, 사별, 이사, 여행, 소득변화도 같은 패턴을 만들 수 있으므로 단일 신호를 진단처럼 해석하면 안 된다.")

    table(doc,
          ["관찰된 변화", "근거 수준", "치매머니 적용 원칙"],
          [
              ["청구서 누락·연체", "상대적으로 강함", "핵심 변화 신호로 사용하되 소득·생활사건 확인"],
              ["계좌·명세서 관리 실수", "상대적으로 강함", "반복·지속 여부와 여러 신호의 결합을 평가"],
              ["과잉지출·현금흐름 악화", "부분 근거", "평소 소득·필수지출 대비 변화로 설명"],
              ["금융사기 취약성", "연관 근거 있음", "사기여부와 금융관리 변화는 별도로 구분"],
              ["중복결제·잊은 구독", "합리적 가설", "인지저하 특이 신호가 아니라 확인이 필요한 오류로 표시"],
              ["시간·지역·ATM 변화", "직접 근거 제한", "여행·입원·이사 등 생활맥락을 먼저 질문"],
              ["동일상품 반복가입", "직접 근거 제한", "불완전판매·판단취약성 보호 항목으로 활용"],
          ], [2350, 1740, 5270], compact=True)
    source_paragraph(doc, [
        ("JAMA Intern Med 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
        ("Neurology 2009", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
        ("JAMA Netw Open 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
        ("미국 NIA 안내", "https://www.nia.nih.gov/health/legal-and-financial-planning/managing-money-problems-people-dementia"),
    ])
    callout(doc, "표현 원칙", "‘치매 가능성 83%’가 아니라 ‘최근 평소와 다른 금융행동 3건이 확인되어 본인 확인이 필요합니다’라고 안내한다.", fill=PALE_RED, accent=RED)


def add_korea_status(doc):
    new_page(doc)
    title_block(doc, "03 Korea", "대한민국의 현재 대응", "사기형 FDS, 후견, 공공·민간 신탁, 사전 신청형 보호서비스가 각자 역할을 담당한다")
    table(doc,
          ["대응 수단", "작동 방식", "치매머니 관점의 한계"],
          [
              ["은행 FDS·ASAP", "접속·거래정보와 공유된 사기정보를 분석해 의심거래 탐지·지연·정지", "보이스피싱 중심이며 장기 금융관리 변화 전체를 보지는 않음"],
              ["치매안심재산관리서비스", "NPS가 신탁받은 현금성 자산을 계획에 따라 생활·의료·요양비로 지급", "신탁재산만 관리하며 기존 전 계좌 이상패턴 탐지 기능은 아님"],
              ["치매공공후견·성년후견", "법원이 정한 권한 안에서 재산·복지·의료·주거 의사결정을 지원", "법원 절차가 필요하고 실시간 조기탐지 수단은 아님"],
              ["고령자 지정인 알림·숙려", "일부 고위험 상품 가입 사실을 지정인에게 알리고 판매과정을 녹취·숙려", "특정 상품·시점에 한정되고 일상거래 변화는 보지 않음"],
              ["안심차단·지연이체", "비대면 계좌·여신·오픈뱅킹 차단, 지연이체, 수취인 지정 등을 사전 신청", "스스로 필요성을 인지하고 설정해야 하며 개인 기준선 AI가 아님"],
              ["민간 치매안심신탁", "사전 계약과 지정대리인에 따라 생활비·병원비 등을 지급", "맡긴 자산만 보호하고 발동조건·수수료·대리인 신뢰 문제가 있음"],
          ], [2150, 3540, 3670], compact=True)
    source_paragraph(doc, [
        ("금융위 FDS·ASAP", "https://www.fsc.go.kr/po010102/86997"),
        ("NPS 치매안심재산관리", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
        ("복지부 서비스 안내", "https://www.mohw.go.kr/menu.es?mid=a10712010600"),
        ("법원 성년후견", "https://www.scourt.go.kr/nm/min_3/min_3_12/index.html"),
        ("금융위 지정인 알림", "https://www.fsc.go.kr/no010101/73764"),
        ("금융위 안심차단", "https://www.fsc.go.kr/no010101/85644"),
    ])

    doc.add_heading("가장 중요한 2026년 변화", level=2)
    add_bullet(doc, "국민연금공단은 2026년 4월 22일부터 2년간 치매안심재산관리서비스 시범사업을 시작했다. 치매·경도인지장애 등으로 재산관리 어려움이 있거나 예상되는 사람을 대상으로 맞춤 재정계획과 지출 모니터링을 지원한다.")
    add_bullet(doc, "2026년 7월 7일 공식자료상 첫 계약을 포함해 총 4건으로, 실제 운영이 시작됐지만 아직 초기·소규모 단계다.")
    add_bullet(doc, "은행 FDS와 ASAP은 이미 큰 규모의 보이스피싱 탐지·차단 성과를 내고 있으므로 ‘한국에 이상거래 탐지가 없다’고 주장해서는 안 된다.")

    doc.add_heading("직접 경쟁 동향", level=2)
    callout(doc, "신한은행", "2026년 ‘치매고객 금융보호 TF’를 운영하고 고령층 특화 금융거래 이상징후 모니터링을 고도화 중이다. 따라서 ‘국내 최초’보다 여러 금융사의 장기 행동변화 분석과 탐지 이후 대응 연결을 차별점으로 제시해야 한다.", fill=PALE_GOLD, accent=GOLD)
    source_paragraph(doc, [
        ("신한은행 진행 보도", "https://www.mk.co.kr/news/special-edition/12088908"),
        ("신한그룹 치매머니 교육", "https://www.shinhangroup.com/kr/archive/press/detail/805"),
        ("NPS 첫 계약 현황", "https://m.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?hmpgBbsCd=BS20240145&hmpgCd=01&menuId=MN24000898&pageIndex=1&pstId=ZZ202600000000000780"),
    ])


def add_overseas(doc):
    new_page(doc)
    title_block(doc, "04 Global", "해외에서는 어떻게 돕고 있나", "공통점은 진단이 아니라 사전동의, 거래별 최소개입, 사람의 확인과 법적 대리의 분리다")
    table(doc,
          ["국가", "대표 방식", "작동 원리", "가져올 설계 원칙"],
          [
              ["미국", "신뢰연락인·거래별 임시보류", "증권사가 착취 의심 시 연락인에게 통지하고 특정 거래를 제한적으로 보류", "연락인과 법적 대리권을 분리하고 의심 연락인은 제외"],
              ["영국", "취약고객 지원·72시간 지급지연·경찰개입·사후배상", "직원 확인과 경찰 협력, 취약소비자 배상 보호를 결합", "탐지 후 사람의 질문·이의제기·피해회복까지 설계"],
              ["일본", "일상생활 자립지원·후견지원신탁", "소액 생활비는 지역 지원, 큰 자산은 법원 감독 아래 보호", "생활비 지갑과 보호자산 금고의 이중 구조"],
              ["싱가포르", "Money Lock·신뢰연락인·Safety Switch", "고객이 자금을 잠그고 중요 거래를 연락인에게 통지, 강한 해제 절차 적용", "정적 잠금과 개인 변화 탐지를 함께 제공"],
              ["호주·캐나다", "취약고객 지침·신뢰연락인·특정 거래 보류", "은행 코드 또는 증권 규칙으로 최소범위 개입", "광범위 자동동결보다 비례적·거래별 개입"],
          ], [900, 2450, 3080, 2930], compact=True)
    source_paragraph(doc, [
        ("FINRA Trusted Contact", "https://www.finra.org/rules-guidance/notices/22-31"),
        ("FINRA Rule 2165", "https://www.finra.org/rules-guidance/rulebooks/finra-rules/2165"),
        ("FCA FG21/1", "https://www.fca.org.uk/publication/finalised-guidance/fg21-1.pdf"),
        ("영국 지급지연", "https://www.gov.uk/government/news/new-powers-for-banks-to-combat-fraudsters"),
        ("일본 일상생활자립지원", "https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/hukushi_kaigo/seikatsuhogo/chiiki-fukusi-yougo/index.html"),
        ("싱가포르 Money Lock", "https://www.mddi.gov.sg/newsroom/measures-to-protect-singaporeans-against-online-scams/"),
    ])

    doc.add_heading("운영 중인 해외 민간서비스", level=2)
    table(doc,
          ["서비스", "핵심 기능", "한계"],
          [
              ["EverSafe·Carefull", "여러 계좌를 읽기 전용으로 연결해 개인 기준선 대비 이상징후와 가족 알림 제공", "자체 거래차단 권한과 공개된 독립 임상성능이 없음"],
              ["Sibstar", "치매 당사자용 선불카드에 가족이 일·월·ATM·채널 한도를 설정하고 실시간 알림", "AI 패턴탐지가 아니며 충전된 금액만 보호"],
              ["True Link", "선불카드의 업종·가맹점·ATM·구독·기간별 한도를 관리", "사전 규칙형이며 카드 밖 거래를 보호하지 못함"],
              ["AARP BankSafe", "은행 직원에게 착취 징후, 고객 질문, 신고·상향절차를 교육", "고객용 모니터링 제품이 아니며 직원 실행에 의존"],
          ], [1800, 4440, 3120], compact=True)
    source_paragraph(doc, [
        ("EverSafe", "https://www.eversafe.com/for-families/"),
        ("Carefull", "https://getcarefull.com/"),
        ("Sibstar", "https://www.sibstar.co.uk/"),
        ("True Link", "https://www.truelinkfinancial.com/"),
        ("AARP BankSafe", "https://www.aarp.org/pri/initiatives/banksafe/"),
    ])
    callout(doc, "해외 사례의 시사점", "탐지형 서비스와 통제형 서비스가 대체로 분리돼 있다. 치매머니는 설명 가능한 탐지 → 본인 재확인 → 동의된 조력자 공동확인 → 적합한 금융·공공 보호수단 연결을 한 흐름으로 묶는다.", fill=MINT, accent=TEAL)


def add_gap_and_definition(doc):
    new_page(doc)
    title_block(doc, "05 Opportunity", "기존 체계의 공백과 치매머니의 위치", "기존 FDS를 대체하지 않고, 사기형 탐지와 후견·신탁 사이의 조기 대응층을 만든다")
    table(doc,
          ["비교 항목", "기존 FDS", "신탁·후견", "치매머니"],
          [
              ["주요 목적", "사기·계정탈취 탐지", "맡긴 재산·법률행위 관리", "금융생활 변화 조기경보와 대응지원"],
              ["분석 단위", "금융사별 거래·접속정보", "계약·심판 대상 재산", "동의 기반 여러 금융사의 장기 개인 기준선"],
              ["작동 시점", "의심거래 발생 시", "사전 계약 또는 판단능력 저하 이후", "일상 변화가 누적되는 초기 단계"],
              ["탐지 대상", "낯선 송금·사기계좌·악성앱", "지급계획·대리권 범위", "중복·미납·구독·현금·상품가입·생활맥락"],
              ["후속 지원", "확인·차단·신고", "지급·대리·감독", "본인 확인, 안전계획, 행원 코파일럿, 제도 연결"],
          ], [1650, 2450, 2450, 2810], compact=True)

    doc.add_heading("치매머니가 맡는 연결층", level=2)
    add_process(doc, [
        ("기존 FDS", "사기형 위험 신호"),
        ("치매머니", "장기 금융생활 변화와 맥락"),
        ("본인·조력자", "확인과 공동대응"),
        ("행원", "최종판단과 보호조치"),
        ("신탁·후견", "지속적 재산관리"),
    ], colors=[LIGHT_GRAY, SKY, MINT, PALE_GOLD, LIGHT_GRAY])

    callout(doc, "포지셔닝", "은행별 ATM·송금 탐지를 넘어 결제·자동이체·구독·금융상품 데이터를 장기 관찰하고, 이상징후 탐지부터 고객의 사전대책 수립과 행원의 후속업무까지 지원하는 B2B2C AI 금융안전 플랫폼", fill=SKY, accent=BLUE)

    doc.add_heading("서비스 설계 원칙", level=2)
    add_number(doc, "치매를 진단하지 않고 금융안전상 확인이 필요한 변화를 설명한다.")
    add_number(doc, "AI는 선별·설명·추천을 담당하고 고객과 행원이 최종 결정한다.")
    add_number(doc, "생활비·의료비 등 필수거래를 해치지 않는 최소·단계적 개입을 사용한다.")
    add_number(doc, "신뢰연락인은 알림·열람·공동확인·법적 대리를 권한별로 분리한다.")
    add_number(doc, "탐지 이후 실제로 행동할 수 있는 국내 보호수단까지 연결한다.")


def add_customer_ai(doc):
    new_page(doc)
    title_block(doc, "06 Customer AI", "고객용 AI: 조기경보에서 안전계획까지", "경고만 보내지 않고, 왜 확인이 필요한지와 다음 행동을 쉬운 말로 제시한다")
    doc.add_heading("AI 1. 금융생활 변화 탐지", level=2)
    table(doc,
          ["관찰 영역", "개인 기준선", "예시 신호"],
          [
              ["송금", "평균 금액·횟수·수취인·시간", "새 수취인에게 단기간 반복·고액 송금"],
              ["납부", "공과금·보험·통신비 주기", "평소 납부하던 항목 누락·연체"],
              ["결제·구독", "정기결제·가맹점·금액", "중복결제, 새 구독, 해지 후 재결제"],
              ["현금", "ATM 위치·금액·빈도", "장기간 현금인출 증가 또는 낯선 지역 인출"],
              ["금융상품", "기존 보유상품·가입주기", "유사상품 반복가입, 단기간 고위험 상품 증가"],
          ], [1800, 3000, 4560], compact=True)

    doc.add_heading("AI 2. 생활맥락 확인과 위험 재평가", level=2)
    callout(doc, "고객 질문 예시", "‘최근 병원비 180만원이 새로 발생했습니다. 본인이 결제한 것이 맞나요?’  ‘최근 이사·여행·입원처럼 지출이 달라질 만한 일이 있었나요?’", fill=MINT, accent=TEAL)
    add_bullet(doc, "고객의 확인 결과와 생활사건을 기록해 같은 정상 변화가 반복 경보로 발생하지 않게 한다.")
    add_bullet(doc, "단일 이상거래보다 여러 신호의 결합, 지속기간, 자산 대비 금액을 함께 평가한다.")
    add_bullet(doc, "고객이 기억하지 못함을 선택하더라도 치매로 단정하지 않고 추가 확인과 상담을 제안한다.")

    doc.add_heading("AI 3. 맞춤형 안전계획", level=2)
    table(doc,
          ["위험 단계", "고객에게 제공하는 행동"],
          [
              ["안내", "거래를 쉬운 문장으로 설명하고 본인 의도를 확인"],
              ["주의", "중복결제·자동이체 정리, 두낫콜, 비대면 한도 축소를 추천"],
              ["높음", "선택형 쿨링오프와 동의된 신뢰연락인 공동확인을 제안"],
              ["긴급", "은행 상담, 112·1332 신고, 본인계좌 지급정지 안내"],
              ["지속 변화", "치매안심센터, NPS 재산관리, 민간신탁·임의후견 상담 연결"],
          ], [1700, 7660], compact=True, header_fill=TEAL)


def add_banker_ai(doc):
    new_page(doc)
    title_block(doc, "07 Banker Copilot", "행원용 AI: 전수조사가 아닌 위험기반 선별", "AI가 사건을 정리하고 질문·설명·기록 초안을 만들며, 행원이 최종 판단한다")
    doc.add_heading("반복업무를 어디까지 줄이는가", level=2)
    table(doc,
          ["업무 단계", "현재 부담", "AI 지원"],
          [
              ["경보 검토", "거래별 경보를 하나씩 확인", "여러 경보를 고객별 사건으로 묶고 우선순위화"],
              ["맥락 파악", "거래내역과 고객정보를 수동 조회", "기준선 대비 변화와 주요 근거를 요약"],
              ["본인 확인", "상황별 질문을 즉석에서 구성", "고령 고객용 쉬운 확인 질문과 설명문 생성"],
              ["조치 선택", "여러 제도와 가입조건을 검색", "검증된 보호서비스 목록에서 적합 순서 추천"],
              ["상담 기록", "통화·조치 내용을 반복 입력", "상담기록 초안과 후속 체크리스트 생성"],
              ["사후 관리", "재연락 일정과 미완료 조치 추적", "미확인 사건과 권장조치 완료 여부 추적"],
          ], [1800, 3200, 4360], compact=True)

    doc.add_heading("행원 화면 예시", level=2)
    callout(doc, "보호 우선순위 높음", "김○○ 고객: 최근 14일간 신규 수취인 3명에게 총 1,850만원 송금. 동일 금액 2회, 기존 공과금 1건 미납, 등록된 여행·입원 없음. 권장 순서: 본인 확인 → 사기의심 질문 → 동의 시 신뢰연락인 공동확인 → 지연이체·한도조정 안내.", fill=PALE_RED, accent=RED)

    doc.add_heading("업무 원칙", level=2)
    add_bullet(doc, "AI 추천은 ‘초안’으로 명시하고 행원이 승인·수정한다.")
    add_bullet(doc, "고객에게 보여줄 근거와 내부 위험특징을 분리해 불필요한 낙인을 방지한다.")
    add_bullet(doc, "법률·의료·금융상품 정보를 자유 생성하지 않고, 검증·갱신된 정책 지식베이스에서만 안내한다.")
    add_bullet(doc, "고객의 이의제기와 경보 해제 이유를 기록해 모델 개선과 감사에 활용한다.")


def add_flow_and_architecture(doc):
    new_page(doc)
    title_block(doc, "08 Service Flow", "전체 서비스 흐름", "탐지 → 맥락 확인 → 단계별 개입 → 사람의 판단 → 국내 보호체계 연결")
    add_process(doc, [
        ("1. 동의·연결", "합성데이터 또는 마이데이터 제휴"),
        ("2. 기준선 학습", "개인별 거래·납부·구독 추세"),
        ("3. 변화 탐지", "규칙+시계열+이상탐지 앙상블"),
        ("4. 본인 확인", "쉬운 설명과 생활맥락 질문"),
        ("5. 공동대응", "연락인·행원·공공기관 연결"),
    ], colors=[LIGHT_GRAY, SKY, SKY, MINT, PALE_GOLD])

    doc.add_heading("단계적 개입 설계", level=2)
    table(doc,
          ["단계", "트리거", "조치", "결정권자"],
          [
              ["L0 기록", "낮은 이상도", "대시보드에만 기록", "고객"],
              ["L1 확인", "단일·낮은 위험", "쉬운 설명과 본인 확인", "고객"],
              ["L2 계획", "반복·중간 위험", "맞춤 안전조치 체크리스트", "고객"],
              ["L3 공동확인", "미확인·고위험", "동의된 연락인 또는 행원 검토", "고객+사람"],
              ["L4 보호연결", "사기 임박·지속 변화", "법적 권한에 따른 보류·신고·신탁·후견 연결", "금융회사·공공기관"],
          ], [900, 2200, 3660, 2600], compact=True)

    doc.add_heading("AI·데이터 구조", level=2)
    table(doc,
          ["계층", "역할", "MVP 구현"],
          [
              ["데이터", "거래, 자동이체, 상품보유, 고객 확인, 생활사건", "가상 고객 6~12개월 합성데이터"],
              ["탐지", "개인 기준선과 규칙형 위험신호 결합", "통계 기준선+Isolation Forest 등 단순 모델 비교"],
              ["설명", "평소 대비 무엇이 얼마나 달라졌는지 제시", "근거 템플릿+생성형 AI 쉬운말 변환"],
              ["대응", "검증된 보호조치와 기관을 조건별 추천", "규칙 기반 정책 지식베이스+LLM 요약"],
              ["거버넌스", "동의, 권한, 로그, 이의제기, 모델감사", "역할별 화면과 감사로그 시뮬레이션"],
          ], [1600, 3880, 3880], compact=True, header_fill=TEAL)

    callout(doc, "중요", "실제 계좌동결·송금거절은 금융회사의 권한과 법적 근거가 필요하다. 공모전 MVP에서는 합성거래를 이용한 ‘보류 요청·본인 확인·공동확인’ 흐름으로 구현한다.", fill=PALE_RED, accent=RED)


def add_mvp(doc):
    new_page(doc)
    title_block(doc, "09 MVP", "MVP 범위와 검증 시나리오", "작동 가능한 웹서비스로 고객용 화면과 행원용 코파일럿을 함께 시연한다")
    doc.add_heading("필수 화면", level=2)
    table(doc,
          ["화면", "핵심 기능", "데모 포인트"],
          [
              ["고객 홈", "오늘의 금융안전 상태와 최근 변화", "치매 점수가 아닌 확인 필요 신호로 표현"],
              ["이상거래 상세", "평소 기준선·변화 이유·본인 확인", "비교 가능한 숫자와 쉬운 설명"],
              ["안전계획", "개인별 보호조치 체크리스트", "어카운트인포·두낫콜·안심차단·상담 연결"],
              ["신뢰연락인", "알림·열람·공동확인 권한 설정", "복수 연락인, 철회, 열람로그"],
              ["행원 대시보드", "사건 우선순위·근거·고객상태", "전수검토 대신 위험기반 큐"],
              ["행원 코파일럿", "질문·설명·조치·상담기록 초안", "사람 승인 전에는 조치가 실행되지 않음"],
          ], [1800, 3650, 3910], compact=True)

    doc.add_heading("대표 시나리오", level=2)
    table(doc,
          ["시나리오", "발생 데이터", "시스템 반응"],
          [
              ["반복송금", "같은 날 동일 수취인·동일 금액 2회", "본인에게 중복 가능성 확인 후 취소요청 안내"],
              ["보이스피싱", "신규 수취인에게 자산 대비 고액을 다회 이체", "행원 우선검토와 동의된 연락인 공동확인"],
              ["잊은 구독", "새 정기결제 3건과 사용 없는 장기구독", "어카운트인포 조회·해지 체크리스트"],
              ["생활변화", "입원 후 병원·간병 지출 급증", "정상 사유 확인 후 경보 억제, 생활사건 기록"],
              ["장기 관리저하", "공과금 미납·ATM 증가·상품 중복이 3개월 지속", "치매안심센터·신탁 상담을 단계적으로 제안"],
          ], [1700, 3670, 3990], compact=True)

    doc.add_heading("검증 지표", level=2)
    add_bullet(doc, "탐지율·오탐률과 고객 1인당 월평균 경보 수")
    add_bullet(doc, "행원의 사건 검토시간과 상담기록 작성시간 감소율")
    add_bullet(doc, "고객의 경보 이해도와 권장 보호조치 완료율")
    add_bullet(doc, "본인 확인부터 사건 종결까지 걸린 시간")
    add_bullet(doc, "연령·성별·디지털 이용수준별 성능 편차와 이의제기 처리시간")


def add_safety(doc):
    new_page(doc)
    title_block(doc, "10 Trust & Safety", "법적·윤리적 안전장치", "금융보호를 명분으로 고객의 자기결정권과 개인정보를 침해하지 않도록 설계한다")
    table(doc,
          ["위험", "안전장치"],
          [
              ["치매 낙인·오진", "치매 확률·인지점수 금지, 금융행동 변화와 확인 필요성만 표현"],
              ["AI 자동차단", "AI는 추천까지만 수행, 고객 확인·행원 승인·법적 근거를 거쳐 조치"],
              ["과도한 개인정보 수집", "목적별 최소수집, 명시적 동의, 보유기간·삭제·철회, 암호화와 접근로그"],
              ["보호자 권한 남용", "알림/제한열람/공동확인/법적대리 분리, 복수 연락인, 즉시 철회와 감사로그"],
              ["연령차별", "나이만으로 위험판단·거래제한 금지, 본인 기준선과 상황정보 중심"],
              ["생성형 AI 환각", "검증된 정책 지식베이스만 사용, 출처·적용조건 표시, 행원 승인"],
              ["오탐·미탐", "다중신호·지속기간·생활맥락 결합, 이의제기·경보해제, 정기 성능감사"],
              ["필수생활 침해", "생활비·의료비·정기청구를 구분하고 최소범위·최단기간 개입"],
          ], [2400, 6960], compact=True)

    doc.add_heading("권한 모델", level=2)
    table(doc,
          ["역할", "허용 권한", "금지 또는 별도 근거 필요"],
          [
              ["고객", "동의·확인·철회·이의제기·연락인 지정", "없음"],
              ["신뢰연락인", "경보 수신, 허용된 범위의 제한 열람, 공동확인", "송금·해지·계좌통제"],
              ["행원", "사건 검토, 고객 확인, 회사 정책에 따른 조치", "AI 결과만으로 능력상실 판정"],
              ["후견인·위임인", "법원 심판·위임계약 범위의 대리", "부여 범위를 벗어난 재산처분"],
              ["AI", "탐지·설명·질문·추천·기록 초안", "진단·법적판단·독립적 거래차단"],
          ], [1700, 3850, 3810], compact=True, header_fill=TEAL)

    callout(doc, "사업화 전 필수 검토", "개인신용정보 처리, 마이데이터·오픈뱅킹 API 자격과 제휴, 자동화된 결정의 설명·이의제기, 금융회사 내부통제와 사고책임, 의료기기 오인 가능성을 법률·준법 전문가와 검토해야 한다.", fill=PALE_RED, accent=RED)


def add_business_and_roadmap(doc):
    new_page(doc)
    title_block(doc, "11 Business", "도입 모델과 단계별 로드맵", "초기에는 합성데이터 MVP로 효과를 증명하고, 이후 금융회사 제휴형 B2B2C로 확장한다")
    doc.add_heading("권장 사업모델", level=2)
    table(doc,
          ["고객", "제공가치", "수익·도입 방식"],
          [
              ["은행·카드·증권", "고위험 고객 선별, 상담업무 절감, 소비자보호 강화", "SaaS·API 또는 화이트라벨 라이선스"],
              ["보험·자산관리", "상품 중복·판단취약성 보호와 가족 공동관리", "고객보호 부가서비스·기관구독"],
              ["지자체·공공기관", "신탁·후견·치매안심센터 필요대상 조기연계", "실증사업·공공조달·위탁운영"],
              ["개인·가족", "다계좌 금융안전 대시보드와 맞춤형 안전계획", "금융기관 제공을 기본으로 하되 선택형 프리미엄"],
          ], [1900, 4050, 3410], compact=True)

    doc.add_heading("단계별 추진", level=2)
    add_process(doc, [
        ("1단계 공모전", "합성데이터·웹 MVP·사용성 검증"),
        ("2단계 실증", "금융사 1곳·상담직원·가상사건 파일럿"),
        ("3단계 제휴", "마이데이터·정책 API·보호업무 통합"),
        ("4단계 확장", "다기관·공공신탁·후견 연계"),
    ], colors=[SKY, MINT, PALE_GOLD, LIGHT_GRAY])

    doc.add_heading("공모전에서 입증할 핵심", level=2)
    add_bullet(doc, "불필요한 경보를 억제하면서 설명 가능한 금융생활 변화를 찾는가.")
    add_bullet(doc, "고객이 경보를 이해하고 실제 보호조치를 완료하도록 돕는가.")
    add_bullet(doc, "행원의 검토·설명·기록 시간을 줄이되 인간의 최종판단을 보존하는가.")
    add_bullet(doc, "신뢰연락인과 금융회사의 권한을 최소화하고 모든 행동을 감사할 수 있는가.")

    callout(doc, "최종 제안 문장", "치매머니는 고객의 금융행동 변화를 조기에 발견해 피해 전에 맞춤형 안전계획을 제시하고, 금융회사에는 고위험 고객 선별부터 상담·보호조치·사후관리까지 지원하는 인간 중심 AI 금융보호 플랫폼이다.", fill=MINT, accent=TEAL)


def add_claims_and_sources(doc):
    new_page(doc)
    title_block(doc, "Appendix A", "과장 방지를 위한 표현 가이드", "사실, 조사상 해석, 서비스 제안을 구분한다")
    table(doc,
          ["피해야 할 주장", "권장 표현"],
          [
              ["거래내역으로 치매를 조기진단한다", "거래내역으로 금융안전상 확인이 필요한 변화를 탐지한다"],
              ["현재 제도는 전부 사후 대응이다", "예방수단도 있지만 스스로 위험을 인식하고 신청해야 하는 한계가 있다"],
              ["은행은 모든 거래를 행원이 조사한다", "FDS가 자동탐지하지만 경보 후 확인·상담·기록 부담이 남아 있다"],
              ["국내 최초·유일이다", "공개자료상 여러 기능의 통합 제공은 확인하기 어렵다고 표현한다"],
              ["AI가 행원 업무를 크게 줄인다", "MVP에서 검토·기록 시간 감소를 지표로 검증한다"],
              ["AI가 의심거래를 자동 차단한다", "금융회사 제휴와 법적 근거를 전제로 보류요청·공동확인을 지원한다"],
              ["보호자가 거래를 취소한다", "법적 대리권이 없으면 알림·제한열람·공동확인만 가능하다"],
              ["마이데이터로 모든 거래를 바로 가져온다", "고객동의 외에 사업자 자격·API 범위·금융회사 제휴가 필요하다"],
              ["공공·민간신탁이 전 계좌를 모니터링한다", "신탁에 맡긴 재산을 계약과 지급계획에 따라 관리한다"],
              ["이상징후가 없으면 안전하다", "미탐 가능성이 있으므로 AI 경보는 여러 보호수단 중 하나다"],
          ], [3800, 5560], compact=True)

    new_page(doc)
    title_block(doc, "Appendix B", "주요 참고자료", "2026년 8월 13일 기준 확인한 공식자료와 대표 연구")
    groups = [
        ("연구 근거", [
            ("Nicholas et al., JAMA Internal Medicine (2020)", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
            ("Triebel et al., Neurology (2009)", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
            ("UK bank data, JAMA Network Open (2025)", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
            ("National Institute on Aging: Managing Money Problems", "https://www.nia.nih.gov/health/legal-and-financial-planning/managing-money-problems-people-dementia"),
        ]),
        ("대한민국", [
            ("금융위원회 FDS·ASAP 운영성과", "https://www.fsc.go.kr/po010102/86997"),
            ("국민연금공단 치매안심재산관리 시범사업", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
            ("보건복지부 치매안심재산관리 안내", "https://www.mohw.go.kr/menu.es?mid=a10712010600"),
            ("대한민국 법원 성년후견제도", "https://www.scourt.go.kr/nm/min_3/min_3_12/index.html"),
            ("금융위원회 고령자 지정인 알림", "https://www.fsc.go.kr/no010101/73764"),
            ("금융위원회 금융거래 안심차단", "https://www.fsc.go.kr/no010101/85644"),
            ("금융결제원 어카운트인포", "https://www.payinfo.or.kr/main/main.do"),
            ("금융권 두낫콜", "https://www.donotcall.or.kr/"),
        ]),
        ("해외 제도·서비스", [
            ("FINRA Trusted Contact", "https://www.finra.org/rules-guidance/notices/22-31"),
            ("FINRA Rule 2165", "https://www.finra.org/rules-guidance/rulebooks/finra-rules/2165"),
            ("FCA Vulnerable Customers FG21/1", "https://www.fca.org.uk/publication/finalised-guidance/fg21-1.pdf"),
            ("일본 후생노동성 일상생활자립지원사업", "https://www.mhlw.go.jp/stf/seisakunitsuite/bunya/hukushi_kaigo/seikatsuhogo/chiiki-fukusi-yougo/index.html"),
            ("싱가포르 Money Lock", "https://www.mddi.gov.sg/newsroom/measures-to-protect-singaporeans-against-online-scams/"),
            ("EverSafe", "https://www.eversafe.com/for-families/"),
            ("Carefull", "https://getcarefull.com/"),
            ("Sibstar", "https://www.sibstar.co.uk/"),
        ]),
    ]
    for group, links in groups:
        doc.add_heading(group, level=2)
        for label, url in links:
            p = doc.add_paragraph(style="Small")
            p.paragraph_format.left_indent = Inches(0.16)
            r = p.add_run("• ")
            set_font(r, size=8.7, color=TEAL)
            hyperlink(p, label, url, size=8.7)
    p = doc.add_paragraph(style="Small")
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("주의: 본 보고서는 공모전 기획을 위한 조사·서비스 제안 문서이며, 의료·법률·투자 자문을 제공하지 않는다. 제도·법령·상품은 변경될 수 있으므로 실제 사업화 시 최신 원문과 전문가 검토가 필요하다.")
    set_font(r, size=8.7, italic=True, color=GRAY)


def audit(doc_path: Path):
    from zipfile import ZipFile
    import xml.etree.ElementTree as ET

    issues = []
    with ZipFile(doc_path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
        ns = {"w": W_NS}
        attr = lambda name: f"{{{W_NS}}}{name}"
        for idx, tbl in enumerate(root.findall(".//w:tbl", ns), 1):
            tbl_w = tbl.find("w:tblPr/w:tblW", ns)
            width = int(tbl_w.get(attr("w"), "0")) if tbl_w is not None else 0
            grid = [int(c.get(attr("w"), "0")) for c in tbl.findall("w:tblGrid/w:gridCol", ns)]
            if width <= 0 or sum(grid) != width:
                issues.append(f"table {idx} width/grid mismatch")
            for row in tbl.findall("w:tr", ns):
                cell_widths = []
                for tc in row.findall("w:tc", ns):
                    tcw = tc.find("w:tcPr/w:tcW", ns)
                    cell_widths.append(int(tcw.get(attr("w"), "0")) if tcw is not None else 0)
                if cell_widths != grid:
                    issues.append(f"table {idx} cell/grid mismatch")
                    break
    if issues:
        raise RuntimeError("; ".join(issues))


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    set_repeatable_header_footer(doc)
    doc.core_properties.title = "치매머니 AI 금융안전 플랫폼 조사·서비스 기획 보고서"
    doc.core_properties.subject = "2026 금융 AI Challenge 공유용 통합본"
    doc.core_properties.author = "Dementia Money Project Team"
    doc.core_properties.keywords = "치매머니, 금융 AI, 고령자 금융보호, FDS, 신뢰연락인, 치매안심재산관리"

    add_cover(doc)
    add_exec_summary(doc)
    add_problem_and_evidence(doc)
    add_korea_status(doc)
    add_overseas(doc)
    add_gap_and_definition(doc)
    add_customer_ai(doc)
    add_banker_ai(doc)
    add_flow_and_architecture(doc)
    add_mvp(doc)
    add_safety(doc)
    add_business_and_roadmap(doc)
    add_claims_and_sources(doc)

    # Explicit page fields should refresh when opened in Word.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUT)
    audit(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
