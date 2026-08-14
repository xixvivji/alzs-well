from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile
import sys
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
import build_dementia_money_report as base  # noqa: E402


OUT = ROOT / "output" / "docx" / "Ansim_Rhythm_2026_Financial_AI_Challenge_Master_Report.docx"

# standard_business_brief preset with a named Korean-font/brand-palette override.
# Named font override for Korean: the renderer is launched with the bundled
# LibreOffice fontconfig file so this system font is resolved and embedded.
FONT = "Apple SD Gothic Neo"
NAVY = "17324D"
BLUE = "2B6F9F"
TEAL = "2C8C82"
SKY = "EAF4F8"
MINT = "EAF6F3"
GOLD = "D69B2D"
PALE_GOLD = "FFF5DF"
RED = "A64242"
PALE_RED = "FAECEC"
INK = "202B35"
GRAY = "5F6B75"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D9E0E5"
WHITE = "FFFFFF"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def patch_base_tokens():
    for name, value in {
        "FONT": FONT,
        "NAVY": NAVY,
        "BLUE": BLUE,
        "TEAL": TEAL,
        "SKY": SKY,
        "MINT": MINT,
        "GOLD": GOLD,
        "PALE_GOLD": PALE_GOLD,
        "RED": RED,
        "PALE_RED": PALE_RED,
        "INK": INK,
        "GRAY": GRAY,
        "LIGHT_GRAY": LIGHT_GRAY,
        "MID_GRAY": MID_GRAY,
        "WHITE": WHITE,
    }.items():
        setattr(base, name, value)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + side))
        if node is None:
            node = OxmlElement("w:" + side)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height = None
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, cell in enumerate(row.cells):
            width = widths[i]
            cell.width = Twips(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("Source", "Lead", "Small", "Table Text", "Kicker"):
        if name not in styles:
            styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    source = styles["Source"]
    source.font.name = FONT
    source._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    source.font.size = Pt(8.1)
    source.font.color.rgb = RGBColor.from_string(GRAY)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)
    source.paragraph_format.line_spacing = 1.10

    lead = styles["Lead"]
    lead.font.name = FONT
    lead._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    lead.font.size = Pt(12)
    lead.font.bold = True
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.15

    small = styles["Small"]
    small.font.name = FONT
    small._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    small.font.size = Pt(8.6)
    small.font.color.rgb = RGBColor.from_string(GRAY)
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.10

    table_text = styles["Table Text"]
    table_text.font.name = FONT
    table_text._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    table_text.font.size = Pt(9)
    table_text.font.color.rgb = RGBColor.from_string(INK)
    table_text.paragraph_format.space_after = Pt(1)
    table_text.paragraph_format.line_spacing = 1.08

    kicker = styles["Kicker"]
    kicker.font.name = FONT
    kicker._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    kicker.font.size = Pt(9)
    kicker.font.bold = True
    kicker.font.color.rgb = RGBColor.from_string(TEAL)
    kicker.paragraph_format.space_after = Pt(3)

    bullet_id, number_id = add_numbering_definitions(doc)
    return bullet_id, number_id


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element

    def next_id(tag, attr):
        values = []
        for el in numbering.findall(qn(tag)):
            raw = el.get(qn(attr))
            if raw is not None and raw.isdigit():
                values.append(int(raw))
        return max(values, default=0) + 1

    def abstract(marker, fmt):
        abstract_id = next_id("w:abstractNum", "w:abstractNumId")
        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract_num.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        fonts.set(qn("w:eastAsia"), FONT)
        rpr.append(fonts)
        lvl.append(rpr)
        abstract_num.append(lvl)
        numbering.append(abstract_num)
        return abstract_id

    def num_for(abstract_id):
        num_id = next_id("w:num", "w:numId")
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abstract_id))
        num.append(aid)
        numbering.append(num)
        return num_id

    return num_for(abstract("•", "bullet")), num_for(abstract("%1.", "decimal"))


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(nid)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.14


def add_bullet(doc, bullet_id, text, *, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, bullet_id)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        base.set_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        base.set_font(r)
    else:
        base.set_font(p.add_run(text))
    return p


def add_number(doc, number_id, text):
    p = doc.add_paragraph()
    apply_num(p, number_id)
    base.set_font(p.add_run(text))
    return p


def set_header_footer(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header
        p = header.paragraphs[0]
        p.text = "안심리듬 | 금융생활 변화 조기알림 및 보호업무 코파일럿"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        base.set_font(p.runs[0], size=8, color=GRAY)
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run("2026.08.13  |  ")
        base.set_font(r, size=8.5, color=GRAY)
        base.add_page_field(p)


def paragraph(doc, text, *, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        base.set_font(r, bold=True, color=NAVY)
        base.set_font(p.add_run(text[len(bold_prefix):]))
    else:
        base.set_font(p.add_run(text))
    return p


def source(doc, items):
    return base.source_paragraph(doc, items)


def section_title(doc, kicker, title, subtitle):
    p = doc.add_paragraph(style="Kicker")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    base.set_font(p.add_run(kicker.upper()), size=9, bold=True, color=TEAL)

    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    base.set_font(p.add_run(title), size=24, bold=True, color=NAVY)

    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(11)
        base.set_font(p.add_run(subtitle), size=11.7, color=GRAY)


def new_page(doc):
    doc.add_page_break()


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("2026 금융 AI Challenge"), size=11, bold=True, color=TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("안심리듬"), size=40, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("ANSIM RHYTHM"), size=16, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    p.paragraph_format.line_spacing = 1.30
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(
        p.add_run("금융생활 변화 조기알림 및 보호업무 코파일럿\n출품 전략·기능명세·MVP 구현 통합본"),
        size=15,
        bold=True,
        color=INK,
    )

    base.callout(
        doc,
        "CORE DEMO",
        "같은 경보, 다른 맥락, 다른 다음 행동",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("질병을 예측하지 않고, 설명이 필요한 금융변화만 사람의 보호업무로 연결한다."), size=11, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(58)
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("기준일  2026. 8. 13."), size=10, color=GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_font(p.add_run("공유용 마스터 문서 | 사실·해석·팀 설계·내부 검증지표를 구분해 작성"), size=8.7, color=GRAY)


def add_exec_summary(doc, bullet_id):
    new_page(doc)
    section_title(doc, "Executive Summary", "안심리듬 한눈에 보기", "대회의 세 축을 하나의 사용자 여정으로 묶는 인간 중심 금융안전 코파일럿")
    base.callout(
        doc,
        "최종 정의",
        "안심리듬은 고객 동의로 평소 금융생활의 기준선을 학습하고, 장기간 누적되는 변화를 생활맥락으로 재확인하여 고객과 행원에게 설명 가능한 보호조치와 후속관리를 제안하는 금융안전 코파일럿이다.",
        fill=SKY,
        accent=BLUE,
    )
    doc.add_heading("대회 적합성", level=2)
    base.table(
        doc,
        ["대회가 요구하는 축", "안심리듬의 구현"],
        [
            ["고령층·취약소비자 포용금융", "낙인 없는 쉬운 설명, 본인 중심 동의, 생활맥락 확인"],
            ["이상금융거래 탐지", "개인별 기준선·MAD·추세·변화점·사유코드 기반 변화 탐지"],
            ["맞춤형 행동요령", "본인 확인→행원 검토→검증된 보호수단 안내의 단계적 흐름"],
            ["임직원 업무 효율화", "사건 묶음, 중립 질문, 근거, 상담기록·후속관리 초안 제공"],
        ],
        [3400, 5960],
        header_fill=NAVY,
    )
    doc.add_heading("수상 전략의 중심", level=2)
    add_bullet(doc, bullet_id, "치매·인지상태를 예측하지 않는다. 출력은 ‘평소와 다른 금융활동’과 확인 근거뿐이다.")
    add_bullet(doc, bullet_id, "같은 alertId·동일 거래·동일 사전판정을 유지하고, 생활맥락 응답만 바꾸어 A/B 다음 행동이 갈리는 장면을 시연한다.")
    add_bullet(doc, bullet_id, "AI는 점수·차단·상품가입을 결정하지 않는다. 탐지는 규칙·통계 엔진, 보호조치는 정책엔진, 최종결정은 고객과 행원이 담당한다.")
    add_bullet(doc, bullet_id, "로그인·API 키 없이 90초 안에 완주하고, LLM이 없어도 결정론적 템플릿으로 동일한 결과를 설명한다.")
    base.callout(doc, "세 가지 성공조건", "작동하는 웹서비스 > 설명 가능한 오탐 감소 > 인간의 안전한 최종판단", fill=MINT, accent=TEAL)


def add_competition(doc, bullet_id):
    new_page(doc)
    section_title(doc, "01 Competition", "대회 요구와 제출 운영계획", "공식 제출규칙과 팀 내부 검증기준을 섞지 않는다")
    base.table(
        doc,
        ["구분", "기한·조건", "안심리듬 대응"],
        [
            ["기획서 PDF", "2026. 9. 7.(월) 10:00", "최종 문제정의·서비스 흐름·근거·사업화 제시"],
            ["기능명세서 PDF", "2026. 9. 7.(월) 10:00", "3개 화면·상태기계·API·수용기준 중심"],
            ["실행 URL", "2026. 9. 7. 11:00~9. 11. 23:59 접근 가능", "무로그인·무키 기본경로, 상태점검과 장애 fallback"],
            ["발표심사", "상위 11팀 내외, 발표 15분+질의응답 5분", "90초 데모를 발표 전반의 증거로 사용"],
            ["최종 산출물", "2026. 10. 8.(목) 23:59, PDF+ZIP", "발표 진출 시 문서·코드·평가로그 고정"],
        ],
        [1900, 3180, 4280],
        compact=True,
    )
    paragraph(doc, "별도 제공 데이터가 없으므로 MVP는 완전 합성데이터로 구성한다. 공개된 세부 평가 배점은 확인되지 않았으므로 임의 점수표를 공식 기준처럼 제시하지 않는다.")
    doc.add_heading("팀 내부 합격선", level=2)
    add_bullet(doc, bullet_id, "심사자가 로그인과 설정 없이 첫 화면에서 문제와 경보를 이해한다.")
    add_bullet(doc, bullet_id, "A/B 분기에서 원시 거래와 사전판정이 동일하며 맥락만 달라진다.")
    add_bullet(doc, bullet_id, "신뢰연락인 미동의 상태에서는 연락 기능과 API가 모두 차단된다.")
    add_bullet(doc, bullet_id, "행원 승인 전 거래 차단·한도변경·신뢰연락인 통보가 발생하지 않는다.")
    add_bullet(doc, bullet_id, "합성데이터 성능은 코드 회귀·시뮬레이션 근거로만 표현하고 실제 금융권·임상 성능으로 확대하지 않는다.")
    source(doc, [("금융보안원 대회 공식 안내", "https://www.fsec.or.kr/bbs/detail?bbsNo=11997&menuNo=66")])


def add_problem(doc, bullet_id):
    new_page(doc)
    section_title(doc, "02 Problem", "출품용 문제정의", "보호수단은 존재하지만, 조기 발견에서 후속 보호업무까지의 연결이 분절돼 있다")
    paragraph(doc, "금융권에는 FDS·ASAP 같은 금융사기 탐지수단과 신탁·후견·안심차단 같은 보호수단이 존재한다. 그러나 고객 본인의 평소 금융생활을 기준으로 장기적인 변화를 발견하고, 생활맥락을 확인한 뒤, 적합한 보호수단과 행원의 후속업무까지 연결하는 과정은 여전히 분절돼 있다.", style="Lead")
    base.add_process(
        doc,
        [
            ("평소 생활", "작은 변화는 흩어져 보임"),
            ("변화 누적", "중복·미납·낯선 거래"),
            ("경보 발생", "정상 변화와 위험 혼재"),
            ("사람 확인", "전화·설명·기록 반복"),
            ("보호 연결", "제도별 신청·후속관리"),
        ],
        colors=[LIGHT_GRAY, PALE_GOLD, PALE_RED, SKY, MINT],
    )
    doc.add_heading("고객의 어려움", level=2)
    add_bullet(doc, bullet_id, "여러 은행·카드·자동이체에 흩어진 작은 변화를 스스로 종합하기 어렵다.")
    add_bullet(doc, bullet_id, "위험을 알아차려도 지연이체·안심차단·어카운트인포·두낫콜·신탁·후견 중 무엇을 언제 써야 하는지 판단하기 어렵다.")
    add_bullet(doc, bullet_id, "기존 제도는 사전 신청 또는 진단·후견 절차가 필요한 경우가 많아 대책을 세울 시점을 놓칠 수 있다.")
    doc.add_heading("행원의 어려움", level=2)
    add_bullet(doc, bullet_id, "FDS 경보 이후의 사실확인·중립 질문·제도 검색·상담기록·재연락은 사람의 반복업무로 남는다.")
    add_bullet(doc, bullet_id, "모든 고객의 장기 거래 변화를 전수조사할 수 없고, 데이터는 거래·서류·규정·법률에 파편화돼 있다.")
    base.callout(doc, "해결할 단위", "치매가 아니라 ‘설명이 필요한 금융생활 변화 사건’을 찾고, 사건별 다음 행동을 완결한다.", fill=PALE_GOLD, accent=GOLD)


def add_evidence(doc):
    new_page(doc)
    section_title(doc, "03 Evidence", "왜 금융생활 변화를 보나", "집단 수준의 연관 근거는 있지만 개인 진단으로 사용할 수는 없다")
    paragraph(doc, "연구에서는 인지기능 저하 전후에 연체, 계좌관리 실수, 지출통제 저하, 사기 취약성이 집단 수준에서 증가하는 경향이 관찰된다. 그러나 은퇴·입원·사별·이사·여행·소득변화도 같은 패턴을 만들 수 있다. 따라서 안심리듬은 질병 라벨을 만들지 않고 본인의 금융생활 기준선 대비 변화만 다룬다.")
    base.table(
        doc,
        ["금융행동", "근거 수준", "제품에서의 안전한 사용"],
        [
            ["청구서 누락·연체", "상대적으로 강함", "지속·반복과 데이터 누락 여부를 함께 확인"],
            ["계좌·명세서 관리 오류", "상대적으로 강함", "다른 신호군과 결합한 사건으로 구성"],
            ["과잉지출·현금흐름 악화", "부분 근거", "소득·필수지출·생활사건 대비 변화로 설명"],
            ["사기 취약성", "연관 근거 있음", "사기의심과 금융관리 변화는 별도 사유코드로 분리"],
            ["중복결제·잊은 구독", "합리적 가설", "진단 신호가 아니라 확인이 필요한 오류로 표시"],
            ["시간·지역·ATM 변화", "직접 근거 제한", "여행·입원·이사 맥락을 먼저 질문"],
            ["동일상품 반복가입", "직접 근거 제한", "상품추천 없이 소비자보호 검토 신호로만 사용"],
        ],
        [2700, 1800, 4860],
        compact=True,
    )
    base.callout(doc, "출력 문구", "‘질병 가능성’이 아니라 ‘최근 평소와 다른 금융활동이 확인되었습니다. 본인 확인이 필요합니다.’", fill=PALE_RED, accent=RED)
    source(doc, [
        ("JAMA Internal Medicine 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
        ("Neurology 2009", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
        ("JAMA Network Open 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
        ("미국 NIA", "https://www.nia.nih.gov/health/legal-and-financial-planning/managing-money-problems-people-dementia"),
    ])


def add_korea(doc, bullet_id):
    new_page(doc)
    section_title(doc, "04 Korea", "대한민국의 대응과 남은 공백", "사기 FDS·신탁·후견·사전신청형 보호수단은 작동하지만 목적과 시점이 다르다")
    base.table(
        doc,
        ["현재 수단", "작동 방식", "안심리듬과의 경계"],
        [
            ["은행 FDS·ASAP", "접속·거래·공유 사기정보로 의심거래를 탐지하고 확인·지연·지급정지", "보이스피싱 중심. 장기 개인 금융생활 변화 확인은 별도 층"],
            ["치매안심 재산관리서비스", "NPS가 신탁받은 현금성 자산을 계획에 따라 지급·관리", "2026년 시범사업. 전 금융거래 모니터링이 아니며 자격은 상담 확인"],
            ["공공후견·성년후견", "법원 결정 범위에서 재산·복지·의료·주거 업무 지원", "실시간 조기탐지가 아니라 신청·심판 기반 법적 지원"],
            ["민간 치매안심신탁", "사전 계약·지정대리인·발동조건에 따라 생활·의료비 지급", "맡긴 자산만 보호하고 가입능력·수수료·대리인 신뢰 이슈 존재"],
            ["지정인 알림·숙려", "일부 고위험 상품·카드대출 이용을 선택 지정인에게 통지", "옵트인·특정 상품·사후 알림 중심이며 일상거래 변화는 미포함"],
            ["안심차단·어카운트인포·두낫콜", "거래 사전차단, 자동이체 일괄조회·해지, 영업연락 차단", "변화 탐지 후 실행할 수 있는 국내 행동수단으로 연결"],
        ],
        [2250, 3650, 3460],
        compact=True,
    )
    doc.add_heading("사실을 정확히 말하는 방법", level=2)
    add_bullet(doc, bullet_id, "ASAP은 2025년 10월 출범 후 2026년 4월까지 31.7만 건의 정보공유를 통해 5,261건의 계좌 지급정지 조치와 474.6억 원의 피해예방 성과를 냈다. AI 단독 자동조치라고 표현하지 않는다.")
    add_bullet(doc, bullet_id, "은행 FDS도 고령자·신규 수취인·단시간 다회 이체 등을 일부 탐지한다. 따라서 ‘한국에는 이상거래 탐지가 없다’고 주장하지 않는다.")
    add_bullet(doc, bullet_id, "신한은행도 2026년 고령층 금융거래 이상징후 모니터링을 개발·고도화 중이다. ‘국내 최초·유일’ 대신 여러 금융사의 장기 변화와 탐지 이후 업무 연결을 차별점으로 둔다.")
    source(doc, [
        ("금융위 FDS·ASAP", "https://www.fsc.go.kr/po010102/86997"),
        ("NPS 치매안심 재산관리서비스", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
        ("보건복지부 안내", "https://www.mohw.go.kr/menu.es?mid=a10712010600"),
        ("금융보안원 FDS 가이드", "https://www.fsec.or.kr/bbs/detail?bbsNo=11355&menuNo=69"),
        ("신한그룹 공식자료", "https://www.shinhangroup.com/kr/archive/press/detail/805"),
    ])


def add_global(doc):
    new_page(doc)
    section_title(doc, "05 Global", "해외의 지원 방식", "진단이 아니라 모니터링·제한형 지갑·신뢰연락인·법적 대리·사람 개입을 조합한다")
    base.table(
        doc,
        ["유형", "대표 사례", "작동 방식", "한계·시사점"],
        [
            ["읽기전용 모니터링", "EverSafe·Carefull", "여러 금융계좌를 연결해 개인 패턴 대비 이상을 본인·선택 가족에게 알림", "자체 차단권한·독립 임상성능 공개 없음"],
            ["통제형 선불카드", "Sibstar·True Link", "한도·ATM·업종·채널을 미리 정하고 승인단계에서 허용·거절", "선불카드 밖 거래는 보호하지 못하고 과잉통제 위험"],
            ["은행 지원권한", "Lloyds·HSBC·Santander", "읽기전용 신뢰인, 별도 생활계좌, 돌봄카드, LPA 안내", "동의·법적 권한·자기결정권을 세분화"],
            ["직원 개입", "AARP BankSafe·UK Banking Protocol", "직원이 이상 인출을 질문하고 내부·경찰·복지기관으로 상향", "AI 이후의 확인·기록·신고 절차가 핵심"],
            ["제도적 보류", "미국·영국·캐나다·싱가포르", "착취·사기 의심 특정 거래를 제한 기간 보류하고 통지·재심", "전국·업권별 범위가 다르며 광범위 자동동결은 아님"],
        ],
        [1450, 2000, 3640, 2270],
        compact=True,
    )
    base.callout(doc, "해외 벤치마크", "탐지형과 통제형이 분리된 경우가 많다. 안심리듬은 개인 기준선 탐지→본인 재확인→동의된 공동확인→직원 판단→검증된 보호수단 안내를 하나의 사건 흐름으로 묶는다.", fill=MINT, accent=TEAL)
    source(doc, [
        ("EverSafe", "https://www.eversafe.com/for-families/"),
        ("Carefull", "https://getcarefull.com/"),
        ("Sibstar", "https://www.sibstar.co.uk/"),
        ("True Link", "https://www.truelinkfinancial.com/"),
        ("AARP BankSafe", "https://www.aarp.org/pri/initiatives/banksafe/"),
        ("FINRA Rule 2165", "https://www.finra.org/rules-guidance/notices/22-05"),
        ("FCA FG21/1", "https://www.fca.org.uk/publication/finalised-guidance/fg21-1.pdf"),
    ])


def add_positioning(doc, number_id):
    new_page(doc)
    section_title(doc, "06 Positioning", "안심리듬의 위치와 차별성", "기존 FDS를 대체하지 않고, 경보와 보호제도 사이의 ‘맥락 확인·후속업무’ 층을 만든다")
    base.table(
        doc,
        ["구분", "기존 FDS·ASAP", "안심리듬"],
        [
            ["목적", "금융사기·의심계좌 탐지와 차단", "개인의 금융생활 변화 확인과 보호계획 수립"],
            ["분석 관점", "거래·계좌·범죄 네트워크 이상징후", "개인별 장기 기준선·추세·생활맥락"],
            ["결과", "추가인증·지연이체·지급정지·사기 대응", "쉬운 설명·본인 확인·행원 검토·보호수단 안내"],
            ["행원 업무", "경보 이후 사실확인과 기록", "사건 묶음·질문·근거·기록·후속관리 지원"],
            ["결정권", "금융회사 정책·법적 절차", "AI가 아니라 고객·행원·권한기관"],
        ],
        [1600, 3830, 3930],
        compact=True,
    )
    doc.add_heading("서비스 설계 원칙", level=2)
    add_number(doc, number_id, "질병·인지상태를 진단·추론하지 않고 금융생활 변화만 탐지한다.")
    add_number(doc, number_id, "정상적인 생활변화는 맥락으로 확인해 불필요한 경보에서 제외한다.")
    add_number(doc, number_id, "설명이 필요한 사건만 행원의 보호업무 큐로 연결한다.")
    add_number(doc, number_id, "보호자·연락인·후견인의 권한을 구분하고 동의되지 않은 정보제공을 차단한다.")
    add_number(doc, number_id, "상품추천·가입·법위반 확정·거래실행은 AI 범위에서 제외한다.")
    base.callout(doc, "최종 포지셔닝", "고령자를 통제하는 AI가 아니라, 설명 가능한 개인 기준선 탐지와 사람의 확인을 연결해 자기결정권을 지키는 금융안전망", fill=SKY, accent=BLUE)


def add_field_view(doc, bullet_id):
    new_page(doc)
    section_title(doc, "07 Field View", "금융권 현업 관점을 반영한 업무 재설계", "AI 기술 소개보다 업무를 세부 태스크로 나누고 사람이 책임지는 전체 프로세스를 보여준다")
    base.table(
        doc,
        ["현업 관점", "안심리듬 반영"],
        [
            ["고객용 추천보다 임직원 효율화에 관심", "행원 검토·질문·설명·기록·후속관리 시간을 줄이는 코파일럿을 B2B 핵심가치로 둔다"],
            ["업무 태스크 분해가 먼저", "동의→데이터정비→탐지→근거봉인→맥락확인→정책판단→사람승인→감사를 별도 태스크로 분리"],
            ["데이터가 이미지·서류·규정·법률에 파편화", "거래 정규화 사전과 승인 문서 RAG를 분리하고 문서 버전·효력일·페이지를 저장"],
            ["망분리·심사·모델변경 부담", "공개·합성정보만 외부 LLM, 실제 거래·고객정보는 내부 또는 승인 전용환경, 모델 registry와 골든셋 검증"],
            ["챗봇에서 전사 Agent 플랫폼으로 이동", "전사 플랫폼을 새로 만들지 않고 플랫폼에 탑재 가능한 단일 소비자보호 업무 Agent로 설계"],
            ["최종판단은 사람", "AI는 근거검색·검수·반복작업 지원, 권리·재산에 영향을 주는 결정은 임직원 승인"],
        ],
        [3200, 6160],
        compact=True,
    )
    doc.add_heading("한 플랫폼, 두 권한영역", level=2)
    add_bullet(doc, bullet_id, "탐지영역: 규칙·통계·시계열 모델이 사유코드와 불변 evidence snapshot을 생성한다. LLM이 위험점수를 만들지 않는다.")
    add_bullet(doc, bullet_id, "안심설명 Copilot: 봉인된 사실과 승인 문서만 사용해 중립 질문·쉬운 설명·상담기록을 만든다. 직원이 수정·승인한다.")
    add_bullet(doc, bullet_id, "이체·동결·한도변경·연락인 변경 도구는 Agent 권한에서 제외한다.")
    source(doc, [("금융위 2026 금융분야 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142")])


def add_user_journey(doc, bullet_id):
    new_page(doc)
    section_title(doc, "08 User Journey", "3개 화면으로 완결되는 사용자 여정", "고객 화면이 본체이고, 행원 화면은 금융사 도입을 완성하는 보호업무 모듈이다")
    base.table(
        doc,
        ["화면", "고객·행원이 보는 것", "핵심 행동"],
        [
            ["1. 고객 금융생활", "12개월 변화 요약, 확인 필요 사건, 비교근거 3개", "경보 상세 열기"],
            ["2. 생활맥락 확인", "동일 alertId, 근거 거래, 중립 질문, 동의 상태", "본인 확인·기억 안 남·응답 보류"],
            ["3. 행원 코파일럿", "사건 묶음, evidence, 확인질문, 보호수단, 기록초안", "검토 승인·오탐 종결·재연락 예약"],
        ],
        [1900, 4450, 3010],
    )
    doc.add_heading("상태기계", level=2)
    base.add_process(
        doc,
        [
            ("OPEN", "개인 변화신호 사건 생성"),
            ("AWAITING_CONTEXT", "본인에게 생활맥락 질문"),
            ("NORMAL CLOSED", "증빙된 정상맥락으로 종결"),
            ("BANK REVIEW", "기억 안 남·무응답·hard signal"),
        ],
        colors=[PALE_GOLD, SKY, MINT, PALE_RED],
    )
    add_bullet(doc, bullet_id, "고객의 단순 ‘괜찮아요’만으로 고액 신규수취인·사기의심 hard signal을 자동 해제하지 않는다.")
    add_bullet(doc, bullet_id, "정상으로 확정된 생활맥락만 제한적으로 다음 기준선 갱신에 반영한다.")
    add_bullet(doc, bullet_id, "최초 판정을 덮어쓰지 않고 preDecision·postDecision·맥락 출처·유효기간·알고리즘 버전을 함께 남긴다.")


def add_demo(doc):
    new_page(doc)
    section_title(doc, "09 Demo", "90초 핵심 데모", "같은 경보, 다른 맥락, 다른 다음 행동")
    base.callout(doc, "비교 원칙", "A와 B는 동일 alertId·동일 거래·동일 evidence·동일 preDecision을 사용하고 생활맥락 응답만 바꾼다.", fill=PALE_GOLD, accent=GOLD)
    base.table(
        doc,
        ["시간", "시연", "화면에 남는 증거"],
        [
            ["0~15초", "무로그인 시작→고객 홈 경보 선택", "신규 부동산 수취인, 반복송금, 공과금 누락을 묶은 동일 alertId"],
            ["15~40초", "A: 이사 계약금·잔금과 주소이전 내역 확인", "세 신호를 설명하는 구조적 근거 일치, CLOSED_NORMAL, 추가 조치 없음"],
            ["40~50초", "같은 경보를 Reset하여 B로 재실행", "원시 거래·preDecision·알고리즘 버전 동일"],
            ["50~65초", "B: 같은 거래에 ‘본인 거래인지 확인할 수 없음’ 선택", "신규수취인·반복송금·공과금 누락 유지, BANK_REVIEW, 신뢰연락인 통보 차단"],
            ["65~85초", "행원 화면에서 근거·질문·조치 초안 검토", "직원 승인 전 차단·한도변경 없음, 상담기록 초안"],
            ["85~90초", "안전원칙 요약", "진단 없음·자동조치 없음·미동의 제3자 제공 없음"],
        ],
        [1300, 3840, 4220],
        compact=True,
    )
    doc.add_heading("데모 수용기준", level=2)
    base.table(
        doc,
        ["항목", "합격 기준"],
        [
            ["접근", "로그인 0회, API 키 0개, 첫 상호작용 5초 이내"],
            ["재현성", "동일 입력·동일 버전의 결정 재현율 100%"],
            ["안전", "미동의 연락 호출 0건, 승인 전 실제 조치 0건"],
            ["장애대응", "키 없음·timeout·429·5xx·스키마 오류에서 템플릿 fallback 100%"],
            ["격리", "익명 세션 간 상태 누출 0건, Reset 멱등성 보장"],
        ],
        [2100, 7260],
        compact=True,
        header_fill=TEAL,
    )


def add_detection(doc, bullet_id):
    new_page(doc)
    section_title(doc, "10 Detection", "개인 기준선 탐지엔진", "AI가 문장을 만드는 장면이 아니라 오탐을 줄이는 계산과 근거를 보여준다")
    base.table(
        doc,
        ["신호", "관측단위·방법", "예외·fallback"],
        [
            ["금액·횟수 급증", "일·주 집계 중앙값과 MAD, 단측 modified z", "MAD=0·희소값이면 IQR·경험분위·절대규칙 또는 LOW_CONFIDENCE"],
            ["점진 추세", "6~8주 주별 집계의 Theil-Sen 기울기 또는 3개 연속창 증가", "계절성·명절·급여일 변화는 hard negative로 검증"],
            ["수준 전환", "미래를 보지 않는 online Page-Hinkley", "최소구간·reset·false alarm·허용 탐지창 고정"],
            ["신규 수취인", "과거 수취인 사전과 최초시점 규칙", "표기변형·계열사·본인계좌 정규화"],
            ["반복송금·중복결제", "동일/유사 금액·수취인·시간창 규칙", "취소·환불·pending 상태 분리"],
            ["정기납부 누락", "주기 추정+예상일 grace period", "데이터 미수집·연결장애를 미납으로 보지 않음"],
        ],
        [2200, 3970, 3190],
        compact=True,
    )
    paragraph(doc, "Modified z-score 예시: 0.6745 × (현재값 - 개인 중앙값) ÷ MAD. 다만 이 수치를 질병 확률이나 단일 위험확률처럼 표시하지 않고, 검증된 사유코드와 정책표로 사건단계를 결정한다.", style="Small")
    doc.add_heading("사건 융합과 기준선 보호", level=2)
    add_bullet(doc, bullet_id, "독립 신호 1개는 L1, 지속 또는 서로 다른 신호군 2개는 L2, 사기 hard rule은 L3처럼 정책표로 결합한다.")
    add_bullet(doc, bullet_id, "같은 고객·기간의 신호를 하나의 incident로 묶고 duplicate suppression·cooldown·월 경보예산을 적용한다.")
    add_bullet(doc, bullet_id, "개인화 준비상태를 READY·LOW_CONFIDENCE·COLD_START로 구분하고 특징별 최소 90일·거래수·납부주기를 둔다.")
    add_bullet(doc, bullet_id, "이상기간을 즉시 정상으로 학습하지 않는다. 고객·행원이 정상으로 확정한 관측만 기준선 갱신 후보로 사용한다.")


def add_comparison(doc):
    new_page(doc)
    section_title(doc, "11 Evaluation Design", "개인 기준선이 오탐을 줄이는지 검증", "전역 임계값과 같은 holdout·같은 경보예산에서 비교한다")
    base.table(
        doc,
        ["비교군", "정의", "확인하려는 질문"],
        [
            ["규칙 only", "신규 수취인·중복·미납 등 고정 규칙", "기본 안전신호만으로 어느 정도 잡히는가"],
            ["전역 임계값", "전체 합성고객 분포에서 정한 단일 기준", "개인차 때문에 정상 생활이 얼마나 과다 경보되는가"],
            ["개인 MAD", "고객별 중앙값·MAD·준비상태", "금액·횟수 개인화가 오탐을 줄이는가"],
            ["개인 ensemble", "규칙+MAD+추세+변화점+사건융합", "다중신호가 recall과 운영 경보량을 함께 개선하는가"],
            ["맥락 재평가", "동일 경보에서 확인된 생활사건만 추가", "안전한 정상종결로 불필요 경보가 얼마나 감소하는가"],
        ],
        [1900, 3600, 3860],
        compact=True,
    )
    doc.add_heading("표시할 내부 검증지표", level=2)
    base.table(
        doc,
        ["영역", "지표", "표현 원칙"],
        [
            ["탐지", "episode precision·recall·F1, PR-AUC, 탐지지연", "거래 단위가 아니라 사건 단위로 계산"],
            ["운영", "사용자-월당 오탐, 중복경보, 검토시간", "정확도 하나보다 행원 부담을 함께 표시"],
            ["맥락", "정상종결률, unsafe downgrade rate, 재경보율", "hard signal의 위험한 하향은 0 목표"],
            ["안전", "미동의 연락 0, 승인 전 조치 0, 허구근거 0", "절대 건수로 공개"],
            ["데모", "90초 완주율, P50/P95 지연, no-key 성공률", "공개 URL에서 반복 smoke test"],
        ],
        [1500, 4000, 3860],
        compact=True,
        header_fill=TEAL,
    )
    base.callout(doc, "주의", "합성데이터 결과는 코드 회귀와 설계 비교를 위한 시뮬레이션이다. 실제 금융사 성능·고령층 효과·임상 예측력으로 주장하지 않는다.", fill=PALE_RED, accent=RED)


def add_policy_and_consent(doc, bullet_id):
    new_page(doc)
    section_title(doc, "12 Policy", "단계적 개입과 신뢰연락인 동의", "탐지보다 더 중요한 것은 누가 어떤 정보를 보고 어떤 행동을 할 수 있는지다")
    base.table(
        doc,
        ["단계", "조건", "안심리듬의 출력", "실행 주체"],
        [
            ["L0 기록", "낮은 단일 변화", "고객 대시보드 기록", "고객"],
            ["L1 확인", "설명이 필요한 단일 사건", "근거 3개·중립 질문", "고객"],
            ["L2 계획", "반복·지속 또는 다중 신호", "보호수단 체크리스트·상담 권유", "고객"],
            ["L3 검토", "기억 안 남·무응답·hard signal", "행원 사건 큐·질문·기록 초안", "행원"],
            ["L4 외부연계", "사기 임박·지속적 관리지원 필요", "112·1332·은행·공공제도 안내", "금융회사·권한기관"],
        ],
        [1200, 2750, 3510, 1900],
        compact=True,
    )
    doc.add_heading("신뢰연락인 하드 규칙", level=2)
    add_bullet(doc, bullet_id, "고객이 제공받는 자·목적·항목·기간·철회방법을 선택해 별도 동의하고, 연락인도 초대를 수락·본인확인한다.")
    add_bullet(doc, bullet_id, "미동의 상태에서는 연락 버튼을 비활성화하고 서버 API도 CONSENT_REQUIRED로 거절한다.")
    add_bullet(doc, bullet_id, "연락인에게는 ‘확인이 필요한 금융활동 발생’ 정도의 최소정보만 전달하며 전체 거래내역·수취인·질병 추정정보를 보내지 않는다.")
    add_bullet(doc, bullet_id, "연락인은 법정대리인이 아니다. 알림·제한열람·공동확인과 송금·해지·동결 권한을 분리한다.")
    add_bullet(doc, bullet_id, "연락인이 가해자일 가능성에 대비해 행원 전용검토·복수 연락인·즉시 철회·열람로그를 제공한다.")
    source(doc, [
        ("개인정보보호법", "https://law.go.kr/LSW/lsInfoP.do?lsiSeq=270351"),
        ("신용정보법 제32조", "https://www.law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1025128253"),
    ])


def add_functional_spec(doc):
    new_page(doc)
    section_title(doc, "13 Functional Spec", "세 화면 기능명세와 수용기준", "심사자가 보는 UI와 서버가 지켜야 할 규칙을 함께 정의한다")
    base.table(
        doc,
        ["ID", "기능", "입력", "출력·수용기준"],
        [
            ["C-01", "고객 홈", "익명 demoSession", "최근 변화사건·근거 3개·기준선 준비상태, 질병 라벨 0개"],
            ["C-02", "경보 상세", "alertId", "evidence transaction ID·비교기간·reasonCodes·preDecision"],
            ["C-03", "맥락 확인", "선택 응답·선택 증빙기간", "postDecision과 설명, 원판정은 보존"],
            ["C-04", "연락인 확인", "consent snapshot", "미동의면 UI/API 차단, 동의면 최소정보 미리보기"],
            ["B-01", "행원 사건 큐", "BANK_REVIEW 사건", "우선순위·사유·미확인 상태·재연락기한"],
            ["B-02", "안심설명", "불변 evidence snapshot", "중립 질문·쉬운 설명·공식근거, 문서 밖 내용은 확인 불가"],
            ["B-03", "직원 결정", "승인·수정·오탐·재연락", "직원 ID·이유·시각·정책버전 감사로그"],
            ["D-01", "Reset", "demoSession·scenario", "동일 alertId 초기상태 복원, 여러 번 실행해도 같은 결과"],
        ],
        [900, 1800, 2340, 4320],
        compact=True,
    )
    paragraph(doc, "고객 홈·맥락 확인·행원 코파일럿의 세 화면만 필수로 둔다. 보호수단 목록·연락인 설정·감사타임라인은 이 세 화면의 패널 또는 모달로 구현해 MVP 범위를 확장하지 않는다.", style="Small")


def add_architecture(doc, bullet_id):
    new_page(doc)
    section_title(doc, "14 Architecture", "Spring Boot 모듈형 모놀리스", "탐지·정책결정과 생성형 설명을 분리하고, 키가 없어도 전체 서비스가 작동한다")
    base.add_process(
        doc,
        [
            ("Browser", "Thymeleaf+HTMX 3화면"),
            ("demo", "익명세션·reset·API"),
            ("case", "상태기계·동의·정책"),
            ("detection", "규칙·MAD·추세·변화점"),
            ("ledger", "합성거래·정규화·품질"),
        ],
        colors=[LIGHT_GRAY, SKY, PALE_GOLD, MINT, LIGHT_GRAY],
    )
    base.table(
        doc,
        ["모듈", "책임", "금지 의존"],
        [
            ["ledger", "거래·수취인·자동이체 정규화, occurredAt/postedAt, 취소·환불·누락 품질", "LLM·웹 DTO"],
            ["detection", "feature·baseline·MAD·Theil-Sen·Page-Hinkley·reasonCodes", "LLM·행원 화면"],
            ["case", "incident·context·상태기계·consent·policy decision", "외부 모델 직접호출"],
            ["explanation", "템플릿 기본, 선택형 LLM, 승인 문서 RAG, 출력검증", "위험등급·행동코드 변경"],
            ["demo", "3화면 controller, signed anonymous session, reset", "업무규칙 직접 구현"],
            ["audit", "결정 이벤트·모델·문서·사람 override append-only 기록", "원문 프롬프트 일반로그"],
        ],
        [1550, 4850, 2960],
        compact=True,
    )
    doc.add_heading("권장 고정 기술구성", level=2)
    add_bullet(doc, bullet_id, "Java 21 · Spring Boot 4.1.0 · Spring AI 2.0.0 · Spring Modulith 2.1.0")
    add_bullet(doc, bullet_id, "Spring MVC · Spring Security · Thymeleaf/HTMX · Spring JDBC 또는 jOOQ")
    add_bullet(doc, bullet_id, "PostgreSQL+pgvector · Flyway · PDF reader/PDFBox · Testcontainers 2.0.5 · Actuator/Micrometer")
    add_bullet(doc, bullet_id, "ApplicationModules.verify()로 모듈 cycle·내부패키지 접근을 CI에서 검증한다.")
    source(doc, [
        ("Spring Boot 요구사항", "https://docs.spring.io/spring-boot/system-requirements.html"),
        ("Spring AI 시작문서", "https://docs.spring.io/spring-ai/reference/getting-started.html"),
        ("Spring Modulith 검증", "https://docs.spring.io/spring-modulith/reference/verification.html"),
    ])


def add_llm_and_rag(doc, bullet_id):
    new_page(doc)
    section_title(doc, "15 AI Routing", "혼합 LLM·RAG·Fallback", "실제 금융정보를 외부 LLM에 보내지 않고, 설명 생성이 실패해도 탐지·정책결정은 유지한다")
    base.table(
        doc,
        ["데이터 등급", "MVP 라우팅", "금융회사 확장"],
        [
            ["PUBLIC", "외부 LLM 또는 템플릿", "승인된 외부·전용 LLM"],
            ["SYNTHETIC", "외부 LLM 또는 템플릿", "승인된 외부·내부 LLM"],
            ["INTERNAL·PSEUDONYMIZED", "차단", "내부 LLM 또는 승인 private endpoint"],
            ["PERSONAL_CREDIT·AUTHENTICATION", "차단", "기본 차단, 별도 법적·보안근거와 최소처리"],
        ],
        [2550, 2750, 4060],
        compact=True,
    )
    base.callout(doc, "유효 등급", "effectiveDataClass = max(입력 데이터 등급, 검색된 문서 등급). 검색결과가 더 민감하면 더 높은 등급으로 라우팅한다.", fill=PALE_GOLD, accent=GOLD)
    doc.add_heading("결정론적 기본경로", level=2)
    add_bullet(doc, bullet_id, "detector→policy decision→structured ExplanationFacts→TemplateExplanationGenerator가 기본이다.")
    add_bullet(doc, bullet_id, "선택형 LLM은 ExplanationFacts의 말투만 쉽게 바꾸며 riskBand·reasonCodes·actionCode·연락권한을 변경할 수 없다.")
    add_bullet(doc, bullet_id, "키 없음·timeout·429·5xx·malformed JSON·schema failure에서 즉시 템플릿으로 fallback한다.")
    doc.add_heading("RAG 근거 통제", level=2)
    add_bullet(doc, bullet_id, "승인된 공식 문서만 수집하고 문서 ID·기관·페이지·조항·효력일·버전·checksum·접근등급을 저장한다.")
    add_bullet(doc, bullet_id, "탐지 시점의 evidence snapshot을 먼저 봉인한 뒤 Copilot은 그 사실과 승인 문서만 문장화한다.")
    add_bullet(doc, bullet_id, "출처 URL은 LLM이 만들지 않고 서버가 검색 메타데이터에서 조립한다. 근거가 없으면 ‘확인 불가’로 남긴다.")


def add_data_and_audit(doc):
    new_page(doc)
    section_title(doc, "16 Data & Audit", "데이터 구조와 감사 가능성", "데이터 정비·효력일·모델버전·사람의 판단을 한 사건 단위로 재현한다")
    base.table(
        doc,
        ["테이블", "핵심 필드"],
        [
            ["demo_session", "session_id, scenario_seed, expires_at, reset_version"],
            ["transaction", "occurred_at, posted_at, amount, merchant/payee_key, channel, status, synthetic_label"],
            ["baseline_snapshot", "feature_code, window, median, MAD, readiness, algorithm_version, data_cutoff"],
            ["alert_incident", "alert_id, pre_decision, post_decision, reason_codes, evidence_ids, state"],
            ["context_event", "context_code, source, evidence_period, confidence, expiry, user_response"],
            ["consent_snapshot", "purpose, recipient, fields, duration, status, granted_at, revoked_at"],
            ["knowledge_chunk", "source_version, page, section, effective_from/to, access_class, sha256, embedding_version"],
            ["decision_audit", "trace_id, feature_hash, policy/model/prompt/schema version, citations, staff_override, timestamps"],
        ],
        [2500, 6860],
        compact=True,
    )
    paragraph(doc, "데이터 필드와 화면 어디에도 질병 위험점수나 인지상태 라벨을 만들지 않는다. 모델이 산출하는 것은 금융생활 변화신호·기준선 준비상태·사유코드·사람 검토 필요 여부뿐이다.", style="Lead")
    base.callout(doc, "로그 원칙", "원문 거래·상담내용·프롬프트는 일반 로그에 저장하지 않는다. 운영 로그는 가명 trace·버전·상태·소요시간을, 업무 감사로그는 사건 근거와 사람 결정을 최소필요 범위로 분리 보관한다.", fill=SKY, accent=BLUE)


def add_synthetic_data(doc, bullet_id):
    new_page(doc)
    section_title(doc, "17 Synthetic Data", "합성 골든셋과 평가설계", "화면용 두 고객과 성능평가용 독립 데이터셋을 분리한다")
    base.table(
        doc,
        ["구성", "내부 목표", "주의"],
        [
            ["화면용 persona", "A/B 시연용 동일 alertId 1쌍", "평가지표 산출에 재사용하지 않음"],
            ["평가 persona", "독립 30명 이상", "고객 단위 train/tune/test 분리"],
            ["시나리오", "유형별 episode 20개 이상", "거래가 아니라 사건 단위 라벨"],
            ["반복실험", "5개 이상 random seed", "generator와 detector의 seed·threshold 분리"],
            ["신뢰구간", "persona bootstrap 95% CI", "작은 표본의 단일 숫자 과신 방지"],
        ],
        [2100, 3000, 4260],
        compact=True,
    )
    doc.add_heading("반드시 포함할 hard negative", level=2)
    add_bullet(doc, bullet_id, "입원·병원비, 여행·해외사용, 이사·보증금, 명절·가족지원, 큰 합법구매, 급여일 변화")
    add_bullet(doc, bullet_id, "계좌 신규연결, 데이터 단절·지연수집, pending→취소·환불, 가맹점명 표기변형")
    add_bullet(doc, bullet_id, "거래이력 부족·값 대부분 0·MAD=0·정기납부 cycle 부족의 cold-start 사례")
    doc.add_heading("검증 분리 원칙", level=2)
    add_bullet(doc, bullet_id, "baseline 기간과 event 주입기간을 시간으로 분리하고 미래정보를 쓰지 않는 causal replay를 수행한다.")
    add_bullet(doc, bullet_id, "합성 generator가 detector의 임계값·사유코드 판정을 읽지 못하게 별도 패키지·설정으로 분리한다.")
    add_bullet(doc, bullet_id, "합성 연령·성별 결과는 실제 공정성 입증이 아니라 시뮬레이션 내 강건성 점검으로만 표시한다.")


def add_security(doc, bullet_id):
    new_page(doc)
    section_title(doc, "18 Trust & Safety", "법적·보안·소비자보호 안전장치", "보호를 이유로 자기결정권·민감정보·금융접근권을 침해하지 않는다")
    base.table(
        doc,
        ["위험", "MVP 안전장치", "상용화 추가요건"],
        [
            ["질병 추론", "질병·인지상태 라벨과 학습목표 자체를 만들지 않음", "건강정보 연계 시 별도 민감정보 동의·접근권한"],
            ["완전자동 결정", "알림·설명·행원 보조만, 실제 조치 없음", "설명·검토·거부요구와 사람의 실질 판단"],
            ["제3자 제공", "연락인 별도동의·최소정보·철회·API 차단", "신용정보법 동의·감사·권한검증"],
            ["외부 LLM 유출", "합성·공개정보만 허용, 실제 거래 입력 차단", "내부/전용모델·DLP·private endpoint·무보존 계약"],
            ["생성형 환각", "템플릿 기본, 승인문서 RAG, 구조화 출력·서버 인용", "골든셋·공격시험·모델변경 승인·kill switch"],
            ["Agent 오권한", "송금·동결·한도·연락인 변경 tool 없음", "최소권한 토큰·step-up 인증·이중승인"],
            ["세션 누출", "signed anonymous session·TTL·rate limit·reset", "MFA·RBAC/ABAC·RLS·KMS/HSM"],
        ],
        [1900, 3650, 3810],
        compact=True,
    )
    doc.add_heading("MVP 보안 체크", level=2)
    add_bullet(doc, bullet_id, "실제 고객 금융정보·진단서·주민번호·계좌번호 입력을 차단하고 합성데이터임을 화면에 표시한다.")
    add_bullet(doc, bullet_id, "PDF·문서 입력이 있다면 파일 형식·크기·페이지·처리시간을 제한하고 외부 문서를 비신뢰 데이터로 격리한다.")
    add_bullet(doc, bullet_id, "prompt/completion 원문 로깅을 끄고 모델·지연·토큰 관측과 업무결정 감사로그를 분리한다.")
    source(doc, [
        ("개인정보보호법 제23조", "https://www.law.go.kr/LSW/lsLinkCommonInfo.do?chrClsCd=010202&lsJoLnkSeq=1027416043"),
        ("자동화된 결정 안내", "https://m.pipc.go.kr/np/cop/bbs/selectBoardArticle.do?bbsId=BS074&mCode=C020010000&nttId=10611"),
        ("금융위 망분리 로드맵", "https://fsc.go.kr/po010102/82885"),
        ("금융보안원 AI 보안 안내", "https://www.fsec.or.kr/bbs/detail?bbsNo=11977&menuNo=222"),
    ])


def add_business(doc, bullet_id):
    new_page(doc)
    section_title(doc, "19 Adoption", "금융회사 도입모델과 사업성", "전사 AI 플랫폼을 새로 만드는 것이 아니라 그 안에 바로 탑재할 수 있는 단일 소비자보호 Agent")
    base.table(
        doc,
        ["대상", "도입가치", "형태"],
        [
            ["은행·카드", "장기 변화사건 선별, 확인·기록시간 감소, 소비자보호 강화", "내부 Agent 플랫폼 모듈·API·화이트라벨"],
            ["증권·보험", "반복가입·불완전판매 가능성의 사람 검토, 쉬운 설명", "직원용 Copilot·사후품질점검"],
            ["마이데이터 사업자", "다기관 거래를 개인 기준선으로 통합해 금융안전 부가가치 제공", "허가사업자 제휴형 B2B2C"],
            ["공공기관", "적합한 상담·후견·재산관리 제도로 조기연결", "실증·공공조달·지역연계"],
        ],
        [1900, 4800, 2660],
        compact=True,
    )
    doc.add_heading("상용화 전제", level=2)
    add_bullet(doc, bullet_id, "마이데이터 2.0은 상용제도지만 스타트업이 전 금융거래를 바로 가져오는 권한은 아니다. 금융회사 또는 허가받은 본인신용정보관리업자와 제휴한다.")
    add_bullet(doc, bullet_id, "D-테스트베드는 원격 분석 실증사업이고, 마이데이터 테스트베드는 API 기능·연동·보안 점검 환경이다. MVP 실데이터 조달수단으로 표현하지 않는다.")
    add_bullet(doc, bullet_id, "거래정보·고객상담·내부규정은 내부 또는 승인 private 환경에서 처리하고 공개자료 업무만 외부 LLM으로 라우팅한다.")
    source(doc, [
        ("마이데이터 2.0", "https://fsc.go.kr/po010101/84780"),
        ("D-테스트베드", "https://www.fsc.go.kr/po010106/86744"),
        ("마이데이터 테스트베드", "https://fsc.go.kr/no010101/75679"),
    ])


def add_roadmap(doc):
    new_page(doc)
    section_title(doc, "20 Roadmap", "제출까지의 실행 우선순위", "문서보다 작동하는 세 화면과 재현 가능한 A/B 증거를 먼저 완성한다")
    base.table(
        doc,
        ["기간", "우선순위", "완료조건"],
        [
            ["8/13~8/15", "1. 출품용 문제정의 확정", "한 문장·금지주장·데모 alertId·정책표 고정"],
            ["8/16~8/19", "2. 세 화면 기능명세", "화면 상태·API·동의 차단·감사 이벤트 확정"],
            ["8/20~8/25", "3. 합성데이터와 탐지엔진", "MAD·추세·변화점·dedup·cold-start 단위테스트"],
            ["8/26~8/30", "4. A/B 데모 완성", "무로그인·Reset·행원승인·템플릿 fallback"],
            ["8/31~9/3", "5. 골든셋·성능비교", "전역/개인 ablation·운영지표·안전건수 검증"],
            ["9/4~9/6", "6. 기획서·배포 안정화", "PDF·기능명세·URL smoke test·복구절차·최종 QA"],
            ["9/7", "7. 제출·URL 유지", "10:00 제출, 11:00부터 외부 접근 확인"],
        ],
        [1800, 2800, 4760],
        compact=True,
    )
    doc.add_heading("출시 전 10개 게이트", level=2)
    checks = [
        "미래정보 없는 causal replay",
        "MAD=0·cold-start·누락/지연 처리",
        "incident fusion·dedup·cooldown",
        "동일 alertId A/B 입력 고정",
        "hard signal의 자기신고 자동해제 금지",
        "no-key·timeout fallback 100%",
        "익명 세션 격리·Reset·rate limit",
        "generator·detector 분리와 persona holdout",
        "알고리즘·맥락·동의·사람결정 감사로그",
        "모듈 검증·Testcontainers 통합테스트·공개 URL smoke monitor",
    ]
    for item in checks:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3.5)
        base.set_font(p.add_run("□ "), size=10, bold=True, color=TEAL)
        base.set_font(p.add_run(item), size=10)


def add_claims(doc):
    new_page(doc)
    section_title(doc, "Appendix A", "표현·기능 금지선", "과장과 낙인을 피하고 심사 질문에 일관되게 답한다")
    base.table(
        doc,
        ["피해야 할 표현·기능", "제출본의 표현·설계"],
        [
            ["치매 위험도·인지저하 가능성 산출", "개인 기준선 대비 금융생활 변화와 확인 필요 사유만 표시"],
            ["기존 국내 탐지체계는 없다", "사기형 FDS와 장기 생활변화·후속업무 사이의 공백"],
            ["국내 최초·유일", "개인 기준선·맥락 재평가·행원 workflow 결합의 차별성"],
            ["AI가 거래를 자동 차단", "실제 조치는 본인 재확인·행원 검토·법적 권한 이후"],
            ["가족에게 자동 통보", "별도 동의·최소정보·철회, 미동의면 UI/API 차단"],
            ["AI가 상품을 추천·가입", "공식 보호수단 정보를 조건과 기준일과 함께 안내, 적용 여부는 직원 확인"],
            ["마이데이터로 즉시 전 금융사 연결", "MVP는 합성데이터, 상용화는 금융회사·허가사업자 제휴"],
            ["합성셋 정확도=실제 성능", "코드 회귀·설계 비교를 위한 시뮬레이션 결과"],
            ["LLM이 위험판정", "Java 규칙·통계엔진이 판정하고 LLM은 봉인된 사실을 쉬운 말로 변환"],
            ["민원예방 Agent", "고객가치를 앞세운 취약소비자 안심설명·이해확인 Copilot"],
        ],
        [3800, 5560],
        compact=True,
    )
    base.callout(doc, "최종 한 줄", "정상적인 생활변화는 오탐에서 제외하고, 설명이 필요한 금융변화만 골라 고객과 행원의 안전한 다음 행동으로 연결한다.", fill=MINT, accent=TEAL)


def add_sources(doc, bullet_id):
    new_page(doc)
    section_title(doc, "Appendix B", "주요 참고자료", "2026년 8월 13일 기준 공식자료·대표 연구·공식 기술문서")
    groups = [
        ("대회·금융정책", [
            ("금융보안원 2026 금융 AI Challenge", "https://www.fsec.or.kr/bbs/detail?bbsNo=11997&menuNo=66"),
            ("금융위원회 FDS·ASAP 운영성과", "https://www.fsc.go.kr/po010102/86997"),
            ("금융위원회 금융분야 AI 가이드라인", "https://www.fsc.go.kr/po010101/87142"),
            ("금융위원회 망분리 개선 로드맵", "https://fsc.go.kr/po010102/82885"),
            ("금융보안원 AI 보안 안내서", "https://www.fsec.or.kr/bbs/detail?bbsNo=11977&menuNo=222"),
        ]),
        ("국내 보호제도", [
            ("국민연금공단 치매안심 재산관리서비스", "https://www.nps.or.kr/pnsgdnc/nscvrgdata/getOHAE0002M1.do?menuId=MN24000898&pstId=ZZ202600000000000453"),
            ("보건복지부 치매안심 재산관리 안내", "https://www.mohw.go.kr/menu.es?mid=a10712010600"),
            ("금융결제원 어카운트인포", "https://www.payinfo.or.kr/main/main.do"),
            ("금융권 두낫콜", "https://www.donotcall.or.kr/"),
            ("금융위원회 마이데이터 2.0", "https://fsc.go.kr/po010101/84780"),
        ]),
        ("대표 연구·해외", [
            ("JAMA Internal Medicine 2020", "https://jamanetwork.com/journals/jamainternalmedicine/fullarticle/2773241"),
            ("Neurology 2009", "https://www.neurology.org/doi/10.1212/WNL.0b013e3181b87971"),
            ("JAMA Network Open 2025", "https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2835294"),
            ("FINRA Trusted Contact", "https://www.finra.org/rules-guidance/notices/22-31"),
            ("FCA Vulnerable Customers FG21/1", "https://www.fca.org.uk/publication/finalised-guidance/fg21-1.pdf"),
            ("싱가포르 CPF Trusted Contact", "https://www.cpf.gov.sg/member/infohub/news/news-releases/cpfb-launches-tc-notification-service-and-cpf-safety-switch1"),
        ]),
        ("기술문서", [
            ("Spring Boot System Requirements", "https://docs.spring.io/spring-boot/system-requirements.html"),
            ("Spring AI Getting Started", "https://docs.spring.io/spring-ai/reference/getting-started.html"),
            ("Spring AI Structured Output", "https://docs.spring.io/spring-ai/reference/api/structured-output-converter.html"),
            ("Spring AI PGvector", "https://docs.spring.io/spring-ai/reference/api/vectordbs/pgvector.html"),
            ("Spring Modulith Verification", "https://docs.spring.io/spring-modulith/reference/verification.html"),
            ("Testcontainers PostgreSQL", "https://java.testcontainers.org/modules/databases/postgres/"),
        ]),
    ]
    for heading, links in groups:
        # Keep the reference appendix predictable in Word and LibreOffice.
        # The automatic split can pull the title area into the header/footer
        # when a long numbered-link group is treated as one pagination block.
        if heading == "대표 연구·해외":
            new_page(doc)
        doc.add_heading(heading, level=2)
        for label, url in links:
            p = doc.add_paragraph()
            apply_num(p, bullet_id)
            base.hyperlink(p, label, url, size=8.7)
    paragraph(doc, "본 문서는 공모전 출품·팀 공유를 위한 통합 기획·기능명세 자료이며 의료·법률·투자자문이 아니다. 법령·상품·사업상태는 변경될 수 있으므로 실제 사업화와 제출 직전 최신 원문을 다시 확인한다.", style="Small")


def audit(doc_path: Path):
    issues = []
    with ZipFile(doc_path) as zf:
        xml = zf.read("word/document.xml")
        root = ET.fromstring(xml)
        ns = {"w": W_NS}
        attr = lambda name: f"{{{W_NS}}}{name}"
        for idx, tbl in enumerate(root.findall(".//w:tbl", ns), 1):
            tbl_w = tbl.find("w:tblPr/w:tblW", ns)
            width = int(tbl_w.get(attr("w"), "0")) if tbl_w is not None else 0
            grid = [int(c.get(attr("w"), "0")) for c in tbl.findall("w:tblGrid/w:gridCol", ns)]
            if width <= 0 or sum(grid) != width:
                issues.append(f"table {idx} width/grid mismatch")
            for row in tbl.findall("w:tr", ns):
                cell_widths = []
                for tc in row.findall("w:tc", ns):
                    tcw = tc.find("w:tcPr/w:tcW", ns)
                    cell_widths.append(int(tcw.get(attr("w"), "0")) if tcw is not None else 0)
                if cell_widths != grid:
                    issues.append(f"table {idx} cell/grid mismatch")
                    break
        full_text = "".join(root.itertext())
        for forbidden in ("치매머니", "Dementia Money", "Dementia_Money"):
            if forbidden in full_text:
                issues.append(f"legacy service name remains: {forbidden}")
    if issues:
        raise RuntimeError("; ".join(issues))


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    patch_base_tokens()
    base.set_cell_margins = set_cell_margins
    base.set_table_geometry = set_table_geometry

    doc = Document()
    bullet_id, number_id = configure_styles(doc)
    set_header_footer(doc)
    doc.core_properties.title = "안심리듬 2026 금융 AI Challenge 출품 전략·기능명세·MVP 구현 통합본"
    doc.core_properties.subject = "금융생활 변화 조기알림 및 보호업무 코파일럿"
    doc.core_properties.author = "Ansim Rhythm Project Team"
    doc.core_properties.keywords = "안심리듬, 금융 AI, 개인 기준선, 이상변화 탐지, 금융소비자 보호, 행원 코파일럿"

    add_cover(doc)
    add_exec_summary(doc, bullet_id)
    add_competition(doc, bullet_id)
    add_problem(doc, bullet_id)
    add_evidence(doc)
    add_korea(doc, bullet_id)
    add_global(doc)
    add_positioning(doc, number_id)
    add_field_view(doc, bullet_id)
    add_user_journey(doc, bullet_id)
    add_demo(doc)
    add_detection(doc, bullet_id)
    add_comparison(doc)
    add_policy_and_consent(doc, bullet_id)
    add_functional_spec(doc)
    add_architecture(doc, bullet_id)
    add_llm_and_rag(doc, bullet_id)
    add_data_and_audit(doc)
    add_synthetic_data(doc, bullet_id)
    add_security(doc, bullet_id)
    add_business(doc, bullet_id)
    add_roadmap(doc)
    add_claims(doc)
    add_sources(doc, bullet_id)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(OUT)
    audit(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
